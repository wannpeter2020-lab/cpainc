"""
Hotel Planner Points — utility module.

Responsibilities:
  - detect_chain(hotel_name)            : map a hotel name to a chain key
  - fill_docx_template(blob, values)    : fill "Label: value" patterns in a .docx
  - build_field_values(...)             : assemble the dict passed to fill / display
  - build_mailto(...)                   : compose a mailto: URL
  - import_tracking_xlsx(file, chain)   : preview rows from a brand tracking workbook
  - default_chain_seeds()                : seven hotel programs with starter config

The .docx templates (Marriott, Hyatt) store fields inline inside a single
paragraph as "Label: value" so the fill strategy is regex replacement of the
trailing value after each known label.
"""

import io
import json
import re
import urllib.parse
from datetime import datetime, date

# ---------------------------------------------------------------------------
# Chain detection
# ---------------------------------------------------------------------------

CHAIN_KEYWORDS = [
    ('Marriott', [
        # Core brands
        'marriott', 'renaissance', 'sheraton', 'westin', 'courtyard',
        'fairfield', 'residence inn', 'springhill', 'aloft', 'ritz-carlton',
        'ritz carlton', 'jw ', 'jw marriott', 'element', 'ac hotel',
        'autograph', 'gaylord', 'le meridien', 'tribute', 'moxy',
        'protea', 'st. regis', 'st regis', 'w hotel', 'edition',
        'townplace', 'four points', 'delta hotel', 'delta guelph',
        'le westin', 'luxury collection', 'design hotels',
        # Common typos / space variants
        'resdence inn', 'residence in', 'spring hill', 'springhill suites',
        'westn', 'four pts', 'four point', 'fp ',
        # Specific Autograph Collection / Sheraton / Westin properties
        # whose names don't contain a recognisable brand keyword
        'hotel park city',     # Autograph Collection
        'ambassador hotel',    # Ambassador Hotel KC — Autograph Collection
        'the henry',           # The Henry — Autograph Collection
        'moana surfrider',     # Westin Moana Surfrider
        'royal hawaiian',      # Sheraton Royal Hawaiian (Luxury Collection)
        # Delta acquired by Marriott in 2015
        'delta ', 'delta hotels',
    ]),
    ('Hilton', [
        # Core brands
        'hilton', 'embassy suites', 'hampton', 'doubletree', 'curio',
        'tapestry', 'conrad', 'waldorf', 'tru ', 'home2', 'homewood',
        'lxr', 'canopy', 'motto', 'spark by hilton', 'tempo by hilton',
        # Common typos / space variants / abbreviations
        'double tree', 'dbl tree', 'doubetree', 'doubltree', 'dblt ',
        'hgi ', 'hilton garden', 'embassy ste', 'homewd suites',
        'home2 suites',
        # Curio Collection independents
        'st louis union station', 'st. louis union station',
        'union station hotel',
    ]),
    ('Hyatt', [
        'hyatt', 'andaz', 'grand hyatt', 'hyatt place', 'hyatt house',
        'thompson', 'park hyatt', 'caption by hyatt', 'unbound', 'alila',
        'destination by hyatt', 'joie de vivre', 'miraval', 'hyatt centric',
        'hyatt regency', 'hyatt ziva', 'hyatt zilara',
    ]),
    ('IHG', [
        'ihg', 'holiday inn', 'crowne plaza', 'staybridge', 'intercontinental',
        'kimpton', 'candlewood', 'avid', 'hotel indigo', 'voco', 'atwell',
        'even hotels', 'six senses', 'regent hotel', 'iberostar',
        'holiday inn express',
    ]),
    ('Omni', ['omni']),
    ('Choice', [
        'comfort inn', 'comfort suites', 'quality inn', 'sleep inn',
        'clarion', 'ascend', 'cambria', 'mainstay', 'suburban', 'econo lodge',
        'rodeway', 'choice hotel', 'woodspring',
    ]),
    ('Sonesta', [
        'sonesta', 'royal sonesta', 'sonesta es', 'sonesta select',
        'sonesta simply',
    ]),
]


def detect_chain(hotel_name):
    """Return the chain key (e.g. 'Marriott') for a hotel name, or None."""
    if not hotel_name:
        return None
    h = str(hotel_name).lower()
    # Strong-signal disambiguators first: "X by <Chain>" phrasing is canonical
    # and outranks generic substring matches (e.g. "DoubleTree by Hilton Hotel
    # Park City" should be Hilton, not Marriott via the Autograph "hotel park
    # city" keyword).
    if 'by hilton'   in h or 'doubletree' in h or 'embassy suites by' in h:
        return 'Hilton'
    if 'by marriott' in h or 'by sheraton' in h or 'by westin' in h:
        return 'Marriott'
    if 'by hyatt'    in h:
        return 'Hyatt'
    if 'by ihg'      in h or 'by intercontinental' in h:
        return 'IHG'
    for chain, keywords in CHAIN_KEYWORDS:
        for kw in keywords:
            if kw in h:
                return chain
    return None


# ---------------------------------------------------------------------------
# Field assembly
# ---------------------------------------------------------------------------

def _fmt_date(value, fmt='%-m/%-d/%Y'):
    """Render an ISO date / datetime / sqlite text as a friendly string."""
    if not value:
        return ''
    try:
        s = str(value)[:10]
        d = datetime.strptime(s, '%Y-%m-%d')
        try:
            return d.strftime(fmt)
        except ValueError:
            return d.strftime('%m/%d/%Y').lstrip('0').replace('/0', '/')
    except Exception:
        return str(value)


def _fmt_date_range(start, end):
    """Render '10/6 – 10/9, 2026' style for Marriott's Meeting Dates field."""
    if not start or not end:
        return ''
    try:
        s = datetime.strptime(str(start)[:10], '%Y-%m-%d')
        e = datetime.strptime(str(end)[:10], '%Y-%m-%d')
    except Exception:
        return f'{start} – {end}'
    if s.year == e.year:
        try:
            sm = s.strftime('%-m/%-d')
            em = e.strftime('%-m/%-d')
        except ValueError:
            sm = s.strftime('%m/%d').lstrip('0').replace('/0', '/')
            em = e.strftime('%m/%d').lstrip('0').replace('/0', '/')
        return f'{sm} – {em}, {s.year}'
    return f'{_fmt_date(start)} – {_fmt_date(end)}'


def build_field_values(user_profile, program, request_row, pickup_config, booking_row):
    """
    Build the dict of data points used both for .docx filling and the
    weblink-helper side panel.

    All values are stringified so callers don't need to worry about None.
    """
    today = date.today().isoformat()
    fv = {}

    # User profile
    fv['user_full_name'] = (user_profile or {}).get('user_full_name', '') or ''
    fv['user_email']     = (user_profile or {}).get('user_email', '') or ''
    fv['user_phone']     = (user_profile or {}).get('user_phone', '') or ''
    fv['member_number']  = (program or {}).get('member_number', '') or ''
    fv['chain_name']     = (program or {}).get('chain_name', '') or ''
    fv['today']          = _fmt_date(today)
    fv['today_iso']      = today

    # Pickup config — hotel + event info
    if pickup_config:
        fv['hotel']                 = pickup_config.get('hotel') or ''
        fv['hotel_contact_name']    = pickup_config.get('hotel_contact') or ''
        fv['hotel_contact_email']   = pickup_config.get('hotel_contact_email') or ''
        fv['event_name']            = (pickup_config.get('event_name')
                                       or pickup_config.get('organization') or '')
        fv['client_organization']   = pickup_config.get('organization') or ''
        ev_start = pickup_config.get('event_start') or ''
        ev_end   = pickup_config.get('event_end') or ''
        # Fall back to first/last contracted block night if event_start/end blank
        try:
            block = json.loads(pickup_config.get('contracted_block') or '{}')
            keys = sorted(block.keys())
            if not ev_start and keys: ev_start = keys[0]
            if not ev_end and keys:   ev_end   = keys[-1]
        except Exception:
            pass
        fv['event_start']           = _fmt_date(ev_start)
        fv['event_end']             = _fmt_date(ev_end)
        fv['meeting_dates_formatted'] = _fmt_date_range(ev_start, ev_end)
        fv['contracted_rate']       = (str(pickup_config.get('contracted_rate'))
                                       if pickup_config.get('contracted_rate') else '')
        # Sum of contracted block = total contracted rooms
        try:
            total_rooms = sum(int(v or 0) for v in
                              json.loads(pickup_config.get('contracted_block') or '{}').values())
            fv['total_contracted_rooms'] = str(total_rooms) if total_rooms else ''
        except Exception:
            fv['total_contracted_rooms'] = ''

    # Booking row (ReportPipeline) — supplies Cvent code, peak rooms, etc.
    # Promagent uses spaced column names ("Booking Id"); CPAinc uses unspaced
    # ("BookingId"). Try both.
    if booking_row:
        if hasattr(booking_row, 'get'):
            _bk_get_raw = booking_row.get
        else:
            _keys = booking_row.keys() if hasattr(booking_row, 'keys') else []
            _bk_get_raw = lambda k: booking_row[k] if k in _keys else None

        def bk_get(spaced, unspaced=None):
            v = _bk_get_raw(spaced)
            if v is not None:
                return v
            if unspaced:
                return _bk_get_raw(unspaced)
            return None

        fv['booking_id']        = str(bk_get('Booking Id', 'BookingId') or '')
        fv['peak_rooms']        = str(bk_get('Peak Rooms', 'PeakRooms') or '') if bk_get('Peak Rooms', 'PeakRooms') else ''
        fv['total_room_nights'] = str(bk_get('Total Room Nights', 'TotalRoomNights') or '') if bk_get('Total Room Nights', 'TotalRoomNights') else ''
        fv['revenue']           = str(bk_get('Revenue') or '') if bk_get('Revenue') else ''
        fv['account_name']      = bk_get('Account Name', 'AccountName') or ''
        # Override event/organization if booking has fuller names
        if not fv.get('event_name'):
            fv['event_name']    = bk_get('Booking Name', 'BookingName') or bk_get('Event Name', 'EventName') or ''
        if not fv.get('client_organization'):
            fv['client_organization'] = bk_get('Account Name', 'AccountName') or ''

    # Request-row overrides — user-edited values take precedence
    if request_row:
        for k in ('cvent_rfp_code', 'contract_signature_date',
                  'incentive_type', 'award_timing',
                  'second_recipient_name', 'second_recipient_email',
                  'second_recipient_phone', 'second_recipient_number'):
            v = request_row.get(k) if hasattr(request_row, 'get') else (
                request_row[k] if k in request_row.keys() else None)
            if v:
                fv[k] = str(v) if not k.endswith('_date') else _fmt_date(v)
            else:
                fv.setdefault(k, '')

    fv.setdefault('cvent_rfp_code', '')
    fv.setdefault('contract_signature_date', '')

    # Pretty-string of total dollar commitment for Hyatt's "Room Rate/F&B"
    if fv.get('contracted_rate'):
        fv.setdefault('room_rate_fnb', f'${fv["contracted_rate"]}')
    else:
        fv.setdefault('room_rate_fnb', '')

    return fv


# ---------------------------------------------------------------------------
# .docx auto-fill
# ---------------------------------------------------------------------------

def _replace_paragraph_text(para, new_text):
    """
    Replace a paragraph's text while preserving the run formatting of the
    first run. Subsequent runs are cleared.
    """
    if not para.runs:
        para.text = new_text
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ''


SECONDARY_SECTION_MARKERS = (
    '2nd recipient', '2nd  recipient', 'second recipient',
    'additional recipient', 'additional member',
)

PRIMARY_KEYS = (
    'user_full_name', 'user_email', 'user_phone',
    'member_number',
)

SECONDARY_KEYS = (
    'second_recipient_name', 'second_recipient_email',
    'second_recipient_phone', 'second_recipient_number',
)


def fill_docx_template(template_blob, field_values, field_mapping):
    """
    Fill a Word template where each field is rendered as
        "Label: <value>"
    inside a table cell or paragraph.

    Section handling: Marriott's form has a "2nd Recipient (if splitting
    points)" block lower in the doc with labels identical to the primary
    block. Walking in document order, once we encounter that marker we
    stop filling primary-recipient keys (user_*, member_number) and only
    fill secondary-recipient keys (second_recipient_*). The hotel/event
    fields after the 2nd-recipient block (Property Name, etc.) continue
    to fill normally.

    Args:
      template_blob       : bytes — the blank .docx template
      field_values        : dict  — data dictionary from build_field_values()
      field_mapping       : dict  — {label_in_doc: data_key_in_field_values}

    Returns: bytes (filled .docx)
    """
    from docx import Document  # python-docx, already used elsewhere

    doc = Document(io.BytesIO(template_blob))

    # Compile regex patterns once
    patterns = []
    for label, data_key in (field_mapping or {}).items():
        if not label:
            continue
        pat = re.compile(rf'({re.escape(label)}\s*:)\s*[^\r\n]*', re.IGNORECASE)
        patterns.append((label, pat, data_key))

    state = {'in_secondary': False}

    def _section_check(text):
        low = (text or '').lower()
        for marker in SECONDARY_SECTION_MARKERS:
            if marker in low:
                state['in_secondary'] = True
                return

    def _apply(text):
        if not text:
            return text
        _section_check(text)
        for label, pat, key in patterns:
            if state['in_secondary']:
                # In the 2nd recipient block, skip primary-recipient keys
                if key in PRIMARY_KEYS:
                    continue
                # Re-route primary labels to secondary keys (e.g. on Marriott
                # the second Member Name slot maps to second_recipient_name).
                if key in PRIMARY_KEYS:
                    continue
            else:
                # In primary block, don't fill secondary slots if any
                if key in SECONDARY_KEYS:
                    continue
            value = str(field_values.get(key, '') or '')
            if pat.search(text):
                text = pat.sub(rf'\1 {value}', text)
        return text

    # Walk every paragraph in document body and every cell in every table
    for para in doc.paragraphs:
        new_text = _apply(para.text)
        if new_text != para.text:
            _replace_paragraph_text(para, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    new_text = _apply(para.text)
                    if new_text != para.text:
                        _replace_paragraph_text(para, new_text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# mailto: builder
# ---------------------------------------------------------------------------

def build_mailto(to_email, subject, body, cc=None):
    """
    Build a mailto: URL string with prefilled subject and body.
    Note: attachments cannot be added via mailto; user attaches manually.
    """
    params = {}
    if subject: params['subject'] = subject
    if body:    params['body']    = body
    if cc:      params['cc']      = cc if isinstance(cc, str) else ','.join(cc)
    qs = urllib.parse.urlencode(params, quote_via=urllib.parse.quote)
    return f'mailto:{to_email or ""}?{qs}'


# ---------------------------------------------------------------------------
# Default chain seeds (used by ensure_pickup_tables on first run)
# ---------------------------------------------------------------------------

DEFAULT_MARRIOTT_MAPPING = {
    'Member Name (First and Last)':     'user_full_name',
    'Bonvoy Account Number':            'member_number',
    'E-mail Address':                   'user_email',
    'Business Phone':                   'user_phone',
    'Date Submitted to Property':       'today',
    'Property Name':                    'hotel',
    'End User Account Name':            'client_organization',
    'Property Sales Contact':           'hotel_contact_name',
    'Cvent RFP Code':                   'cvent_rfp_code',
    'Meeting Dates (day/month/year)':   'meeting_dates_formatted',
    'Meeting Name':                     'event_name',
}

DEFAULT_HYATT_MAPPING = {
    'Property Name':                    'hotel',
    'Total # Contracted Rooms':         'total_contracted_rooms',
    'Event Name':                       'event_name',
    'Room Rate/F&B Commitment':         'room_rate_fnb',
    'Event Start Date':                 'event_start',
    'Event End Date':                   'event_end',
    'Hotel Sales Contact':              'hotel_contact_name',
    'Contract Signature Date':          'contract_signature_date',
    'FIRST NAME':                       'user_full_name',
    'LAST NAME':                        'user_full_name',
    'WORLD OF HYATT ACCOUNT NUMBER':    'member_number',
    'EMAIL ADDRESS':                    'user_email',
    'PHONE NUMBER':                     'user_phone',
}

DEFAULT_HILTON_MAPPING = {
    'Hilton Honors Number':             'member_number',
    'First Name':                       'user_full_name',
    'Last Name':                        'user_full_name',
    'Email':                            'user_email',
    'Phone':                            'user_phone',
    'Hotel Name':                       'hotel',
    'Event Name':                       'event_name',
    'Account / Organization':           'client_organization',
    'Event Start Date':                 'event_start',
    'Event End Date':                   'event_end',
    'Peak Room Nights':                 'peak_rooms',
    'Total Room Nights':                'total_room_nights',
    'Hotel Sales Contact':              'hotel_contact_name',
}

DEFAULT_IHG_MAPPING = {
    'IHG One Rewards Number':           'member_number',
    'Member Name':                      'user_full_name',
    'Email':                            'user_email',
    'Phone':                            'user_phone',
    'Hotel Name':                       'hotel',
    'Event Name':                       'event_name',
    'Cvent ID':                         'cvent_rfp_code',
    'Account / Organization':           'client_organization',
    'Event Start Date':                 'event_start',
    'Event End Date':                   'event_end',
    'Total Room Nights':                'total_room_nights',
    'Date Contract Signed':             'contract_signature_date',
}

DEFAULT_OMNI_MAPPING = {
    'Omni Member Number':               'member_number',
    'Member Name':                      'user_full_name',
    'Email':                            'user_email',
    'Phone':                            'user_phone',
    'Hotel Name':                       'hotel',
    'Event Name':                       'event_name',
    'Event Start Date':                 'event_start',
    'Event End Date':                   'event_end',
    'Total Room Nights':                'total_room_nights',
}


def default_chain_seeds():
    """Seven default chain rows to insert on first run."""
    return [
        {
            'chain_name': 'Marriott',
            'submission_type': 'docx_template',
            'form_url': None,
            'field_mapping_json': json.dumps(DEFAULT_MARRIOTT_MAPPING),
            'submission_window_days': 90,
            'receipt_window_days': 60,
            'notes': 'Bonvoy Rewarding Events. Submit within 90 days of event.',
        },
        {
            'chain_name': 'Hilton',
            'submission_type': 'weblink',
            'form_url': 'https://hiltonsales.my.site.com/HWS/s/3p-request?language=en_US',
            'field_mapping_json': json.dumps(DEFAULT_HILTON_MAPPING),
            'submission_window_days': 90,
            'receipt_window_days': 60,
            'notes': 'Hilton Honors — submit online via the Salesforce form.',
        },
        {
            'chain_name': 'Hyatt',
            'submission_type': 'docx_template',
            'form_url': None,
            'field_mapping_json': json.dumps(DEFAULT_HYATT_MAPPING),
            'submission_window_days': 30,
            'receipt_window_days': 60,
            'notes': 'World of Hyatt — submit within 30 days of contract signature.',
        },
        {
            'chain_name': 'IHG',
            'submission_type': 'weblink',
            'form_url': ('https://forms.office.com/pages/responsepage.aspx?id='
                         'P8RiJ8EpikSJ9nrJA8-DFlDq2Hu6U1RIsb0yLBjhOfhUNE9TNTVDSzBBRklMQU1ZVldHREI1RldQNC4u'
                         '&route=shorturl'),
            'field_mapping_json': json.dumps(DEFAULT_IHG_MAPPING),
            'submission_window_days': 60,
            'receipt_window_days': 60,
            'notes': 'IHG One Rewards — submit online via Microsoft Forms.',
        },
        {
            'chain_name': 'Omni',
            'submission_type': 'manual',
            'form_url': None,
            'field_mapping_json': json.dumps(DEFAULT_OMNI_MAPPING),
            'submission_window_days': 90,
            'receipt_window_days': 60,
            'notes': 'Omni Select Guest — manual email to hotel salesperson.',
        },
        {
            'chain_name': 'Choice',
            'submission_type': 'manual',
            'form_url': None,
            'field_mapping_json': json.dumps({}),
            'submission_window_days': 90,
            'receipt_window_days': 60,
            'notes': 'Choice Privileges — manual process.',
        },
        {
            'chain_name': 'Sonesta',
            'submission_type': 'manual',
            'form_url': None,
            'field_mapping_json': json.dumps({}),
            'submission_window_days': 90,
            'receipt_window_days': 60,
            'notes': 'Sonesta Travel Pass — manual process.',
        },
    ]


# ---------------------------------------------------------------------------
# Tracking xlsx import (historical data)
# ---------------------------------------------------------------------------

def _safe_int(value):
    if value is None:
        return None
    try:
        s = str(value).strip().replace(',', '')
        if s == '' or s.lower() in ('nan', 'none'):
            return None
        return int(float(s))
    except Exception:
        return None


def _parse_excel_date(value):
    """Normalise to ISO date string, or None."""
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    s = str(value).strip()
    if not s or s.lower() in ('nan', 'nat', 'none'):
        return None
    for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m/%d/%y', '%d/%m/%Y',
                '%Y-%m-%d %H:%M:%S', '%m/%d/%Y %H:%M:%S'):
        try:
            return datetime.strptime(s, fmt).strftime('%Y-%m-%d')
        except Exception:
            pass
    return None


def _normalise_status(date_or_status_cell, points_value):
    """
    'Date Points Posted' is mixed: date OR status string (Cancelled, Disallowed,
    Not Qualified, etc.). Return (iso_date_or_None, status).
    """
    if date_or_status_cell is None:
        return None, ('received' if points_value else 'submitted')
    iso = _parse_excel_date(date_or_status_cell)
    if iso:
        return iso, 'received' if points_value else 'submitted'
    s = str(date_or_status_cell).strip().lower()
    if not s:
        return None, ('received' if points_value else 'submitted')
    if 'cancel' in s:
        return None, 'cancelled'
    if 'disallow' in s or 'past 90' in s or 'past submission' in s:
        return None, 'disallowed'
    if 'not qualif' in s or 'not in cvent' in s:
        return None, 'disallowed'
    return None, 'submitted'


def import_tracking_xlsx(file_bytes, chain_name):
    """
    Parse a brand-tracking workbook into a flat list of dicts.
    Each dict is a candidate hotel_points_request row.

    Returns: list of dicts with keys —
       event_name, hotel, start_date (iso), end_date (iso),
       form_sent_date (iso), points_received_date (iso),
       points_awarded (int|None), status, rewards_form_link, notes
    """
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    rows_out = []

    for sheet_name in wb.sheetnames:
        if sheet_name in ('Sheet1', 'Marriott_Starwood'):
            continue
        # Heuristic — skip sheets whose names aren't 4-digit years
        if not re.fullmatch(r'\d{4}', sheet_name.strip()):
            # Allow "2024 ", "2024-Final", etc.
            if not re.match(r'^\d{4}', sheet_name.strip()):
                continue

        ws = wb[sheet_name]
        # Find header row — look in first 5 rows for "EVENT NAME" or "Event Name"
        header_row_idx = None
        for r in range(1, min(6, ws.max_row + 1)):
            cells = [str(c.value or '').strip() for c in ws[r]]
            if any(c.lower() == 'event name' for c in cells):
                header_row_idx = r
                break
        if not header_row_idx:
            continue
        headers = [str(c.value or '').strip().lower() for c in ws[header_row_idx]]

        def col_idx(*names):
            for n in names:
                if n.lower() in headers:
                    return headers.index(n.lower())
            return None

        i_event   = col_idx('event name')
        i_start   = col_idx('start date', 'event start date')
        i_end     = col_idx('end date', 'event end date')
        i_hotel   = col_idx('hotel')
        i_sent    = col_idx('form sent', 'date form sent')
        i_posted  = col_idx('date points posted', 'points posted')
        i_points  = col_idx('points')
        i_form    = col_idx('rewards form')
        i_ct1     = col_idx('contract')

        for r in range(header_row_idx + 1, ws.max_row + 1):
            row = [c.value for c in ws[r]]
            def _cell(i):
                return row[i] if (i is not None and i < len(row)) else None
            event_name = _cell(i_event)
            if not event_name:
                continue
            ev = str(event_name).strip()
            if not ev or ev.lower() == 'event name':
                continue
            points_val = _safe_int(_cell(i_points))
            posted_iso, status = _normalise_status(_cell(i_posted), points_val)

            # Compose notes from any extra free-text columns
            extras = []
            for h_lower, val in zip(headers, row):
                if val is None or h_lower in ('', 'event name', 'start date', 'end date',
                                              'hotel', 'brand', 'form sent',
                                              'date points posted', 'points',
                                              'rewards form', 'contract'):
                    continue
                s = str(val).strip()
                if s and s.lower() not in ('nan', 'none'):
                    extras.append(f'{h_lower}: {s}')

            rows_out.append({
                'sheet': sheet_name,
                'chain': chain_name,
                'event_name': ev,
                'hotel': str(_cell(i_hotel) or '').strip(),
                'start_date': _parse_excel_date(_cell(i_start)),
                'end_date':   _parse_excel_date(_cell(i_end)),
                'form_sent_date':       _parse_excel_date(_cell(i_sent)),
                'points_received_date': posted_iso,
                'points_awarded':       points_val,
                'status':               status,
                'rewards_form_link':    (str(_cell(i_form) or '').strip()
                                         if _cell(i_form) else ''),
                'notes':                ' | '.join(extras)[:2000],
            })

    return rows_out
