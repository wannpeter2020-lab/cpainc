"""
pickup_utils.py — Utility functions for the Pickup Tracking module.

Functions:
  import_pickup_tab(ws, tab_name)         Parse one Excel tab → config + weekly entries
  parse_rooming_list_pdf(file_bytes)      Extract guests from PDF → nights_by_date
  reconcile(pickup_by_night, rl_nights)   Compare pickup vs rooming list
  format_export_row(config, weekly)       Tab-separated row for Excel paste-back
  hotel_request_mailto(config)            Pre-filled mailto: for hotel pickup request
  client_summary_mailto(config, weekly, rl_status)  mailto: for client summary
"""

import re
import json
import io
from datetime import datetime, timedelta
from urllib.parse import quote


# ── Spreadsheet importer ──────────────────────────────────────────────────────

def _shared_helpers():
    """Return (sv, to_iso) helpers — defined once, reused by both parsers."""
    def sv(v):
        return str(v).strip() if v is not None else ''

    def to_iso(v):
        if v is None:
            return None
        if hasattr(v, 'strftime'):
            return v.strftime('%Y-%m-%d')
        s = str(v).strip()
        if re.match(r'^\d{4}-\d{2}-\d{2}', s):
            return s[:10]
        if re.match(r'^\d{2}/\d{2}/\d{2}$', s):
            m, d, y = s.split('/')
            yr = 2000 + int(y) if int(y) < 50 else 1900 + int(y)
            return f'{yr}-{m}-{d}'
        return None

    return sv, to_iso


def _parse_section(rows, tab_name):
    """
    Parse a list of rows that represents one hotel's pickup table.
    rows[0] must be the ORGANIZATION: row (col A = 'ORGANIZATION:' or similar).

    Layout (0-indexed within the slice):
      0  ORGANIZATION:      | org_name
      1  HOTEL & LOCATION:  | hotel_name
      2  NAME & DATE …      | event_name
      3  GROUP CONTACT      | gc_name  [col C may hold gc email]
      4  HOTEL CONTACT:     | hc_name_or_email
      5  CONTACT'S NUMBER:  | phone_or_email
      6  Booking            | booking_id
      7  Cut-off Date:      | cutoff  …  Attrition: | pct
      8  Dates:             | d1 | d2 | …
      9  Day:               | Mon | Tue | …
     10  Block:             | b1 | b2 | …
     11+ weekly rows
    """
    sv, to_iso = _shared_helpers()

    def cell(r, c):
        try:
            return rows[r][c]
        except (IndexError, TypeError):
            return None

    def _int(v):
        try:
            return int(v) if v not in (None, '', '#DIV/0!', 'NA', 'N/A') else None
        except Exception:
            return None

    def _flt(v):
        try:
            return float(v) if v not in (None, '', '#DIV/0!', 'NA', 'N/A') else None
        except Exception:
            return None

    organization = sv(cell(0, 1))
    hotel        = sv(cell(1, 1))
    event_name   = sv(cell(2, 1))

    # Group contact: col B; email sometimes in col C
    group_contact = sv(cell(3, 1))
    gc_col_c = sv(cell(3, 2))
    group_contact_email = gc_col_c if '@' in gc_col_c else ''

    # Hotel contact: name in col B; email may be directly in col B, or in row 5
    hc_raw = sv(cell(4, 1))
    hotel_contact_email = ''
    em = re.search(r'<([^>]+@[^>]+)>', hc_raw)
    if em:
        hotel_contact_email = em.group(1)
        hotel_contact = hc_raw[:hc_raw.find('<')].strip()
    elif '@' in hc_raw:
        hotel_contact_email = hc_raw
        hotel_contact = ''
    else:
        hotel_contact = hc_raw
    # Fallback: email in row 5 (CONTACT'S NUMBER)
    if not hotel_contact_email:
        row5 = sv(cell(5, 1))
        if '@' in row5:
            hotel_contact_email = row5

    # Booking ID
    bid_raw = cell(6, 1)
    try:
        booking_id = str(int(float(str(bid_raw).strip()))) if bid_raw else None
    except Exception:
        booking_id = sv(bid_raw) or None

    # Cut-off date and attrition (row 7)
    cutoff_date = to_iso(cell(7, 1))
    attrition_pct = None
    if len(rows) > 7 and rows[7]:
        for v in rows[7]:
            if isinstance(v, (int, float)) and 0 < v <= 1:
                attrition_pct = float(v)
                break

    # Event dates (row 8, cols B+)
    event_dates = []
    if len(rows) > 8 and rows[8]:
        for v in (rows[8][1:] if len(rows[8]) > 1 else []):
            d = to_iso(v)
            if d:
                event_dates.append(d)
            elif v is None and event_dates:
                break
    n_nights = len(event_dates)

    # Contracted block (row 10, cols B+)
    contracted_block = {}
    if len(rows) > 10 and n_nights:
        for i, date in enumerate(event_dates):
            v = cell(10, i + 1)
            try:
                contracted_block[date] = int(v) if v not in (None, '', 'N/A') else 0
            except Exception:
                contracted_block[date] = 0

    # Weekly entries (row 11+)
    _SKIP = {'remaining', 'pick-up update', 'balance',
             'revised block', 'day', 'dates', 'block', 'organization',
             'hotel & location', 'name & date of event', 'group contact',
             'hotel contact', "contact's number", 'booking', 'cut-off date',
             'attrition'}
    LABEL_PREFIXES = ('final history', 'final pickup', 'pending history',
                      'pending pickup', 'revised')
    weekly_entries = []
    rows_list = list(rows)   # ensure indexable for .index()

    for ri, row in enumerate(rows_list[11:], start=11):
        if not row:
            continue
        col_a = row[0]
        if col_a is None:
            continue

        label = None
        report_date = None

        if isinstance(col_a, str):
            stripped = col_a.strip()
            stripped_lo = stripped.lower()
            if not stripped or stripped_lo in _SKIP:
                continue

            # Check for label prefix rows (Final History, Pending History, etc.)
            matched_prefix = None
            for pfx in LABEL_PREFIXES:
                if stripped_lo.startswith(pfx):
                    matched_prefix = pfx
                    break

            if matched_prefix is not None:
                label = stripped  # preserve original casing
                # Try to extract a date from remainder of text or col B
                remainder = stripped[len(matched_prefix):].strip().strip(':-')
                report_date = to_iso(remainder) or to_iso(cell(ri, 1))
                # For label rows without a parseable date, use last event date
                if not report_date and event_dates:
                    report_date = event_dates[-1]
                if not report_date:
                    continue
            else:
                report_date = to_iso(stripped)
                if not report_date:
                    continue
        elif hasattr(col_a, 'strftime'):
            report_date = col_a.strftime('%Y-%m-%d')
        else:
            report_date = to_iso(col_a)
            if not report_date:
                continue

        if not report_date:
            continue

        pickup_by_night = {}
        if n_nights:
            for i, date in enumerate(event_dates):
                v = row[i + 1] if (i + 1) < len(row) else None
                try:
                    pickup_by_night[date] = int(v) if v not in (None, '') else None
                except Exception:
                    pickup_by_night[date] = None

        col_total   = n_nights + 1
        col_change  = n_nights + 2
        col_pct_blk = n_nights + 3
        col_pct_att = n_nights + 4
        col_ota     = n_nights + 5

        total_rooms          = _int(row[col_total]   if col_total   < len(row) else None)
        change_from_last     = _int(row[col_change]  if col_change  < len(row) else None)
        pct_of_block_raw     = _flt(row[col_pct_blk] if col_pct_blk < len(row) else None)
        pct_of_attrition_raw = _flt(row[col_pct_att] if col_pct_att < len(row) else None)
        pct_of_block     = round(pct_of_block_raw * 100, 1) if (pct_of_block_raw is not None and pct_of_block_raw <= 1.5) else pct_of_block_raw
        pct_of_attrition = round(pct_of_attrition_raw * 100, 1) if (pct_of_attrition_raw is not None and pct_of_attrition_raw <= 1.5) else pct_of_attrition_raw
        ota_rate         = _flt(row[col_ota]     if col_ota     < len(row) else None)

        if not pickup_by_night and total_rooms is None:
            continue

        weekly_entries.append({
            'report_date':      report_date,
            'pickup_by_night':  json.dumps(pickup_by_night),
            'total_rooms':      total_rooms,
            'change_from_last': change_from_last,
            'pct_of_block':     pct_of_block,
            'pct_of_attrition': pct_of_attrition,
            'ota_rate':         ota_rate,
            'label':            label,
        })

    return {
        'config': {
            'tab_name':             tab_name,
            'organization':         organization,
            'event_name':           event_name,
            'hotel':                hotel,
            'hotel_contact':        hotel_contact,
            'hotel_contact_email':  hotel_contact_email,
            'group_contact':        group_contact,
            'group_contact_email':  group_contact_email,
            'booking_id':           booking_id,
            'cutoff_date':          cutoff_date,
            'attrition_pct':        attrition_pct,
            'contracted_block':     contracted_block,
        },
        'weekly_entries': weekly_entries,
    }


def import_pickup_tab_multi(ws, tab_name):
    """
    Parse one openpyxl worksheet that may contain multiple hotel tables.

    Each table starts with a row where col A = 'ORGANIZATION:'.
    Returns a list of dicts: [{config, weekly_entries}, ...] — one per table.
    Skips summary/rollup tables that have no contracted_block and no weekly entries.
    """
    sv, to_iso = _shared_helpers()
    all_rows = list(ws.iter_rows(values_only=True))

    # Find every row where col A = 'ORGANIZATION:' (case-insensitive, strip colon)
    org_indices = [
        i for i, row in enumerate(all_rows)
        if row and sv(row[0]).upper().rstrip(':').strip() == 'ORGANIZATION'
    ]

    if not org_indices:
        return []

    # Split at each ORGANIZATION: boundary
    sections = []
    for n, start in enumerate(org_indices):
        end = org_indices[n + 1] if n + 1 < len(org_indices) else len(all_rows)
        sections.append(all_rows[start:end])

    results = []
    for n, section_rows in enumerate(sections):
        result = _parse_section(section_rows, tab_name)
        if result is None:
            continue
        cfg = result['config']
        # Skip pure rollup sections: no booking_id AND no weekly entries AND
        # hotel name looks like a summary ("All Hotels", "Total", etc.)
        hotel_lo = (cfg.get('hotel') or '').lower()
        is_rollup = (not cfg.get('booking_id')
                     and not result['weekly_entries']
                     and (not cfg.get('contracted_block') or
                          hotel_lo in ('all hotels', 'all', 'total', 'totals', '')))
        if is_rollup:
            continue
        if not cfg.get('organization'):
            continue
        # Tag with hotel suffix when multiple tables exist
        if len(sections) > 1:
            hotel_tag = cfg.get('hotel') or f'Table {n+1}'
            cfg['tab_name'] = f"{tab_name} [{hotel_tag}]"
        results.append(result)

    return results


def import_pickup_tab(ws, tab_name):
    """
    Parse one openpyxl worksheet tab from the pickup spreadsheet.

    Returns a dict:
      {
        'config': {organization, event_name, hotel, hotel_contact,
                   hotel_contact_email, group_contact, booking_id,
                   cutoff_date, attrition_pct, contracted_block (JSON str)},
        'weekly_entries': [{report_date, pickup_by_night (JSON str),
                            total_rooms, change_from_last, pct_of_block,
                            pct_of_attrition, ota_rate, label}]
      }
    """
    rows = list(ws.iter_rows(values_only=True))

    def cell(r, c):
        """0-indexed row, 0-indexed col."""
        try:
            return rows[r][c]
        except IndexError:
            return None

    def sv(v):
        """Safe string — strip whitespace, return '' for None."""
        return str(v).strip() if v is not None else ''

    def to_iso(v):
        """Convert datetime / 'YYYY-MM-DD HH:MM:SS' string / None to ISO date string."""
        if v is None:
            return None
        if hasattr(v, 'strftime'):
            return v.strftime('%Y-%m-%d')
        s = str(v).strip()
        # Accept 'YYYY-MM-DD ...' or 'MM/DD/YY'
        if re.match(r'^\d{4}-\d{2}-\d{2}', s):
            return s[:10]
        if re.match(r'^\d{2}/\d{2}/\d{2}$', s):
            m, d, y = s.split('/')
            yr = 2000 + int(y) if int(y) < 50 else 1900 + int(y)
            return f'{yr}-{m}-{d}'
        return None

    # ── Header fields (rows 0-11, 0-indexed) ─────────────────────────────────
    organization  = sv(cell(1, 1))
    hotel         = sv(cell(2, 1))
    event_name    = sv(cell(3, 1))

    # Group contact may span cols B + C
    gc = ' '.join(sv(cell(4, c)) for c in range(1, 3)).strip()
    group_contact = gc

    # Hotel contact: may contain email in angle brackets
    hc_raw = sv(cell(5, 1))
    hotel_contact_email = ''
    em = re.search(r'<([^>]+@[^>]+)>', hc_raw)
    if em:
        hotel_contact_email = em.group(1)
        hotel_contact = hc_raw[:hc_raw.find('<')].strip()
    else:
        hotel_contact = hc_raw

    # Booking ID
    bid_raw = cell(7, 1)
    try:
        booking_id = str(int(float(str(bid_raw).strip()))) if bid_raw else None
    except Exception:
        booking_id = sv(bid_raw) or None

    # Cut-off date (row 8 col B)
    cutoff_date = to_iso(cell(8, 1))

    # Attrition % — find the first float between 0 and 1 in row 8
    attrition_pct = None
    if len(rows) > 8:
        for v in rows[8]:
            if isinstance(v, (int, float)) and 0 < v < 1:
                attrition_pct = float(v)
                break

    # Event dates (row 9, cols B onwards until None)
    event_dates = []
    if len(rows) > 9:
        for v in rows[9][1:]:
            d = to_iso(v)
            if d:
                event_dates.append(d)
            elif v is None and event_dates:
                break   # stop at first gap after dates started

    n_nights = len(event_dates)

    # Contracted block (row 11, cols B .. B+n_nights-1)
    contracted_block = {}
    if len(rows) > 11 and n_nights:
        for i, date in enumerate(event_dates):
            v = cell(11, i + 1)
            try:
                contracted_block[date] = int(v) if v not in (None, '', 'N/A') else 0
            except Exception:
                contracted_block[date] = 0

    # ── Weekly pickup rows (row 12 onwards) ──────────────────────────────────
    _SKIP_LABELS = {'Remaining', 'Final History', 'PICK-UP UPDATE',
                    'remaining', 'final history'}
    weekly_entries = []

    for row in rows[12:]:
        if not row:
            continue
        col_a = row[0]
        if col_a is None:
            continue

        # Determine label and report_date
        label = None
        report_date = None

        if isinstance(col_a, str):
            stripped = col_a.strip()
            if stripped in _SKIP_LABELS or not stripped:
                continue
            if stripped.lower() == 'final history':
                label = 'Final History'
                # date may be in col B for 'Final History'
                report_date = to_iso(cell(rows.index(row), 1)) if hasattr(rows, 'index') else None
                if not report_date:
                    continue
            else:
                report_date = to_iso(stripped)
                if not report_date:
                    continue
        elif hasattr(col_a, 'strftime'):
            report_date = col_a.strftime('%Y-%m-%d')
        else:
            report_date = to_iso(col_a)
            if not report_date:
                continue

        if not report_date:
            continue

        # Pickup per night (cols 1 .. n_nights)
        pickup_by_night = {}
        if n_nights:
            for i, date in enumerate(event_dates):
                v = row[i + 1] if (i + 1) < len(row) else None
                try:
                    pickup_by_night[date] = int(v) if v not in (None, '') else None
                except Exception:
                    pickup_by_night[date] = None

        def _int(v):
            try:
                return int(v) if v not in (None, '', '#DIV/0!', 'NA', 'N/A') else None
            except Exception:
                return None

        def _flt(v):
            try:
                return float(v) if v not in (None, '', '#DIV/0!', 'NA', 'N/A') else None
            except Exception:
                return None

        col_total    = n_nights + 1
        col_change   = n_nights + 2
        col_pct_blk  = n_nights + 3
        col_pct_att  = n_nights + 4
        col_ota      = n_nights + 5

        total_rooms      = _int(row[col_total]   if col_total   < len(row) else None)
        change_from_last = _int(row[col_change]  if col_change  < len(row) else None)
        pct_of_block_raw     = _flt(row[col_pct_blk] if col_pct_blk < len(row) else None)
        pct_of_attrition_raw = _flt(row[col_pct_att] if col_pct_att < len(row) else None)
        # Normalize: Excel stores % as decimal (0.76); we keep 0-100 range in DB
        pct_of_block     = round(pct_of_block_raw * 100, 1) if pct_of_block_raw is not None and pct_of_block_raw <= 1.5 else pct_of_block_raw
        pct_of_attrition = round(pct_of_attrition_raw * 100, 1) if pct_of_attrition_raw is not None and pct_of_attrition_raw <= 1.5 else pct_of_attrition_raw
        ota_rate         = _flt(row[col_ota]     if col_ota     < len(row) else None)

        if not pickup_by_night and total_rooms is None:
            continue  # empty row

        weekly_entries.append({
            'report_date':      report_date,
            'pickup_by_night':  json.dumps(pickup_by_night),
            'total_rooms':      total_rooms,
            'change_from_last': change_from_last,
            'pct_of_block':     pct_of_block,
            'pct_of_attrition': pct_of_attrition,
            'ota_rate':         ota_rate,
            'label':            label,
        })

    return {
        'config': {
            'tab_name':             tab_name,
            'organization':         organization,
            'event_name':           event_name,
            'hotel':                hotel,
            'hotel_contact':        hotel_contact,
            'hotel_contact_email':  hotel_contact_email,
            'group_contact':        group_contact,
            'booking_id':           booking_id,
            'cutoff_date':          cutoff_date,
            'attrition_pct':        attrition_pct,
            'contracted_block':     contracted_block,   # dict, not JSON string
        },
        'weekly_entries': weekly_entries,
    }


# ── PDF rooming list parser ───────────────────────────────────────────────────

# Matches a guest data line: Name(s) + 8-digit conf# + arr MM/DD/YY + dep MM/DD/YY
# followed by optional room type, status, persons, nights, rooms columns.
# Name group uses [ \t] (not \s) so it cannot cross a line boundary — this prevents
# "Block Code XYZ" section headers from being absorbed into the first guest's name.
# An optional leading room number (e.g. "1122 Larson,Thad") is consumed before the name.
_GUEST_RE = re.compile(
    r'^\d*[ \t]*'                                    # optional leading room number
    r'([A-Za-z][A-Za-z,\'\-\.\ \t]{1,40}?)[ \t]+'  # Last,First[,Title] — no newlines
    r'(\d{8})[ \t]+'                                 # 8-digit conf #
    r'(\d{2}/\d{2}/\d{2})[ \t]+'                    # arrival MM/DD/YY
    r'(\d{2}/\d{2}/\d{2})[ \t]+'                    # departure MM/DD/YY
    r'\S+[ \t]+'                                     # room type
    r'\d+[ \t]+'                                     # reservation status (numeric code)
    r'(\d+)[ \t]+'                                   # persons
    r'(\d+)[ \t]+'                                   # nights
    r'(\d+)',                                        # rooms
    re.MULTILINE
)

# Cambria/Choice Hotels "Group Reservation List" format:
#   Name  Account(9-12 digits)  Status(letter)  Arrival(M/DD/YYYY)  Departure  People  RoomType  Rate  Balance
# No "nights" column — computed from arrival/departure.
_GUEST_RE_CAMBRIA = re.compile(
    r'^([A-Za-z][A-Za-z,\'\-\.\ \t]{1,40}?)[ \t]+'  # Guest Name (Last, First)
    r'(\d{9,12})[ \t]+'                               # Account / conf# (9–12 digits)
    r'[A-Z][ \t]+'                                    # Status letter (R, C, etc.)
    r'(\d{1,2}/\d{2}/\d{4})[ \t]+'                   # Arrival   M/DD/YYYY
    r'(\d{1,2}/\d{2}/\d{4})[ \t]+'                   # Departure M/DD/YYYY
    r'(\d+)[ \t]+'                                    # People
    r'[A-Z]+[ \t]+'                                   # Room Type
    r'[\d,]+\.\d+[ \t]+'                              # Rate
    r'[\d,]+\.\d+',                                   # Balance
    re.MULTILINE
)

# Royal Sonesta "Group Rooming List" format:
#   Last,First  Conf#(9 digits)  Arrival(MM-DD-YY)  Departure(MM-DD-YY)  RoomType[+Status]  [Status]  Adl  Chl  Nts  Rms  ...
# Room type and status may be merged into one token (e.g. "U1KNRGC") or separate ("S1K GC").
# "Res. Notes:" continuation lines are ignored because they don't match this pattern.
_GUEST_RE_SONESTA = re.compile(
    r'^([A-Za-z][A-Za-z,\'\-\.\ \t]{1,40}?)[ \t]+'  # Last,First name
    r'(\d{9})[ \t]+'                                  # 9-digit conf#
    r'(\d{2}-\d{2}-\d{2})[ \t]+'                     # Arrival   MM-DD-YY (dashes)
    r'(\d{2}-\d{2}-\d{2})[ \t]+'                     # Departure MM-DD-YY (dashes)
    r'\S+[ \t]+'                                      # Room type token (may include status)
    r'(?:[A-Za-z]+[ \t]+)?'                           # Optional separate status code
    r'\d+[ \t]+'                                      # Adults
    r'\d+[ \t]+'                                      # Children
    r'(\d+)[ \t]+'                                    # Nights (explicit)
    r'(\d+)',                                          # Rooms
    re.MULTILINE
)

# Holiday Inn / IHG "Group Rooming List" format:
#   [RoomNo]  Last,First  Conf#(7-10 digits)  Arrival(MM-DD-YY)  Departure(MM-DD-YY)  RoomType  [Status]  Adl  Chl  Nts  Rms
# Room number is OPTIONAL — IHG only prints it on the first guest in each room.
_GUEST_RE_HOLIDAY_INN = re.compile(
    r'^(?:\d+[ \t]+)?'                               # Room number (optional, e.g. 101)
    r'([A-Za-z][A-Za-z,\'\-\.\ \t]{1,40}?)[ \t]+'  # Last,First name
    r'(\d{7,10})[ \t]+'                              # Conf# (7-10 digits)
    r'(\d{2}-\d{2}-\d{2})[ \t]+'                    # Arrival   MM-DD-YY (dashes)
    r'(\d{2}-\d{2}-\d{2})[ \t]+'                    # Departure MM-DD-YY (dashes)
    r'\S+[ \t]+'                                     # Room type token
    r'(?:[A-Za-z/]+[ \t]+)?'                         # Optional status / carrier code
    r'\d+[ \t]+'                                     # Adults
    r'\d+[ \t]+'                                     # Children
    r'(\d+)[ \t]+'                                   # Nights (explicit)
    r'(\d+)',                                         # Rooms
    re.MULTILINE
)

# IHG/Ana "Group Rooming List" format (conf# FIRST, then optional sharewithnum, then name):
#   Conf#(7-10 digits)  [ShareWith#]  Last, First  M/D/YY  M/D/YY  RoomType  $Rate  [Billing]
# Unlike all other parsers, the conf# appears BEFORE the guest name.
# Dates use M/D/YY with possible single-digit month/day (e.g. 8/23/26).
_GUEST_RE_IHG_CONF_FIRST = re.compile(
    r'^(\d{7,10})[ \t]+'                         # Conf# (first)
    r'(?:[\d,]{5,15}[ \t]+)?'                    # Optional sharewithnum (e.g. 50,019,561)
    r'([A-Za-z][A-Za-z,\'\-\.\ ]+?)[ \t]*'      # Last, First name (lazy — stops before digit)
    r'(\d{1,2}/\d{1,2}/\d{2})[ \t]+'            # Arrival   M/D/YY
    r'(\d{1,2}/\d{1,2}/\d{2})[ \t]+'            # Departure M/D/YY
    r'[A-Z]{2,6}[ \t]+'                          # Room type (TRAD, TRDD, PRKN, etc.)
    r'\$[\d.]+',                                  # Rate
    re.MULTILINE
)

# Hilton/Embassy Suites GPRMLSTS "Group Member Status Report" format
_GUEST_RE_HILTON_PMS = re.compile(
    r'^([A-Za-z][A-Za-z/\-\'\s.,]{1,50}?)'   # Guest name (Last/First or First Last)
    r'[ \t]+'
    r'(?:[A-Z0-9]+[ \t]+)?'                   # Optional room type (NKS, NK, KS, Q, etc.)
    r'N[ \t]+N[ \t]+'                          # DCI DK columns
    r'(?:\d+[ \t]*,[ \t]*0[ \t]+\S+[ \t]+)?'  # Optional: #guests, 0, rate-plan-code
    r'\$[\d.]+[ \t]+'                          # Room rate
    r'N[ \t]+'                                 # RTG column
    r'(?:[A-Z]+[ \t]+)?'                       # Optional MOP (CC, CS, etc.)
    r'(NA|AR|DP)[ \t]+'                        # Guest STATUS (active guests only)
    r'(\d{1,2}/\d{1,2}/\d{4})[ \t]+'          # Arrival  M/D/YYYY
    r'(\d{1,2}/\d{1,2}/\d{4})[ \t]*$'         # Departure M/D/YYYY (end of line)
    r'\n(\d{6,10})',                            # Confirmation number on next line
    re.MULTILINE
)


def _repair_truncated_json(text):
    """
    Recover a JSON array that was cut off mid-stream (e.g. by a token limit).
    Finds the last complete {...} object and closes the array.
    Returns a valid JSON string, or None if no complete object found.
    """
    last_brace = text.rfind('}')
    if last_brace == -1:
        return None
    truncated = text[:last_brace + 1]
    if '[' in truncated and not truncated.rstrip().endswith(']'):
        truncated = truncated.rstrip().rstrip(',') + ']'
    try:
        json.loads(truncated)
        return truncated
    except Exception:
        return None


def _mddyyyy_to_iso(s):
    """Convert M/DD/YYYY or MM/DD/YYYY to YYYY-MM-DD."""
    parts = s.strip().split('/')
    m, d, y = int(parts[0]), int(parts[1]), int(parts[2])
    return f'{y}-{m:02d}-{d:02d}'

def _mmddyy_to_iso(s):
    """Convert MM/DD/YY to YYYY-MM-DD."""
    m, d, y = s.split('/')
    yr = 2000 + int(y) if int(y) < 50 else 1900 + int(y)
    return f'{yr}-{int(m):02d}-{int(d):02d}'

def _mmddyy_dashes_to_iso(s):
    """Convert MM-DD-YY (dashes) to YYYY-MM-DD."""
    m, d, y = s.split('-')
    yr = 2000 + int(y) if int(y) < 50 else 1900 + int(y)
    return f'{yr}-{int(m):02d}-{int(d):02d}'


def _ai_parse_rooming_list_vision(images, api_key):
    """
    Vision fallback: send rendered PDF page images to Claude and extract guest records.
    Used when pdfplumber returns no text (scanned / image-only PDF).

    Returns (guests_list, error_str) — same shape as _ai_parse_rooming_list.
    """
    try:
        import anthropic
    except ImportError:
        return [], 'anthropic package not installed'

    rooming_prompt = (
        "You are parsing scanned pages of a hotel group rooming list.\n"
        "Extract every individual guest reservation visible across all the images.\n\n"
        "Return ONLY a valid JSON array — no other text, no markdown fences, no explanation.\n"
        "Each element must have exactly these keys:\n"
        "  name       — guest name as it appears (e.g. \"Smith,John\" or \"John Smith\")\n"
        "  conf_no    — reservation/confirmation number (string)\n"
        "  arrival    — arrival date in YYYY-MM-DD format\n"
        "  departure  — departure date in YYYY-MM-DD format\n"
        "  nights     — integer number of nights\n"
        "  rooms      — integer number of rooms (default 1 if not shown)\n\n"
        "Rules:\n"
        "- Skip header rows, page headers/footers, summary/total rows, and note lines.\n"
        "- If 'nights' is not explicit, compute it as (departure - arrival) in days.\n"
        "- Convert all dates to YYYY-MM-DD regardless of the source format.\n"
        "- If a line has an obvious continuation (e.g. 'Res. Notes:'), skip it.\n"
        "- Include every guest across all pages — do not truncate the list."
    )

    content = [{'type': 'text', 'text': rooming_prompt}]
    for b64 in images:
        content.append({
            'type': 'image',
            'source': {'type': 'base64', 'media_type': 'image/jpeg', 'data': b64}
        })

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model='claude-opus-4-5',
            max_tokens=8192,
            messages=[{'role': 'user', 'content': content}]
        )
        response_text = message.content[0].text.strip()
        response_text = re.sub(r'^```[a-z]*\s*', '', response_text)
        response_text = re.sub(r'\s*```$', '', response_text)

        try:
            guests_raw = json.loads(response_text)
        except json.JSONDecodeError:
            repaired = _repair_truncated_json(response_text)
            if repaired:
                guests_raw = json.loads(repaired)
            else:
                raise

        result = []
        for g in guests_raw:
            try:
                arr_str = str(g.get('arrival', '')).strip()
                dep_str = str(g.get('departure', '')).strip()
                nights  = int(g.get('nights', 0))
                if nights <= 0:
                    nights = (datetime.strptime(dep_str, '%Y-%m-%d') -
                              datetime.strptime(arr_str, '%Y-%m-%d')).days
                if nights <= 0:
                    continue
                result.append({
                    'name':      str(g.get('name', '')).strip(),
                    'conf_no':   str(g.get('conf_no', '')).strip(),
                    'arrival':   arr_str,
                    'departure': dep_str,
                    'nights':    nights,
                    'rooms':     int(g.get('rooms', 1)),
                })
            except Exception:
                continue

        return result, None

    except Exception as e:
        return [], str(e)


def _ai_parse_rooming_list(all_text):
    """
    AI fallback: send raw PDF text to Claude and ask it to extract guest records.
    Used when no regex pattern recognizes the hotel's format.

    Returns a list of guest dicts (same shape as regex parsers), or [] if
    the API key is not configured or the call fails.
    """
    # Load API key from config.py, fall back to environment variable
    api_key = ''
    try:
        import importlib, sys
        # Force reload so edits to config.py take effect without restarting
        if 'config' in sys.modules:
            importlib.reload(sys.modules['config'])
        else:
            import config as _cfg
            sys.modules['config'] = _cfg
        api_key = sys.modules['config'].ANTHROPIC_API_KEY.strip()
    except Exception:
        pass
    if not api_key:
        import os
        api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
    if not api_key:
        return [], None

    try:
        import anthropic
    except ImportError:
        return [], None

    prompt = (
        "You are parsing a hotel group rooming list exported from a property management system.\n"
        "Extract every individual guest reservation from the text below.\n\n"
        "Return ONLY a valid JSON array — no other text, no markdown fences, no explanation.\n"
        "Each element must have exactly these keys:\n"
        "  name       — guest name as it appears (e.g. \"Smith,John\" or \"John Smith\")\n"
        "  conf_no    — reservation/confirmation number (string)\n"
        "  arrival    — arrival date in YYYY-MM-DD format\n"
        "  departure  — departure date in YYYY-MM-DD format\n"
        "  nights     — integer number of nights\n"
        "  rooms      — integer number of rooms (default 1 if not shown)\n\n"
        "Rules:\n"
        "- Skip header rows, page headers/footers, summary/total rows, and note lines.\n"
        "- If 'nights' is not explicit, compute it as (departure - arrival) in days.\n"
        "- Convert all dates to YYYY-MM-DD regardless of the source format.\n"
        "- If a line has an obvious continuation (e.g. 'Res. Notes:'), skip it.\n\n"
        "Rooming list text:\n"
        + all_text[:12000]   # ~3k tokens; enough for 100+ guests
    )

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model='claude-opus-4-5',
            max_tokens=8192,
            messages=[{'role': 'user', 'content': prompt}]
        )
        response_text = message.content[0].text.strip()
        # Strip markdown code fences if the model wrapped the JSON
        response_text = re.sub(r'^```[a-z]*\s*', '', response_text)
        response_text = re.sub(r'\s*```$', '', response_text)

        try:
            guests_raw = json.loads(response_text)
        except json.JSONDecodeError:
            # Response was truncated (hit token limit) — recover partial results
            repaired = _repair_truncated_json(response_text)
            if repaired:
                guests_raw = json.loads(repaired)
            else:
                raise

        result = []
        for g in guests_raw:
            try:
                arr_str = str(g.get('arrival', '')).strip()
                dep_str = str(g.get('departure', '')).strip()
                nights  = int(g.get('nights', 0))
                if nights <= 0:
                    nights = (datetime.strptime(dep_str, '%Y-%m-%d') -
                              datetime.strptime(arr_str, '%Y-%m-%d')).days
                if nights <= 0:
                    continue
                result.append({
                    'name':      str(g.get('name', '')).strip(),
                    'conf_no':   str(g.get('conf_no', '')).strip(),
                    'arrival':   arr_str,
                    'departure': dep_str,
                    'nights':    nights,
                    'rooms':     int(g.get('rooms', 1)),
                })
            except Exception:
                continue

        return result, None

    except Exception as e:
        return [], str(e)

def parse_rooming_list_pdf(file_bytes):
    """
    Parse a hotel rooming list PDF (pdfplumber).

    Returns:
      {
        'guests':        [{name, conf_no, arrival, departure, nights, rooms}],
        'nights_by_date': {"YYYY-MM-DD": N, ...},   # rooms per night
        'total_guests':  N,
        'error':         str or None
      }
    """
    try:
        import pdfplumber
    except ImportError:
        return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                'error': 'pdfplumber not installed'}

    try:
        guests = []
        all_text = ''

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ''
                all_text += '\n' + text

        for m in _GUEST_RE.finditer(all_text):
            name_raw  = m.group(1).strip().rstrip(',')
            conf_no   = m.group(2)
            arr_raw   = m.group(3)
            dep_raw   = m.group(4)
            nights    = int(m.group(6))
            rooms     = int(m.group(7))

            try:
                arrival    = _mmddyy_to_iso(arr_raw)
                departure  = _mmddyy_to_iso(dep_raw)
            except Exception:
                continue

            guests.append({
                'name':      name_raw,
                'conf_no':   conf_no,
                'arrival':   arrival,
                'departure': departure,
                'nights':    nights,
                'rooms':     rooms,
            })

        # ── Fallback: Cambria/Choice Hotels "Group Reservation List" format ──
        if not guests:
            for m in _GUEST_RE_CAMBRIA.finditer(all_text):
                name_raw = m.group(1).strip().rstrip(',')
                conf_no  = m.group(2)
                arr_raw  = m.group(3)
                dep_raw  = m.group(4)
                try:
                    arrival   = _mddyyyy_to_iso(arr_raw)
                    departure = _mddyyyy_to_iso(dep_raw)
                    nights    = (datetime.strptime(departure, '%Y-%m-%d') -
                                 datetime.strptime(arrival,   '%Y-%m-%d')).days
                except Exception:
                    continue
                if nights <= 0:
                    continue
                guests.append({
                    'name':      name_raw,
                    'conf_no':   conf_no,
                    'arrival':   arrival,
                    'departure': departure,
                    'nights':    nights,
                    'rooms':     1,
                })

        # ── Fallback: Royal Sonesta "Group Rooming List" format (MM-DD-YY dashes) ──
        if not guests:
            for m in _GUEST_RE_SONESTA.finditer(all_text):
                name_raw = m.group(1).strip().rstrip(',')
                conf_no  = m.group(2)
                arr_raw  = m.group(3)
                dep_raw  = m.group(4)
                nights   = int(m.group(5))
                rooms    = int(m.group(6))
                try:
                    arrival   = _mmddyy_dashes_to_iso(arr_raw)
                    departure = _mmddyy_dashes_to_iso(dep_raw)
                except Exception:
                    continue
                if nights <= 0:
                    continue
                guests.append({
                    'name':      name_raw,
                    'conf_no':   conf_no,
                    'arrival':   arrival,
                    'departure': departure,
                    'nights':    nights,
                    'rooms':     rooms,
                })

        # ── Fallback: Hilton/Embassy Suites GPRMLSTS "Group Member Status Report" ──
        if not guests:
            seen_conf = set()
            for m in _GUEST_RE_HILTON_PMS.finditer(all_text):
                name_raw = m.group(1).strip().rstrip(',').rstrip()
                arr_raw  = m.group(3)
                dep_raw  = m.group(4)
                conf_no  = m.group(5)
                key = (conf_no, arr_raw, dep_raw)
                if key in seen_conf:
                    continue
                seen_conf.add(key)
                try:
                    arrival   = _mddyyyy_to_iso(arr_raw)
                    departure = _mddyyyy_to_iso(dep_raw)
                    nights    = (datetime.strptime(departure, '%Y-%m-%d') -
                                 datetime.strptime(arrival,   '%Y-%m-%d')).days
                except Exception:
                    continue
                if nights <= 0:
                    continue
                guests.append({
                    'name':      name_raw,
                    'conf_no':   conf_no,
                    'arrival':   arrival,
                    'departure': departure,
                    'nights':    nights,
                    'rooms':     1,
                })

        # ── Fallback: Holiday Inn / IHG "Group Rooming List" (room# prefix, MM-DD-YY dashes) ──
        if not guests:
            for m in _GUEST_RE_HOLIDAY_INN.finditer(all_text):
                name_raw = m.group(1).strip().rstrip(',')
                conf_no  = m.group(2)
                arr_raw  = m.group(3)
                dep_raw  = m.group(4)
                nights   = int(m.group(5))
                rooms    = int(m.group(6))
                try:
                    arrival   = _mmddyy_dashes_to_iso(arr_raw)
                    departure = _mmddyy_dashes_to_iso(dep_raw)
                except Exception:
                    continue
                if nights <= 0:
                    continue
                guests.append({
                    'name':      name_raw,
                    'conf_no':   conf_no,
                    'arrival':   arrival,
                    'departure': departure,
                    'nights':    nights,
                    'rooms':     rooms,
                })

        # ── Fallback: IHG/Ana "Group Rooming List" — conf# first, M/D/YY dates ──
        if not guests:
            for m in _GUEST_RE_IHG_CONF_FIRST.finditer(all_text):
                conf_no  = m.group(1)
                name_raw = m.group(2).strip().rstrip(',').rstrip()
                arr_raw  = m.group(3)
                dep_raw  = m.group(4)
                if not name_raw:
                    continue
                try:
                    arrival   = _mmddyy_to_iso(arr_raw)
                    departure = _mmddyy_to_iso(dep_raw)
                    nights    = (datetime.strptime(departure, '%Y-%m-%d') -
                                 datetime.strptime(arrival,   '%Y-%m-%d')).days
                except Exception:
                    continue
                if nights <= 0:
                    continue
                guests.append({
                    'name':      name_raw,
                    'conf_no':   conf_no,
                    'arrival':   arrival,
                    'departure': departure,
                    'nights':    nights,
                    'rooms':     1,
                })

        # ── Final fallback: AI parser for unrecognized formats ──
        ai_parsed = False
        ai_error  = None
        if not guests:
            if not all_text.strip():
                # Scanned / image-only PDF — use vision
                api_key = ''
                try:
                    import importlib, sys as _sys
                    if 'config' in _sys.modules:
                        importlib.reload(_sys.modules['config'])
                    else:
                        import config as _cfg
                        _sys.modules['config'] = _cfg
                    api_key = _sys.modules['config'].ANTHROPIC_API_KEY.strip()
                except Exception:
                    pass
                if not api_key:
                    import os
                    api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
                if api_key:
                    images = _pdf_pages_to_images(file_bytes, max_pages=20, dpi=100)
                    if images:
                        ai_guests, ai_error = _ai_parse_rooming_list_vision(images, api_key)
                        if ai_guests:
                            guests    = ai_guests
                            ai_parsed = True
                    else:
                        ai_error = 'Could not render PDF pages — pymupdf may not be installed.'
                else:
                    ai_error = 'Scanned PDF detected but Anthropic API key is not configured.'
            else:
                ai_guests, ai_error = _ai_parse_rooming_list(all_text)
                if ai_guests:
                    guests    = ai_guests
                    ai_parsed = True

        # Compute rooms per night from arrival/departure dates
        nights_by_date = {}
        for g in guests:
            try:
                arr = datetime.strptime(g['arrival'],   '%Y-%m-%d')
                dep = datetime.strptime(g['departure'], '%Y-%m-%d')
                cur = arr
                while cur < dep:
                    key = cur.strftime('%Y-%m-%d')
                    nights_by_date[key] = nights_by_date.get(key, 0) + g['rooms']
                    cur += timedelta(days=1)
            except Exception:
                continue

        return {
            'guests':         guests,
            'nights_by_date': nights_by_date,
            'total_guests':   len(guests),
            'ai_parsed':      ai_parsed,
            'ai_error':       ai_error,
            'error':          None,
        }

    except Exception as e:
        return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                'ai_parsed': False, 'error': str(e)}


# ── CSV / XLS rooming list parser ────────────────────────────────────────────

def _parse_date_flexible(val):
    """Try multiple date formats and return YYYY-MM-DD, or None."""
    if val is None:
        return None
    # openpyxl returns datetime objects directly — handle first
    from datetime import datetime as _dt, date as _date_cls
    if isinstance(val, (_dt, _date_cls)):
        try:
            return val.strftime('%Y-%m-%d')
        except Exception:
            return None
    s = str(val).strip()
    # Handle 'YYYY-MM-DD HH:MM:SS' from str(datetime) conversion
    if len(s) > 10 and s[10] == ' ':
        s = s[:10]
    for fmt in ('%m/%d/%Y', '%m/%d/%y', '%Y-%m-%d', '%d-%b-%Y', '%b-%d-%Y',
                '%d/%m/%Y', '%B %d, %Y', '%b %d, %Y', '%B-%d-%Y'):
        try:
            return _dt.strptime(s, fmt).strftime('%Y-%m-%d')
        except ValueError:
            continue
    # Excel serial date (float/int)
    try:
        serial = float(s)
        from datetime import timedelta as _td
        base = _dt(1899, 12, 30)
        return (base + _td(days=serial)).strftime('%Y-%m-%d')
    except Exception:
        pass
    return None


def _col_index(headers, *candidates):
    """Return the index of the first matching header (case-insensitive), or None."""
    lower = [h.lower().strip() for h in headers]
    for c in candidates:
        for i, h in enumerate(lower):
            if c.lower() in h:
                return i
    return None


def _build_guest_result(guests):
    """
    Compute nights_by_date from a parsed guest list, deduplicating by
    Hotel Confirmation # so that two guests sharing a room count as 1 room-night.

    Returns dedup_summary: list of {conf_no, names, nights} for shared rooms.
    """
    # Group guests by conf_no (non-empty)
    from collections import OrderedDict
    conf_groups = OrderedDict()  # conf_no -> list of guests
    no_conf = []
    for g in guests:
        cn = (g.get('conf_no') or '').strip()
        if cn:
            conf_groups.setdefault(cn, []).append(g)
        else:
            no_conf.append(g)

    # Build dedup summary (conf_nos with >1 guest = shared room)
    dedup_summary = []
    for cn, grp in conf_groups.items():
        if len(grp) > 1:
            dedup_summary.append({
                'conf_no': cn,
                'count':   len(grp),
                'names':   ', '.join(g['name'] for g in grp),
                'arrival': grp[0].get('arrival', ''),
                'departure': grp[0].get('departure', ''),
            })

    # Compute nights_by_date using ONE room per unique conf_no
    nights_by_date = {}
    # Each unique conf_no = 1 room
    representative_guests = [grp[0] for grp in conf_groups.values()] + no_conf
    for g in representative_guests:
        try:
            arr = datetime.strptime(g['arrival'],   '%Y-%m-%d')
            dep = datetime.strptime(g['departure'], '%Y-%m-%d')
            cur = arr
            while cur < dep:
                key = cur.strftime('%Y-%m-%d')
                nights_by_date[key] = nights_by_date.get(key, 0) + 1
                cur += timedelta(days=1)
        except Exception:
            continue

    # Unique room count = unique conf_nos + guests with no conf_no
    unique_rooms = len(conf_groups) + len(no_conf)

    return {
        'guests':         guests,
        'nights_by_date': nights_by_date,
        'total_guests':   len(guests),
        'unique_rooms':   unique_rooms,
        'dedup_summary':  dedup_summary,
        'error':          None,
    }


def _parse_date_range(val):
    """
    Parse a combined date-range string like 'May 23, 2026 - May 29, 2026'
    Returns (arrival_iso, departure_iso) or (None, None).
    """
    if not val:
        return None, None
    # Try splitting on ' - ' or ' – ' (en-dash)
    for sep in (' - ', ' – ', '-'):
        parts = str(val).split(sep, 1)
        if len(parts) == 2:
            arr = _parse_date_flexible(parts[0].strip())
            dep = _parse_date_flexible(parts[1].strip())
            if arr and dep:
                return arr, dep
    return None, None


def parse_rooming_list_csv(file_bytes):
    """
    Parse a hotel rooming list CSV.

    Handles two date layouts:
      A) Separate arrival + departure columns
      B) Single combined column e.g. 'Guest Reservation Dates': 'May 23, 2026 - May 29, 2026'
    """
    import csv, io as _io
    try:
        text = file_bytes.decode('utf-8-sig', errors='replace')
        # Handle tab-separated files exported as .csv
        sample = text[:2000]
        delimiter = '\t' if sample.count('\t') > sample.count(',') else ','
        reader = csv.reader(_io.StringIO(text), delimiter=delimiter)
        rows = list(reader)
        if not rows:
            return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                    'error': 'Empty CSV file'}

        # Find header row
        header_row_idx = 0
        headers = rows[0]
        for i, row in enumerate(rows[:10]):
            joined = ' '.join(row).lower()
            if any(k in joined for k in ('arrival', 'check-in', 'checkin', 'check in',
                                          'last name', 'first name', 'guest name', 'name',
                                          'reservation date', 'stay date')):
                headers = row
                header_row_idx = i
                break

        # Map columns
        i_last   = _col_index(headers, 'last name', 'last')
        i_first  = _col_index(headers, 'first name', 'first')
        i_name   = _col_index(headers, 'guest name', 'full name', 'name')
        i_arr    = _col_index(headers, 'arrival', 'check-in', 'checkin', 'check in',
                                       'arrive', 'arr date', 'arrival date')
        i_dep    = _col_index(headers, 'departure', 'check-out', 'checkout', 'check out',
                                       'depart', 'dep date', 'departure date')
        # Combined date range column (e.g. Passkey/Cvent export)
        i_range  = _col_index(headers, 'reservation date', 'stay date', 'guest reservation',
                                       'dates', 'date range', 'hotel date')
        i_conf   = _col_index(headers, 'hotel confirmation', 'confirmation #', 'hotel conf',
                                       'conf #', 'conf no', 'confirmation', 'reservation #',
                                       'passkey', 'booking ref', 'res id')
        i_rooms  = _col_index(headers, 'rooms', 'room count', 'qty', 'quantity')
        i_nights = _col_index(headers, 'nights', 'night count', 'duration', 'los')

        # Need either (arrival + departure) OR a combined date-range column
        has_split_dates = (i_arr is not None and i_dep is not None)
        has_range_date  = (i_range is not None)
        if not has_split_dates and not has_range_date:
            return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                    'error': 'Could not find date columns. '
                             f'Headers found: {list(headers)}'}

        guests = []
        for row in rows[header_row_idx + 1:]:
            if not row or all(c.strip() == '' for c in row):
                continue
            def cell(idx):
                return row[idx].strip() if idx is not None and idx < len(row) else ''

            # Build name
            if i_last is not None and i_first is not None:
                name = f"{cell(i_last)}, {cell(i_first)}".strip(', ')
            elif i_name is not None:
                name = cell(i_name)
            else:
                name = 'Unknown'

            # Parse dates
            if has_split_dates:
                arrival   = _parse_date_flexible(cell(i_arr))
                departure = _parse_date_flexible(cell(i_dep))
            else:
                arrival, departure = _parse_date_range(cell(i_range))
            if not arrival or not departure:
                continue

            conf_no = cell(i_conf) if i_conf is not None else ''
            rooms   = int(cell(i_rooms)) if i_rooms is not None and cell(i_rooms).isdigit() else 1
            if i_nights is not None and cell(i_nights).isdigit():
                nights = int(cell(i_nights))
            else:
                try:
                    arr_d = datetime.strptime(arrival,   '%Y-%m-%d')
                    dep_d = datetime.strptime(departure, '%Y-%m-%d')
                    nights = (dep_d - arr_d).days
                except Exception:
                    nights = 0

            guests.append({
                'name':      name,
                'conf_no':   conf_no,
                'arrival':   arrival,
                'departure': departure,
                'nights':    nights,
                'rooms':     rooms,
            })

        return _build_guest_result(guests)

    except Exception as e:
        return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                'error': str(e)}


# ── Omni / IHG multi-tab rooming list XLSX ───────────────────────────────────

def _parse_omni_rooming_list_xlsx(wb):
    """
    Detect and parse the Omni/IHG multi-tab rooming list workbook.

    Layout:
      Sheet 1  (Pick-up)       — pickup report; used only for hotel name / date
      Sheet 2+ (Rooming List)  — guest table starting at row 4:
        A: section label or None  B: Last Name  C: First Name
        D: Conf #                 E: Status      F: Arrival (datetime)
        G: Departure (datetime)   H: Nights      I: Adults
        J: Children               K: Rate        L: Billing notes

    Returns a dict with keys: guests, hotel_name, report_date
    — or None if the workbook does not match this layout.
    """
    rl_ws = None
    for name in wb.sheetnames:
        if 'rooming' in name.lower():
            rl_ws = wb[name]
            break
    if rl_ws is None:
        return None

    hotel_name  = str(wb.worksheets[0].cell(1, 1).value or '').strip()
    report_date = wb.worksheets[0].cell(2, 1).value

    guests = []
    current_section = ''

    for row in rl_ws.iter_rows(min_row=4, values_only=True):
        r = (list(row) + [None] * 12)[:12]
        col_a, last, first, conf, status, arr, dep, nts, adults, children, rate, notes = r

        if not conf and not last:
            continue

        if col_a:
            parts = [p.strip() for p in str(col_a).split('\n') if p.strip()]
            current_section = parts[-1] if parts else str(col_a).strip()
            if not conf:
                continue

        if not arr or not dep:
            continue

        last_s  = str(last  or '').strip()
        first_s = str(first or '').strip()
        name = f"{last_s}, {first_s}".strip(', ') if first_s else last_s
        if not name:
            continue

        arrival   = arr.strftime('%Y-%m-%d') if hasattr(arr, 'strftime') else _parse_date_flexible(str(arr))
        departure = dep.strftime('%Y-%m-%d') if hasattr(dep, 'strftime') else _parse_date_flexible(str(dep))
        if not arrival or not departure:
            continue

        nights = int(nts) if nts and str(nts).strip().isdigit() else 0
        if nights <= 0:
            try:
                nights = (datetime.strptime(departure, '%Y-%m-%d') -
                          datetime.strptime(arrival,   '%Y-%m-%d')).days
            except Exception:
                continue
        if nights <= 0:
            continue

        guests.append({
            'name':      name,
            'conf_no':   str(conf or '').strip(),
            'arrival':   arrival,
            'departure': departure,
            'nights':    nights,
            'rooms':     1,
            'section':   current_section,
            'rate':      float(rate) if rate else None,
        })

    if not guests:
        return None

    return {
        'guests':      guests,
        'hotel_name':  hotel_name,
        'report_date': report_date,
    }


def generate_rooming_list_pdf(guests, title='Group Rooming List',
                               hotel_name='', report_date=None):
    """
    Render a clean rooming-list PDF from a parsed guest list (reportlab).

    guests: list of dicts with keys: name, conf_no, arrival, departure,
            nights, rooms, section (optional), rate (optional)

    Returns bytes (PDF), or None if reportlab is not installed.
    """
    try:
        from reportlab.lib.pagesizes import landscape, letter
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                         Paragraph, Spacer)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        return None

    import io as _io
    from collections import OrderedDict

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(letter),
        leftMargin=0.45 * inch, rightMargin=0.45 * inch,
        topMargin=0.45 * inch, bottomMargin=0.45 * inch,
        title=title,
    )

    styles = getSampleStyleSheet()
    title_sty = ParagraphStyle('rl_title', parent=styles['Heading1'],
                               fontSize=13, spaceAfter=3, alignment=TA_CENTER)
    sub_sty   = ParagraphStyle('rl_sub',   parent=styles['Normal'],
                               fontSize=9,  spaceAfter=2, alignment=TA_CENTER)
    sec_sty   = ParagraphStyle('rl_sec',   parent=styles['Normal'],
                               fontSize=8,  leading=10)

    NAVY   = colors.HexColor('#1A3C5E')
    STRIPE = colors.HexColor('#F0F5FB')
    SEC_BG = colors.HexColor('#D9E8F5')
    TOT_BG = colors.HexColor('#C6D9F0')
    GRID_C = colors.HexColor('#B0C4D8')

    has_rate = any(g.get('rate') for g in guests)

    if has_rate:
        col_heads  = ['#', 'Last Name', 'First Name', 'Confirmation #',
                      'Arrival', 'Departure', 'Nts', 'Adl', 'Rate']
        col_widths = [0.28*inch, 1.35*inch, 1.15*inch, 1.45*inch,
                      1.0*inch,  1.0*inch,   0.38*inch, 0.38*inch, 0.65*inch]
    else:
        col_heads  = ['#', 'Last Name', 'First Name', 'Confirmation #',
                      'Arrival', 'Departure', 'Nts', 'Adl']
        col_widths = [0.28*inch, 1.5*inch, 1.3*inch, 1.6*inch,
                      1.1*inch,  1.1*inch,  0.45*inch, 0.45*inch]

    ncols = len(col_heads)

    sections = OrderedDict()
    for g in guests:
        sec = g.get('section') or ''
        sections.setdefault(sec, []).append(g)

    table_data = [col_heads]
    row_styles = []
    rn = 1

    for sec, sec_guests in sections.items():
        if sec:
            table_data.append([Paragraph(f'<b>{sec}</b>', sec_sty)] + [''] * (ncols - 1))
            row_styles += [
                ('BACKGROUND', (0, rn), (-1, rn), SEC_BG),
                ('SPAN',       (0, rn), (-1, rn)),
                ('TOPPADDING', (0, rn), (-1, rn), 3),
                ('BOTTOMPADDING', (0, rn), (-1, rn), 3),
            ]
            rn += 1

        for i, g in enumerate(sec_guests, 1):
            parts    = g['name'].split(',', 1)
            last_nm  = parts[0].strip()
            first_nm = parts[1].strip() if len(parts) > 1 else ''
            try:
                arr_s = datetime.strptime(g['arrival'],   '%Y-%m-%d').strftime('%m/%d/%y')
                dep_s = datetime.strptime(g['departure'], '%Y-%m-%d').strftime('%m/%d/%y')
            except Exception:
                arr_s = g.get('arrival',   '')
                dep_s = g.get('departure', '')

            rate_v = g.get('rate')
            rate_s = f'${rate_v:,.0f}' if rate_v else ''

            data_row = [str(i), last_nm, first_nm, g.get('conf_no', ''),
                        arr_s, dep_s, str(g.get('nights', '')),
                        str(g.get('rooms', 1))]
            if has_rate:
                data_row.append(rate_s)

            table_data.append(data_row)
            if i % 2 == 0:
                row_styles.append(('BACKGROUND', (0, rn), (-1, rn), STRIPE))
            rn += 1

        sub_row = [''] * ncols
        sub_row[ncols - (3 if has_rate else 2)] = Paragraph(
            f'<b>{sec} — {len(sec_guests)} room{"s" if len(sec_guests) != 1 else ""}</b>', sec_sty)
        table_data.append(sub_row)
        row_styles += [
            ('SPAN',         (0, rn), (-1, rn)),
            ('ALIGN',        (0, rn), (-1, rn), 'RIGHT'),
            ('TOPPADDING',   (0, rn), (-1, rn), 2),
            ('BOTTOMPADDING',(0, rn), (-1, rn), 2),
        ]
        rn += 1

    gt_row = [''] * ncols
    gt_row[ncols - (3 if has_rate else 2)] = Paragraph(
        f'<b>Grand Total: {len(guests)} room{"s" if len(guests) != 1 else ""}</b>', sec_sty)
    table_data.append(gt_row)
    row_styles += [
        ('BACKGROUND',    (0, rn), (-1, rn), TOT_BG),
        ('SPAN',          (0, rn), (-1, rn)),
        ('ALIGN',         (0, rn), (-1, rn), 'RIGHT'),
        ('TOPPADDING',    (0, rn), (-1, rn), 3),
        ('BOTTOMPADDING', (0, rn), (-1, rn), 3),
    ]

    base_style = [
        ('FONTNAME',      (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE',      (0, 0), (-1, -1), 8),
        ('FONTNAME',      (0, 0), (-1,  0), 'Helvetica-Bold'),
        ('BACKGROUND',    (0, 0), (-1,  0), NAVY),
        ('TEXTCOLOR',     (0, 0), (-1,  0), colors.white),
        ('GRID',          (0, 0), (-1, -1), 0.3, GRID_C),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN',         (0, 0), (0,  -1), 'CENTER'),
        ('ALIGN',         (6, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING',    (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING',   (0, 0), (-1, -1), 4),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 4),
        ('ROWBACKGROUND', (0, 1), (-1, -2), [colors.white, STRIPE]),
    ]

    t = Table(table_data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle(base_style + row_styles))

    # Rooms-per-night summary
    nbd = {}
    for g in guests:
        try:
            cur = datetime.strptime(g['arrival'],   '%Y-%m-%d')
            end = datetime.strptime(g['departure'], '%Y-%m-%d')
            while cur < end:
                k = cur.strftime('%Y-%m-%d')
                nbd[k] = nbd.get(k, 0) + g.get('rooms', 1)
                cur += timedelta(days=1)
        except Exception:
            pass

    story = []
    story.append(Paragraph(hotel_name or '', title_sty))
    story.append(Paragraph(title, sub_sty))
    if report_date:
        rd_s = (report_date.strftime('%B %d, %Y')
                if hasattr(report_date, 'strftime') else str(report_date))
        story.append(Paragraph(f'As of {rd_s}', sub_sty))
    story.append(Spacer(1, 0.12 * inch))
    story.append(t)

    if nbd:
        cell_sty = ParagraphStyle('rl_cell', parent=styles['Normal'], fontSize=8, leading=10)
        story.append(Spacer(1, 0.18 * inch))
        story.append(Paragraph('<b>Rooms Per Night</b>', cell_sty))
        story.append(Spacer(1, 0.04 * inch))
        nbd_data = [['Date', 'Rooms']]
        for d in sorted(nbd):
            try:
                lbl = datetime.strptime(d, '%Y-%m-%d').strftime('%a %m/%d/%y')
            except Exception:
                lbl = d
            nbd_data.append([lbl, str(nbd[d])])
        nbd_table = Table(nbd_data, colWidths=[1.3 * inch, 0.65 * inch])
        nbd_table.setStyle(TableStyle([
            ('FONTNAME',      (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE',      (0, 0), (-1, -1), 8),
            ('FONTNAME',      (0, 0), (-1,  0), 'Helvetica-Bold'),
            ('BACKGROUND',    (0, 0), (-1,  0), NAVY),
            ('TEXTCOLOR',     (0, 0), (-1,  0), colors.white),
            ('GRID',          (0, 0), (-1, -1), 0.3, GRID_C),
            ('ALIGN',         (1, 0), (1,  -1), 'CENTER'),
            ('TOPPADDING',    (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ('LEFTPADDING',   (0, 0), (-1, -1), 4),
        ]))
        story.append(nbd_table)

    doc.build(story)
    return buf.getvalue()


def parse_rooming_list_xls(file_bytes, filename=''):
    """
    Parse a hotel rooming list XLS or XLSX file.

    Uses openpyxl for .xlsx and xlrd for .xls.
    """
    ext = (filename.rsplit('.', 1)[-1].lower()) if '.' in filename else 'xlsx'
    try:
        if ext == 'xls':
            try:
                import xlrd
            except ImportError:
                return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                        'error': 'xlrd not installed — run: pip install xlrd'}
            import io as _io
            wb = xlrd.open_workbook(file_contents=file_bytes)
            ws = wb.sheet_by_index(0)
            rows = [[str(ws.cell_value(r, c)).strip() for c in range(ws.ncols)]
                    for r in range(ws.nrows)]
        else:
            import openpyxl
            import io as _io
            wb = openpyxl.load_workbook(_io.BytesIO(file_bytes), data_only=True)

            # ── Omni/IHG structured multi-tab format ──────────────────────────
            omni = _parse_omni_rooming_list_xlsx(wb)
            if omni:
                result = _build_guest_result(omni['guests'])
                pdf = generate_rooming_list_pdf(
                    omni['guests'],
                    title='Group Rooming List',
                    hotel_name=omni.get('hotel_name', ''),
                    report_date=omni.get('report_date'),
                )
                if pdf:
                    result['pdf_bytes'] = pdf
                return result
            # ─────────────────────────────────────────────────────────────────

            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(c).strip() if c is not None else '' for c in row])

        if not rows:
            return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                    'error': 'Empty spreadsheet'}

        # Find header row
        header_row_idx = 0
        headers = rows[0]
        for i, row in enumerate(rows[:15]):
            joined = ' '.join(row).lower()
            if any(k in joined for k in ('arrival', 'check-in', 'checkin', 'check in',
                                          'last name', 'first name', 'guest name', 'name')):
                headers = row
                header_row_idx = i
                break

        # Map columns (same logic as CSV parser)
        i_last    = _col_index(headers, 'last name', 'last')
        i_first   = _col_index(headers, 'first name', 'first')
        i_name    = _col_index(headers, 'guest name', 'full name', 'name')
        i_arr     = _col_index(headers, 'arrival', 'check-in', 'checkin', 'check in',
                                        'arrive', 'arr date', 'arrival date')
        i_dep     = _col_index(headers, 'departure', 'check-out', 'checkout', 'check out',
                                        'depart', 'dep date', 'departure date')
        i_conf    = _col_index(headers, 'conf', 'confirmation', 'reservation', 'res no',
                                        'booking ref', 'res #', 'res id')
        i_rooms   = _col_index(headers, 'rooms', 'room count', 'qty', 'quantity')
        i_nights  = _col_index(headers, 'nights', 'night count', 'duration', 'los',
                                        'length of stay')

        if i_arr is None or i_dep is None:
            return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                    'error': 'Could not find arrival/departure columns. '
                             f'Headers found: {headers}'}

        guests = []
        for row in rows[header_row_idx + 1:]:
            if not row or all(c == '' or c == 'None' for c in row):
                continue
            def cell(i):
                return row[i].strip() if i is not None and i < len(row) else ''

            if i_last is not None and i_first is not None:
                name = f"{cell(i_last)}, {cell(i_first)}".strip(', ')
            elif i_name is not None:
                name = cell(i_name)
            else:
                name = 'Unknown'

            arrival   = _parse_date_flexible(cell(i_arr))
            departure = _parse_date_flexible(cell(i_dep))
            if not arrival or not departure:
                continue

            conf_no = cell(i_conf) if i_conf is not None else ''
            rooms   = int(cell(i_rooms)) if i_rooms is not None and cell(i_rooms).isdigit() else 1
            if i_nights is not None and cell(i_nights).isdigit():
                nights = int(cell(i_nights))
            else:
                try:
                    arr_d = datetime.strptime(arrival,   '%Y-%m-%d')
                    dep_d = datetime.strptime(departure, '%Y-%m-%d')
                    nights = (dep_d - arr_d).days
                except Exception:
                    nights = 0

            guests.append({
                'name':      name,
                'conf_no':   conf_no,
                'arrival':   arrival,
                'departure': departure,
                'nights':    nights,
                'rooms':     rooms,
            })

        return _build_guest_result(guests)

    except Exception as e:
        return {'guests': [], 'nights_by_date': {}, 'total_guests': 0,
                'error': str(e)}


# ── Reconciliation ────────────────────────────────────────────────────────────

def reconcile(pickup_by_night, rooming_nights_by_date):
    """
    Compare the entered weekly pickup (dict date→rooms) against the rooming
    list extract (dict date→rooms).

    Returns:
      {'status': 'match'|'discrepancy', 'notes': str}
    """
    diffs = []
    all_dates = sorted(set(list(pickup_by_night.keys()) +
                           list(rooming_nights_by_date.keys())))
    for date in all_dates:
        p = pickup_by_night.get(date)
        r = rooming_nights_by_date.get(date, 0)
        if p is None:
            continue   # night not in contracted block
        if p != r:
            diffs.append(
                f"{date}: pickup says {p}, rooming list shows {r} "
                f"({'over' if r > p else 'under'} by {abs(r - p)})"
            )

    if diffs:
        return {'status': 'discrepancy', 'notes': '\n'.join(diffs)}
    return {'status': 'match', 'notes': 'Rooming list matches pickup report.'}


# ── Excel export row ──────────────────────────────────────────────────────────

def format_export_row(config_row, weekly_row):
    """
    Return a tab-separated string matching the pickup spreadsheet row layout:
      Date | night1 | night2 | ... | total | change | %total | %attrition | webrate

    config_row  — sqlite3.Row from pickup_config
    weekly_row  — sqlite3.Row from pickup_weekly
    """
    try:
        block   = json.loads(config_row['contracted_block'] or '{}')
        pickup  = json.loads(weekly_row['pickup_by_night']  or '{}')
        dates   = sorted(block.keys())

        parts = [weekly_row['report_date']]
        for d in dates:
            v = pickup.get(d)
            parts.append(str(v) if v is not None else '')

        parts.append(str(weekly_row['total_rooms']      or ''))
        parts.append(str(weekly_row['change_from_last'] or ''))

        pob = weekly_row['pct_of_block']
        parts.append(f'{pob:.4f}' if pob is not None else '')

        poa = weekly_row['pct_of_attrition']
        parts.append(f'{poa:.4f}' if poa is not None else '')

        ota = weekly_row['ota_rate']
        parts.append(str(int(ota)) if ota is not None else '')

        return '\t'.join(parts)
    except Exception as e:
        return f'Error: {e}'


# ── Email generators ──────────────────────────────────────────────────────────

_SIGNATURE = ""

def _fmt_date(iso):
    """'2026-01-13' → 'January 13, 2026'"""
    try:
        return datetime.strptime(iso, '%Y-%m-%d').strftime('%B %d, %Y')
    except Exception:
        return iso

def _fmt_short(iso):
    """'2026-01-13' → 'MM/DD/YYYY'"""
    try:
        return datetime.strptime(str(iso)[:10], '%Y-%m-%d').strftime('%m/%d/%Y')
    except Exception:
        return iso or ''

def _build_cc(config_row):
    """Return comma-separated CC email string.
    Includes secondary hotel contact email (CC recipient) plus the cc_emails list."""
    try:
        addrs = []
        # Secondary hotel contact goes first in CC
        hc2_email = (config_row['hotel_contact2_email'] or '').strip()
        if hc2_email:
            addrs.append(hc2_email)
        entries = json.loads(config_row['cc_emails'] or '[]')
        for e in entries:
            if isinstance(e, dict):
                addr = (e.get('email') or '').strip()
            else:
                addr = str(e).strip()
            if addr and addr not in addrs:
                addrs.append(addr)
        return ', '.join(addrs)
    except Exception:
        return ''


def _build_cc_recipients(config_row):
    """Return list of {'name': str, 'email': str} dicts for all CC contacts.
    Includes secondary hotel contact as first entry."""
    try:
        result = []
        seen = set()
        # Secondary hotel contact goes first
        hc2_email = (config_row['hotel_contact2_email'] or '').strip()
        hc2_name  = (config_row['hotel_contact2']       or '').strip()
        if hc2_email and hc2_email not in seen:
            result.append({'name': hc2_name, 'email': hc2_email})
            seen.add(hc2_email)
        entries = json.loads(config_row['cc_emails'] or '[]')
        for e in entries:
            if isinstance(e, dict):
                addr = (e.get('email') or '').strip()
                name = (e.get('name')  or '').strip()
            else:
                addr = str(e).strip()
                name = ''
            if addr and addr not in seen:
                result.append({'name': name, 'email': addr})
                seen.add(addr)
        return result
    except Exception:
        return []


def build_hotel_email(config_row):
    """
    Build the hotel pickup-request email.
    Returns {'to': str, 'subject': str, 'body': str}.
    """
    try:
        block  = json.loads(config_row['contracted_block'] or '{}')
        dates  = sorted(block.keys())

        if dates:
            pre  = int(config_row['shoulder_pre']  or 3)
            post = int(config_row['shoulder_post'] or 3)
            first = datetime.strptime(dates[0], '%Y-%m-%d')
            last  = datetime.strptime(dates[-1], '%Y-%m-%d')
            shoulder_start = (first - timedelta(days=pre)).strftime('%B %d, %Y')
            shoulder_end   = (last  + timedelta(days=post)).strftime('%B %d, %Y')
            date_range = (f"{_fmt_date(dates[0])} – {_fmt_date(dates[-1])}"
                          if len(dates) > 1 else _fmt_date(dates[0]))
            block_lines = '\n'.join(
                f"  {_fmt_date(d)}: {block[d]} rooms contracted"
                for d in dates
            )
        else:
            shoulder_start = shoulder_end = date_range = ''
            block_lines = '  (no contracted block dates entered)'

        org   = config_row['organization'] or ''
        hotel = config_row['hotel'] or ''
        to    = config_row['hotel_contact_email'] or ''

        # Extract first name from hotel_contact for salutation
        _hc = (config_row['hotel_contact'] or '').strip()
        if ',' in _hc:
            _hc_first = _hc.split(',', 1)[1].strip().split()[0]
        elif _hc:
            _hc_first = _hc.split()[0]
        else:
            _hc_first = 'Team'

        subject = f"Pickup Request – {org} / {hotel} / {date_range}"

        body = (
            f"Hello {_hc_first},\n\n"
            f"I hope you are doing well.  Can you please send me the current group pickup report "
            f"and group rooming list for the following group:\n\n"
            f"  Organization: {org}\n"
            f"  Event: {config_row['event_name'] or ''}\n"
            f"  Hotel: {hotel}\n"
            f"  Meeting Dates: {date_range}\n\n"
            f"Contracted Block:\n{block_lines}\n\n"
            f"Please include shoulder nights ({shoulder_start} – {shoulder_end}) "
            f"in both the pickup report and rooming list.\n\n"
            f"Please send:\n"
            f"  1. Day-by-day pickup report\n"
            f"  2. Group rooming list\n\n"
            f"\n"
            f""
        )

        cc = _build_cc(config_row)
        return {'to': to, 'cc': cc, 'subject': subject, 'body': body}
    except Exception as e:
        return {'to': '', 'cc': '', 'subject': 'Pickup Request', 'body': f'Error generating email: {e}'}


def hotel_request_mailto(config_row):
    """Return a mailto: URL for the hotel pickup request (kept for backwards compat)."""
    e = build_hotel_email(config_row)
    return f"mailto:{quote(e['to'])}?subject={quote(e['subject'])}&body={quote(e['body'])}"


def build_hotel_rate_issue_email(config_row, ota_rate, ota_url=None):
    """
    Build the hotel rate-parity issue email.
    Returns {'to': str, 'cc': str, 'subject': str, 'body': str, 'body_html': str}.
    """
    try:
        block  = json.loads(config_row['contracted_block'] or '{}')
        dates  = sorted(block.keys())
        date_range = (f"{_fmt_date(dates[0])} – {_fmt_date(dates[-1])}"
                      if len(dates) > 1 else (_fmt_date(dates[0]) if dates else 'the event dates'))

        org        = config_row['organization'] or ''
        event_name = config_row['event_name'] or org
        hotel      = config_row['hotel'] or ''
        to         = config_row['hotel_contact_email'] or ''
        c_rate     = config_row['contracted_rate']
        ota_url    = ota_url or config_row.get('ota_url') or ''

        _hc = (config_row['hotel_contact'] or '').strip()
        if ',' in _hc:
            _hc_first = _hc.split(',', 1)[1].strip().split()[0]
        elif _hc:
            _hc_first = _hc.split()[0]
        else:
            _hc_first = 'Team'

        c_rate_str   = f"${float(c_rate):.2f}"    if c_rate   else 'N/A'
        ota_rate_str = f"${float(ota_rate):.2f}"  if ota_rate else 'N/A'

        subject = f"Rate Parity Issue – {event_name or org} / {hotel}"

        # Plain-text body
        body = (
            f"Hello {_hc_first},\n\n"
            f"I hope your week is going well so far.\n\n"
            f"We check weekly to ensure that {event_name or org} has the lowest negotiated rates "
            f"for {date_range}. When we checked this week we are finding lower online rates:\n\n"
            f"  Hotel Name:       {hotel}\n"
            f"  Contracted Rate:  {c_rate_str}\n"
            f"  OTA Rate Found:   {ota_rate_str}\n"
        )
        if ota_url:
            body += f"  OTA Link:         {ota_url}\n"
        body += (
            f"\nCan you please look into this and let me know how you would like to resolve "
            f"this rate parity issue?\n\n"
        )

        # HTML body — styled table for Outlook
        P  = 'style="margin:0 0 10px 0;font-family:Arial,sans-serif;font-size:13px;"'
        TD = 'style="padding:6px 12px;border:1px solid #dee2e6;font-family:Arial,sans-serif;font-size:13px;"'
        TH = 'style="padding:6px 12px;border:1px solid #dee2e6;background:#f8f9fa;font-weight:bold;font-family:Arial,sans-serif;font-size:13px;"'

        ota_link_cell = (
            f'<a href="{ota_url}" style="color:#004B97;">{ota_url}</a>'
            if ota_url else '—'
        )

        body_html = (
            f'<p {P}>Hello {_hc_first},</p>'
            f'<p {P}>I hope your week is going well so far.</p>'
            f'<p {P}>We check weekly to ensure that <b>{event_name or org}</b> has the lowest '
            f'negotiated rates for <b>{date_range}</b>. When we checked this week we are finding '
            f'lower online rates:</p>'
            f'<table style="border-collapse:collapse;margin-bottom:14px;">'
            f'<thead><tr>'
            f'<th {TH}>Hotel Name</th>'
            f'<th {TH}>Contracted Rate</th>'
            f'<th {TH}>OTA Rate Found</th>'
            f'<th {TH}>OTA Link</th>'
            f'</tr></thead>'
            f'<tbody><tr>'
            f'<td {TD}>{hotel}</td>'
            f'<td {TD}>{c_rate_str}</td>'
            f'<td {TD} style="color:#dc3545;font-weight:bold;padding:6px 12px;border:1px solid #dee2e6;">{ota_rate_str}</td>'
            f'<td {TD}>{ota_link_cell}</td>'
            f'</tr></tbody></table>'
            f'<p {P}>Can you please look into this and let me know how you would like to resolve '
            f'this rate parity issue?</p>'
        )

        cc = _build_cc(config_row)
        return {'to': to, 'cc': cc, 'subject': subject, 'body': body, 'body_html': body_html}
    except Exception as e:
        return {'to': '', 'cc': '', 'subject': 'Rate Parity Issue',
                'body': f'Error generating email: {e}', 'body_html': ''}


def build_client_email(config_row, weekly_row, rl_status=None, weekly_list=None):
    """
    Build the client pickup-summary email with a spreadsheet-style grid.
    Returns {'to': str, 'subject': str, 'body': str}.
    weekly_row may be None if no entries exist yet.
    weekly_list is the full list of weekly entry dicts (most-recent first).
    """
    try:
        block      = json.loads(config_row['contracted_block'] or '{}')
        org        = config_row['organization'] or ''
        event_name = config_row['event_name'] or ''
        hotel      = config_row['hotel'] or ''
        to         = config_row['group_contact_email'] or ''

        # Extract first name from group_contact for salutation
        # Handles "First Last", "Last, First", or a single name
        _gc = (config_row['group_contact'] or '').strip()
        if ',' in _gc:
            # "Last, First" format — take everything after the comma
            _first = _gc.split(',', 1)[1].strip().split()[0]
        elif _gc:
            _first = _gc.split()[0]
        else:
            _first = 'Team'

        if not weekly_row:
            subject = f"Pickup Update – {event_name or org} | {hotel}"
            body = (
                f"Hello {_first},\n\n"
                f"I wanted to touch base regarding the upcoming group at {hotel}.\n"
                f"No pickup data has been recorded yet — I will follow up once the "
                f"first report is available.\n\n"
                f""
            )
            return {'to': to, 'subject': subject, 'body': body}

        pickup    = json.loads(weekly_row['pickup_by_night'] or '{}')
        report_dt = weekly_row['report_date'] or ''
        # Include shoulder nights (in pickup but not in contracted block)
        dates     = sorted(set(block.keys()) | set(pickup.keys()))
        total_blk = sum(block.get(d, 0) for d in dates)
        ota       = weekly_row['ota_rate']
        c_rate    = config_row['contracted_rate']
        _atr_pct_dec = config_row['attrition_pct'] or 0  # decimal e.g. 0.80
        _atr_rooms   = total_blk * float(_atr_pct_dec) if total_blk else 0

        # Always recompute totals/pcts from pickup_by_night (source of truth)
        def _recompute(w):
            """Return a plain dict with corrected total_rooms / pct_of_block /
               pct_of_attrition.  change_from_last is filled in a second pass."""
            pbn = json.loads(w['pickup_by_night'] or '{}')
            t = sum(v for v in pbn.values() if v is not None and v != '')
            pob_ = round(t / total_blk * 100, 1) if total_blk else None
            poa_ = round(t / _atr_rooms * 100, 1) if _atr_rooms else None
            d = dict(w)
            d['total_rooms']      = t
            d['pct_of_block']     = pob_
            d['pct_of_attrition'] = poa_
            return d

        wr = _recompute(weekly_row)
        total = wr['total_rooms']
        pob   = wr['pct_of_block']
        poa   = wr['pct_of_attrition']

        pob_str = f"{pob:.1f}%" if pob is not None else 'N/A'
        poa_str = f"{poa:.1f}%" if poa is not None else 'N/A'

        # Attrition balance — rooms still needed to reach the attrition floor
        if _atr_pct_dec and total_blk:
            _atr_target  = round(total_blk * float(_atr_pct_dec))
            _atr_balance = _atr_target - total
            if _atr_balance > 0:
                _atr_note = f"  ({_atr_balance} room-nights still needed to reach attrition target)"
            else:
                _atr_note = f"  (attrition target reached ✓)"
        else:
            _atr_note = ""

        if pob is not None and pob >= 80:
            status_line = "Your group is ON PACE with the contracted block."
        elif pob is not None and pob >= 60:
            status_line = "Your group is slightly BELOW PACE — monitoring closely."
        else:
            status_line = ""

        # ── Build label-per-line format (works in any font, no alignment needed) ──
        def fmt_night(iso):
            try:
                dt = datetime.strptime(iso, '%Y-%m-%d')
                return dt.strftime('%a %m/%d/%y')
            except Exception:
                return iso

        grid = 'NIGHT-BY-NIGHT PICKUP\n'
        for d in dates:
            b = block.get(d, 0) or 0
            p = pickup.get(d, 0) or 0
            night = fmt_night(d)
            if not b:
                grid += f"  {night}:  {p} rooms  (shoulder)\n"
            else:
                diff = p - b
                diff_s = f'+{diff}' if diff > 0 else str(diff)
                grid += f"  {night}:  Block {b} / Pickup {p}  ({diff_s})\n"
        total_diff = total - total_blk
        total_diff_s = f'+{total_diff}' if total_diff > 0 else str(total_diff)
        grid += f"  Total:  Block {total_blk:,} / Pickup {total:,}  ({total_diff_s})\n"
        grid += '\n'

        # Weekly history — recompute all entries from pickup_by_night then
        # fill change_from_last correctly (oldest → newest pass)
        raw_entries = weekly_list or [dict(weekly_row)]
        # Sort oldest-first for WoW calculation, then flip back
        recomputed = [_recompute(w) for w in reversed(list(raw_entries))]
        prev_t = None
        for e in recomputed:
            e['change_from_last'] = (e['total_rooms'] - prev_t) if prev_t is not None else None
            prev_t = e['total_rooms']
        recomputed.reverse()   # back to most-recent-first
        entries = recomputed
        history_lines = []
        for w in entries:
            w_total  = w['total_rooms']      if isinstance(w, dict) else w.get('total_rooms')
            w_change = w['change_from_last'] if isinstance(w, dict) else w.get('change_from_last')
            w_pob    = w['pct_of_block']     if isinstance(w, dict) else w.get('pct_of_block')
            w_poa    = w['pct_of_attrition'] if isinstance(w, dict) else w.get('pct_of_attrition')
            w_date   = (w['report_date']     if isinstance(w, dict) else w.get('report_date')) or ''
            w_label  = w.get('label', '') if isinstance(w, dict) else ''
            try:
                w_date_fmt = datetime.strptime(w_date, '%Y-%m-%d').strftime('%m/%d/%y')
            except Exception:
                w_date_fmt = w_date
            chg_str = (f'+{w_change}' if w_change and w_change > 0
                       else str(w_change)) if w_change is not None else '—'
            pob_s = f'{w_pob:.1f}%'  if w_pob  is not None else '—'
            poa_s = f'{w_poa:.1f}%'  if w_poa  is not None else '—'
            tot_s = f'{w_total:,}'   if w_total is not None else '—'
            lbl   = f' ({w_label})'  if w_label else ''
            history_lines.append(
                f"  {w_date_fmt}{lbl}:  {tot_s} rooms  ({chg_str})  |  {pob_s} of block  |  {poa_s} of attrition\n"
            )

        history_block = 'WEEKLY HISTORY\n' + ''.join(history_lines)

        # OTA note intentionally omitted from client email (internal use only)

        # Rooming list note
        rl_note = ''
        if rl_status == 'match':
            rl_note = '\nRooming list verified — matches pickup report.'
        elif rl_status == 'discrepancy':
            rl_note = '\nRooming list discrepancy noted — following up with hotel.'

        subject = f"Pickup Update – {event_name or org} | {hotel} | As of {_fmt_short(report_dt)}"

        try:
            # Use the recomputed entries list — first entry is most recent
            _wow = entries[0]['change_from_last'] if entries else None
        except Exception:
            _wow = None
        wow_s = (f'+{_wow}' if _wow and _wow > 0 else str(_wow)) if _wow is not None else '—'

        body = (
            f"Hello {_first},\n\n"
            f"Here is your weekly pickup update for {event_name or org} at {hotel}.\n"
            f"Report date: {_fmt_short(report_dt)}   |   Cut-off: {_fmt_short(config_row['cutoff_date']) if config_row['cutoff_date'] else 'N/A'}\n\n"
            f"SUMMARY\n"
            f"  Total Rooms:      {total:,} of {total_blk:,}\n"
            f"  % of Block:       {pob_str}\n"
            f"  % of Attrition:   {poa_str}{_atr_note}\n"
            f"  Week-over-Week:   {wow_s}\n\n"
            + (f"{status_line}\n\n" if status_line else "")
            + f"{grid}\n"
            + f"{history_block}\n"
            + (f"{rl_note}\n" if rl_note else "")
            + "\nPlease feel free to reach out with any questions.\n\n"
        )

        # ── HTML version — spreadsheet-style table matching weekly pickup history ─
        F   = 'font-family:Calibri,Arial,sans-serif; font-size:10pt;'
        B   = f'{F} border:1px solid #aaa; padding:3px 6px;'
        def td(extra=''):  return f'style="{B} {extra}"'
        def th(extra=''):  return f'style="{B} background:#dce6f1; font-weight:bold; {extra}"'
        def tf(extra=''):  return f'style="{B} background:#f2f2f2; font-weight:bold; {extra}"'

        def diff_color(d):
            return 'color:green;' if d >= 0 else 'color:#cc0000;'

        # Build the single combined spreadsheet table:
        # Rows: header (dates) | Day | Block | [weekly entries, newest first] | Remaining
        # Columns: label | date1 | date2 | ... | Total | WoW | % Block | % Attr
        n = len(dates)
        summary_hdrs = ['Total', 'WoW', '% Block', '% Attr']
        TS = f'style="border-collapse:collapse; {F}"'

        # Header row — dates
        date_headers = ''.join(
            f'<th {th("text-align:center;")}>'
            + (datetime.strptime(d, "%Y-%m-%d").strftime("%m/%d") if len(d) == 10 else d)
            + '</th>'
            for d in dates
        )
        summary_headers = ''.join(f'<th {th("text-align:right;")}>{h}</th>' for h in summary_hdrs)
        header_row = f'<tr><th {th()}></th>{date_headers}{summary_headers}</tr>'

        # Day-of-week row
        day_cells = ''.join(
            f'<td {td("text-align:center; color:#555;")}>'
            + (datetime.strptime(d, "%Y-%m-%d").strftime("%a") if len(d) == 10 else '')
            + '</td>'
            for d in dates
        )
        day_row = f'<tr><td {td("color:#555;")}><i>Day</i></td>{day_cells}<td {td()} colspan="{len(summary_hdrs)}"></td></tr>'

        # Block row
        block_cells = ''.join(
            f'<td {td("text-align:right;")}>{block.get(d,0) or "—"}</td>'
            for d in dates
        )
        block_row = (
            f'<tr><td {tf()}>Block</td>{block_cells}'
            f'<td {tf("text-align:right;")}>{total_blk:,}</td>'
            f'<td {tf()} colspan="{len(summary_hdrs)-1}"></td></tr>'
        )

        # Weekly data rows (entries already ordered most-recent first)
        data_rows = []
        for w in entries:
            w_pickup = json.loads(w['pickup_by_night'] if isinstance(w, dict)
                                  else w.get('pickup_by_night') or '{}')
            w_total  = w['total_rooms']      if isinstance(w, dict) else w.get('total_rooms')
            w_change = w['change_from_last'] if isinstance(w, dict) else w.get('change_from_last')
            w_pob    = w['pct_of_block']     if isinstance(w, dict) else w.get('pct_of_block')
            w_poa    = w['pct_of_attrition'] if isinstance(w, dict) else w.get('pct_of_attrition')
            w_date   = (w['report_date']     if isinstance(w, dict) else w.get('report_date')) or ''
            w_label  = w.get('label', '')    if isinstance(w, dict) else ''
            try:
                w_date_fmt = datetime.strptime(w_date, '%Y-%m-%d').strftime('%m/%d/%y')
            except Exception:
                w_date_fmt = w_date
            lbl = f' ({w_label})' if w_label else ''
            chg_str = (f'+{w_change}' if w_change and w_change > 0 else str(w_change)) if w_change is not None else '—'
            pob_s   = f'{w_pob:.1f}%' if w_pob  is not None else '—'
            poa_s   = f'{w_poa:.1f}%' if w_poa  is not None else '—'
            tot_s   = f'{w_total:,}'  if w_total is not None else '—'
            chg_dc  = diff_color(w_change) if w_change is not None else ''
            night_cells = ''.join(
                f'<td {td("text-align:right;")}>{w_pickup.get(d, "—")}</td>'
                for d in dates
            )
            data_rows.append(
                f'<tr>'
                f'<td {td()}>{w_date_fmt}{lbl}</td>'
                + night_cells +
                f'<td {td("text-align:right; font-weight:bold;")}>{tot_s}</td>'
                f'<td style="{B} text-align:right; {chg_dc}">{chg_str}</td>'
                f'<td {td("text-align:right;")}>{pob_s}</td>'
                f'<td {td("text-align:right;")}>{poa_s}</td>'
                f'</tr>'
            )

        # Remaining row (block − latest pickup, per night)
        rem_cells = ''.join(
            f'<td {td("text-align:right;")}>{(block.get(d,0) or 0) - (pickup.get(d) or 0)}</td>'
            for d in dates
        )
        rem_total = total_blk - total
        rem_dc = diff_color(-rem_total)  # negative remaining = over block = green
        rem_row = (
            f'<tr><td {td("color:#555;")}><i>Remaining</i></td>{rem_cells}'
            f'<td style="{B} text-align:right; font-style:italic; {rem_dc}">{rem_total:+,}</td>'
            f'<td {td()} colspan="{len(summary_hdrs)-1}"></td></tr>'
        )

        pickup_table = (
            f'<table {TS} border="1" cellpadding="0" cellspacing="0">'
            f'<thead>{header_row}{day_row}</thead>'
            f'<tbody>{block_row}{"".join(data_rows)}{rem_row}</tbody>'
            f'</table>'
        )

        rl_html = f'<p style="{F}">{rl_note.strip()}</p>' if rl_note else ''
        status_html = (f'<p style="{F} font-weight:bold;">{status_line}</p>'
                       if status_line else '')
        P = f'style="{F} margin:6px 0;"'

        html_body = (
            f'<p {P}>Hello {_first},</p>'
            f'<p {P}>Here is your weekly pickup update for <b>{event_name or org}</b> at <b>{hotel}</b>.<br>'
            f'Report date: {report_dt}&nbsp;&nbsp;|&nbsp;&nbsp;Cut-off: {config_row["cutoff_date"] or "N/A"}</p>'
            f'<p {P}><b>SUMMARY</b><br>'
            f'Total Rooms: <b>{total:,}</b> of {total_blk:,}<br>'
            f'% of Block: <b>{pob_str}</b><br>'
            f'% of Attrition: <b>{poa_str}</b>{_atr_note}<br>'
            f'Week-over-Week: <b>{wow_s}</b></p>'
            + status_html
            + f'<p {P}><b>WEEKLY PICKUP REPORT</b></p>'
            + pickup_table
            + rl_html
            + f'<p {P}>Please feel free to reach out with any questions.</p>'
        )

        cc = _build_cc(config_row)
        return {'to': to, 'cc': cc, 'subject': subject, 'body': body, 'html_body': html_body}
    except Exception as e:
        return {'to': '', 'cc': '', 'subject': 'Pickup Update', 'body': f'Error generating email: {e}'}


def client_summary_mailto(config_row, weekly_row, rl_status=None):
    """Return a mailto: URL for the client pickup summary (kept for backwards compat)."""
    e = build_client_email(config_row, weekly_row, rl_status)
    return f"mailto:{quote(e['to'])}?subject={quote(e['subject'])}&body={quote(e['body'])}"


# ── Contract Template Parsing ─────────────────────────────────────────────────

_OPTIONAL_SECTION_KEYWORDS = [
    'meeting room', 'meeting space', 'food', 'beverage', 'f&b', 'catering',
    'exhibit', 'banquet', 'function space', 'audio', 'visual', 'av ',
]


def extract_template_metadata(file_bytes):
    """
    Parse a Word contract template (.docx or .docm) and extract:
      - merge_fields: list of MERGEFIELD names found in XML
      - sections: list of {name, index, optional} dicts
    """
    import zipfile
    try:
        from lxml import etree as _et
    except ImportError:
        return {'merge_fields': [], 'sections': []}

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            names = z.namelist()
            doc_name = 'word/document.xml' if 'word/document.xml' in names else next(
                (n for n in names if 'document' in n and n.endswith('.xml')), None)
            if doc_name is None:
                return {'merge_fields': [], 'sections': []}
            xml_content = z.read(doc_name).decode('utf-8', errors='replace')
    except Exception:
        return {'merge_fields': [], 'sections': []}

    fields = re.findall(r'MERGEFIELD\s+"?([A-Za-z_][A-Za-z0-9_ ]*)"?', xml_content)
    fields += re.findall(r'MERGEFIELD\s+([A-Za-z_][A-Za-z0-9_]+)\s', xml_content)
    merge_fields = sorted(set(f.strip().replace(' ', '_') for f in fields if f.strip()))

    sections = []
    idx = 0
    try:
        root = _et.fromstring(xml_content.encode('utf-8'))
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for para in root.findall('.//w:p', ns):
            runs = para.findall('.//w:r', ns)
            text = ''.join(r.findtext('w:t', '', ns) for r in runs).strip()
            if not text or len(text) < 4:
                continue
            pPr = para.find('w:pPr', ns)
            style_val = ''
            if pPr is not None:
                pStyle = pPr.find('w:pStyle', ns)
                if pStyle is not None:
                    style_val = pStyle.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
            is_heading_style = 'heading' in style_val.lower()
            is_bold = len(para.findall('.//w:b', ns)) > 0
            leading = text.split(':')[0].strip()
            all_words = leading.split()
            consecutive_caps = 0
            for w in all_words:
                alpha = re.sub(r'[^A-Za-z]', '', w)
                if alpha and len(alpha) >= 2 and alpha.isupper():
                    consecutive_caps += 1
                else:
                    break
            is_caps_heading = consecutive_caps >= 2 and is_bold
            if is_heading_style or is_caps_heading:
                section_name = leading if len(leading) >= 4 else text
                skip_phrases = ['in witness whereof', 'hotel use only', 'undersigned expressly']
                if re.match(r'^[X☐☑✓]\s*[-–]', section_name):
                    continue
                if any(sp in section_name.lower() for sp in skip_phrases):
                    continue
                tl = section_name.lower()
                optional = any(kw in tl for kw in _OPTIONAL_SECTION_KEYWORDS)
                sections.append({'name': section_name, 'index': idx, 'optional': optional})
                idx += 1
    except Exception:
        pass

    return {'merge_fields': merge_fields, 'sections': sections}


# ── Cvent RFP Document Parser ─────────────────────────────────────────────────

def parse_rfp_docx(file_bytes):
    """
    Parse a Cvent RFP Word document (.docx) and extract key fields.
    Returns a dict with any of these keys populated (None if not found):
        rfp_name, rfp_code, event_name, client_org,
        start_date, end_date, response_due_date, decision_due_date,
        total_attendees, f_and_b_budget, total_room_nights, peak_rooms
    """
    import io as _io
    import re as _re
    import zipfile
    from lxml import etree as _et
    from datetime import datetime as _dt

    try:
        with zipfile.ZipFile(_io.BytesIO(file_bytes)) as z:
            xml = z.read('word/document.xml')
    except Exception:
        return {}

    root = _et.fromstring(xml)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    lines = []
    for p in root.findall('.//w:p', ns):
        line = ''.join(t.text or '' for t in p.findall('.//w:t', ns)).strip()
        if line:
            lines.append(line)

    def _after(label_keywords, max_gap=2):
        for i, l in enumerate(lines):
            ll = l.lower()
            if all(kw.lower() in ll for kw in label_keywords):
                for j in range(1, max_gap + 1):
                    if i + j < len(lines):
                        v = lines[i + j].strip()
                        if v:
                            return v
        return None

    def _parse_date(s):
        if not s:
            return None
        # Strip day-of-week prefix e.g. "Wed, "
        s = _re.sub(r'^[A-Za-z]{3},\s*', '', s.strip())
        # Strip trailing annotations like " + 2 alternate dates"
        s = _re.sub(r'\s*\+.*$', '', s).strip()
        for fmt in ('%b %d, %Y', '%B %d, %Y', '%m/%d/%Y', '%Y-%m-%d'):
            try:
                return _dt.strptime(s.strip(), fmt).strftime('%Y-%m-%d')
            except ValueError:
                pass
        return None

    def _parse_date_range(s):
        if not s:
            return None, None
        # Strip trailing annotations before splitting
        s = _re.sub(r'\s*\+.*$', '', s).strip()
        parts = _re.split(r'\s*[-–]\s*', s, maxsplit=1)
        return _parse_date(parts[0]), _parse_date(parts[1]) if len(parts) > 1 else None

    def _parse_currency(s):
        if not s:
            return None
        m = _re.search(r'[\d,]+\.?\d*', s)
        return float(m.group().replace(',', '')) if m else None

    def _parse_int(s):
        if not s:
            return None
        m = _re.search(r'\d[\d,]*', s)
        return int(m.group().replace(',', '')) if m else None

    rfp_name          = _after(['RFP Name'])
    rfp_code          = _after(['RFP Code'])
    response_due      = _parse_date(_after(['Response Due Date']))
    decision_due      = _parse_date(_after(['Decision Due Date']))
    total_attendees   = _parse_int(_after(['Total Attendees']))
    f_and_b_budget    = _parse_currency(_after(['Food and Beverage Budget']))
    total_room_nights = _parse_int(_after(['Total Room Nights']))
    peak_rooms        = _parse_int(_after(['Peak Room Nights']))
    client_org        = _after(['Organization Name'])

    # Prefer "Planner Preferred" row over "Event Dates" (avoids "+ N alternate dates" suffix)
    preferred_raw  = _after(['Planner Preferred'])
    event_dates_raw = _after(['Event Dates'])
    start_date, end_date = _parse_date_range(preferred_raw or event_dates_raw)

    # Collect alternate date rows (multiple "Alternate Date" labels may appear)
    alt_dates = []
    for i, line in enumerate(lines):
        if line.strip().lower() == 'alternate date':
            for j in range(1, 3):
                if i + j < len(lines):
                    v = lines[i + j].strip()
                    if v and _re.search(r'\d{4}', v):
                        alt_dates.append(v)
                        break
    alt_start_date,   alt_end_date   = _parse_date_range(alt_dates[0]) if len(alt_dates) > 0 else (None, None)
    alt_start_date_2, alt_end_date_2 = _parse_date_range(alt_dates[1]) if len(alt_dates) > 1 else (None, None)
    alt_start_date_3, alt_end_date_3 = _parse_date_range(alt_dates[2]) if len(alt_dates) > 2 else (None, None)

    return {
        'rfp_name':          rfp_name,
        'rfp_code':          rfp_code,
        'event_name':        rfp_name,
        'client_org':        client_org,
        'response_due_date': response_due,
        'decision_due_date': decision_due,
        'total_attendees':   total_attendees,
        'f_and_b_budget':    f_and_b_budget,
        'total_room_nights': total_room_nights,
        'peak_rooms':        peak_rooms,
        'start_date':        start_date,
        'end_date':          end_date,
        'alt_start_date':    alt_start_date,
        'alt_end_date':      alt_end_date,
        'alt_start_date_2':  alt_start_date_2,
        'alt_end_date_2':    alt_end_date_2,
        'alt_start_date_3':  alt_start_date_3,
        'alt_end_date_3':    alt_end_date_3,
    }


# ── Cvent CRF (Consolidated Response Form) Excel Parser ──────────────────────

def parse_crf_excel(file_bytes):
    """
    Parse a Cvent Consolidated Response Form (CRF) Excel file.
    Returns:
        {'rfp_meta': {event_name, response_due_date, decision_due_date},
         'hotels': [{hotel_name, city, state, status, proposed_rate,
                     f_and_b_minimum, commission_pct, attrition_pct,
                     cutoff_days, concessions, notes, contact_name,
                     contact_email, contact_phone, contact_title, crf_row_data}]}
    """
    import io as _io
    import re as _re
    import json as _json
    import openpyxl

    wb = openpyxl.load_workbook(_io.BytesIO(file_bytes), read_only=True, data_only=True)

    def _sv(v):
        if v is None:
            return ''
        s = str(v).replace('_x000d_', ' ').replace('\r', ' ')
        return _re.sub(r'\s+', ' ', s).strip()

    def _parse_currency(val):
        s = _sv(val)
        if not s or s.lower() in ('n/a', 'na', '—', '-'):
            return None
        cleaned = _re.sub(r'[USD$,]', '', s)
        m = _re.search(r'[\d]+\.?\d*', cleaned)
        return float(m.group()) if m else None

    def _parse_commission(val):
        s = _sv(val)
        m = _re.search(r'(\d+(?:\.\d+)?)\s*%', s)
        return float(m.group(1)) / 100.0 if m else None

    def _parse_address(val):
        if val is None:
            return None, None
        raw = str(val).replace('_x000d_', '\n').replace('\r\n', '\n').replace('\r', '\n')
        lines = [ln.strip() for ln in raw.split('\n') if ln.strip() and ln.strip().upper() != 'USA']
        city_state_line = None
        for line in lines:
            if _re.search(r'[A-Za-z].+,\s*[A-Za-z]', line):
                city_state_line = line
                break
        if not city_state_line:
            city_state_line = ' '.join(lines)
        parts = [p.strip() for p in city_state_line.split(',')]
        if len(parts) >= 2:
            city  = parts[0].strip()
            state = _re.sub(r'\s+[\d\-]+\s*$', '', parts[1].strip()).strip()
            state = _re.sub(r'\s*(USA|US|United States)\s*$', '', state, flags=_re.I).strip()
            return city, state
        return None, None

    def _map_status(val):
        s = _sv(val).lower()
        if 'submitted' in s or 'awarded' in s or 'selected' in s:
            return 'proposal_received'
        if 'turned down' in s or 'withdrawn' in s or 'declined' in s:
            return 'declined'
        return 'pending'

    # ── RFP metadata sheet ────────────────────────────────────────────────────
    rfp_meta = {}
    if 'RFP' in wb.sheetnames:
        rfp_ws = wb['RFP']
        rfp_rows = list(rfp_ws.rows)
        for row in rfp_rows[:25]:
            vals = [c.value for c in row]
            if len(vals) < 2:
                continue
            label = _sv(vals[0]).lower()
            val = next((_sv(v) for v in vals[1:4] if _sv(v)), '')
            if ('response' in label and 'due' in label) and val and len(val) <= 40:
                rfp_meta['response_due_date'] = val
            elif 'decision' in label and 'due' in label and val and len(val) <= 40:
                rfp_meta['decision_due_date'] = val
        for row in rfp_rows[:5]:
            for cell in row:
                v = _sv(cell.value)
                if v and len(v) > 8 and 'request for proposal' not in v.lower() and 'rfp' not in v.lower()[:3]:
                    rfp_meta.setdefault('event_name', v)
                    break

    # ── Summary sheet ─────────────────────────────────────────────────────────
    if 'Summary' not in wb.sheetnames:
        return {'rfp_meta': rfp_meta, 'hotels': []}

    ws = wb['Summary']
    all_rows = list(ws.rows)

    header_row_idx = None
    for i, row in enumerate(all_rows[:10]):
        if _sv(row[0].value).lower() == 'hotel':
            header_row_idx = i
            break
    if header_row_idx is None:
        return {'rfp_meta': rfp_meta, 'hotels': []}

    headers = [_sv(all_rows[header_row_idx][i].value) for i in range(len(all_rows[header_row_idx]))]

    def _find_col(*keywords):
        for j, h in enumerate(headers):
            hl = h.lower()
            if all(kw.lower() in hl for kw in keywords):
                return j
        return None

    col_address    = _find_col('address')
    col_status     = _find_col('status')
    col_notes      = _find_col('notes')
    col_rate       = _find_col('room rate')
    col_fab        = _find_col('f&b') or _find_col('food') or _find_col('beverage') or _find_col('minimum')
    col_attrition  = _find_col('80%') or _find_col('attrition')
    col_cutoff     = _find_col('21 day') or _find_col('cut off') or _find_col('cutoff')
    col_commission = _find_col('commission') or _find_col('10%')

    def _gv(vals, idx):
        if idx is None or idx >= len(vals):
            return None
        return vals[idx]

    _SEVEN_PCT_BRANDS = [
        'marriott','westin','sheraton','w hotel','jw marriott','renaissance',
        'courtyard','residence inn','fairfield','delta hotel','tribute portfolio',
        'autograph collection','hilton','doubletree','embassy suites','hampton inn',
        'homewood','home2','curio collection','tapestry','waldorf astoria',
        'conrad','canopy','tempo by hilton','signia','graduate',
    ]
    _CONC_KEYWORDS = [
        'resort fee','comp room','1:40','1 per 40','parking','wifi',
        'wireless','internet','fitness','suite','amenity','projector','audio',
        'profit calc','f&b profit',
    ]

    skip_labels = {'available','not available','hotel',''}
    hotels = []

    for row in all_rows[header_row_idx + 1:]:
        vals = [c.value for c in row]
        hotel_name = _sv(vals[0]) if vals else ''
        if not hotel_name or hotel_name.lower() in skip_labels:
            continue
        if sum(1 for v in vals if v is not None) <= 2:
            continue

        city, state   = _parse_address(_gv(vals, col_address))
        status        = _map_status(_gv(vals, col_status))
        notes         = _sv(_gv(vals, col_notes))
        proposed_rate = _parse_currency(_gv(vals, col_rate))
        f_and_b_min   = _parse_currency(_gv(vals, col_fab))
        commission_pct = _parse_commission(_gv(vals, col_commission))

        if commission_pct is None:
            comm_raw = _sv(_gv(vals, col_commission)).lower()
            if comm_raw.startswith('yes') or '10%' in comm_raw or comm_raw == '10':
                commission_pct = 0.10
            elif any(b in hotel_name.lower() for b in _SEVEN_PCT_BRANDS):
                commission_pct = 0.07

        attrition_raw = _sv(_gv(vals, col_attrition)).lower()
        attrition_pct = 0.80 if attrition_raw.startswith('yes') else None
        cutoff_raw  = _sv(_gv(vals, col_cutoff)).lower()
        cutoff_days = 21 if cutoff_raw.startswith('yes') else None

        concession_items = []
        for j, (h, v) in enumerate(zip(headers, vals)):
            if not h or v is None:
                continue
            hl, vl = h.lower(), _sv(v).lower()
            if not any(kw in hl for kw in _CONC_KEYWORDS):
                continue
            if not vl or vl in ('no','n/a','na','—','-'):
                continue
            if vl.startswith('yes'):
                short_h = h.split('?')[0].strip()[:60]
                extra = _sv(v)[3:].strip(' ,:') if len(_sv(v)) > 3 else ''
                concession_items.append(f'✓ {short_h}: {extra}' if extra else f'✓ {short_h}')
        concessions = '\n'.join(concession_items) if concession_items else None

        crf_row_data = {h: _sv(v)[:400] for h, v in zip(headers, vals) if h and v is not None}

        hotel_dict = {
            'hotel_name': hotel_name, 'city': city, 'state': state,
            'status': status, 'proposed_rate': proposed_rate,
            'f_and_b_minimum': f_and_b_min, 'commission_pct': commission_pct,
            'attrition_pct': attrition_pct, 'cutoff_days': cutoff_days,
            'concessions': concessions, 'notes': notes or None,
            'contact_name': None, 'contact_email': None,
            'contact_phone': None, 'contact_title': None,
            'crf_row_data': _json.dumps(crf_row_data),
        }

        # Pull contact info from matching hotel tab
        h_words = set(w.lower() for w in hotel_name.split() if len(w) > 3)
        for sheet_name in wb.sheetnames:
            if sheet_name.lower() in ('summary', 'rfp'):
                continue
            s_words = set(w.lower() for w in sheet_name.split() if len(w) > 3)
            if h_words & s_words:
                try:
                    h_ws = wb[sheet_name]
                    for h_row in list(h_ws.rows)[5:22]:
                        h_vals = [c.value for c in h_row]
                        if len(h_vals) < 2:
                            continue
                        lbl = _sv(h_vals[0]).lower()
                        val = next((_sv(v) for v in h_vals[1:3] if v is not None and _sv(v)), '')
                        if not val:
                            continue
                        if 'contact name' in lbl or lbl == 'contact':
                            hotel_dict['contact_name'] = val
                        elif 'title' in lbl and 'hotel' not in lbl:
                            hotel_dict['contact_title'] = val
                        elif 'email' in lbl:
                            hotel_dict['contact_email'] = val
                        elif 'phone' in lbl or 'telephone' in lbl:
                            hotel_dict['contact_phone'] = val
                except Exception:
                    pass
                break

        hotels.append(hotel_dict)

    # ── Fallback: read individual hotel sheets when Summary has no data rows ──
    # This happens when the CRF was exported before hotels filled in the Summary
    # grid (e.g. single-hotel CRFs, or early exports). Each hotel tab still has
    # its full proposal data — Proposal Status, rates, contact info, address.
    if not hotels:
        for sheet_name in wb.sheetnames:
            if sheet_name.lower() in ('summary', 'rfp'):
                continue
            try:
                h_ws = wb[sheet_name]
                h_rows = list(h_ws.rows)
                if len(h_rows) < 5:
                    continue
                kv = {}
                kv_all = []
                for h_row in h_rows:
                    h_vals = [c.value for c in h_row]
                    if len(h_vals) < 2:
                        continue
                    lbl = _sv(h_vals[0]).lower().strip().rstrip(':')
                    all_vals = [_sv(v) for v in h_vals[1:6] if v is not None and _sv(v)]
                    val = all_vals[0] if all_vals else ''
                    if lbl and val:
                        kv[lbl] = val
                    if lbl:
                        kv_all.append((lbl, all_vals))

                status_raw = kv.get('proposal status') or kv.get('status') or ''
                if not status_raw:
                    continue
                status = _map_status(status_raw)

                hotel_name = sheet_name.strip()

                raw_addr = kv.get('address') or kv.get('hotel address') or ''
                city, state = _parse_address(raw_addr)
                if city and (len(city) > 40 or any(c.isdigit() for c in city)):
                    import re as _re_addr
                    m = _re_addr.search(
                        r'([A-Za-z][A-Za-z\s\-\.]{1,30}),\s*([A-Za-z][A-Za-z\s]{2,20})',
                        raw_addr)
                    if m:
                        city  = m.group(1).strip().split()[-1] if ' ' in m.group(1) else m.group(1).strip()
                        state = m.group(2).strip().split(',')[0].strip()

                proposed_rate = _parse_currency(
                    kv.get('room rate') or kv.get('single room rate') or kv.get('rate') or
                    kv.get('rates') or '')
                if not proposed_rate:
                    _RATE_LBLS = ('rate', 'price', 'preferred', 'alternate', 'dates')
                    _SKIP_LBLS = ('information', 'introduction', 'note', 'description',
                                  'total', 'tax', 'address', 'phone', 'email', 'name')
                    for _lbl, _vals in kv_all:
                        if any(s in _lbl for s in _SKIP_LBLS):
                            continue
                        for _v in _vals:
                            if _v and ('usd' in _v.lower() or _v.strip().startswith('$')):
                                r = _parse_currency(_v)
                                if r and 50 < r < 5000:
                                    proposed_rate = r
                                    break
                        if proposed_rate:
                            break

                commission_pct = _parse_commission(
                    kv.get('commission rate') or kv.get('commission') or
                    kv.get('commissionable') or '')
                f_and_b_minimum = _parse_currency(kv.get('f&b minimum') or kv.get('food & beverage minimum') or '')
                contact_name  = kv.get('contact name') or kv.get('contact') or None
                contact_email = (kv.get('email address') or kv.get('email') or
                                 kv.get('contact email') or kv.get('e-mail') or None)
                contact_phone = (kv.get('phone') or kv.get('telephone') or
                                 kv.get('contact phone') or kv.get('mobile') or None)
                contact_title = kv.get('title') or None
                notes         = kv.get('notes') or kv.get('recommendations') or None

                if commission_pct is None:
                    _SEVEN_PCT_BRANDS = [
                        'marriott','westin','sheraton','hilton','doubletree','embassy',
                        'hampton','homewood','curio','tapestry','waldorf','renaissance',
                        'autograph','courtyard','residence','fairfield','home2',
                    ]
                    if any(b in hotel_name.lower() for b in _SEVEN_PCT_BRANDS):
                        commission_pct = 0.07

                _CC_KEYWORDS = ['convention center','convention ctr','exhibition hall',
                                'exhibit hall','event center','civic center',
                                'conference center','arena','coliseum']
                venue_type = 'convention_center' if any(
                    kw in hotel_name.lower() for kw in _CC_KEYWORDS) else 'hotel'

                hotels.append({
                    'hotel_name':      hotel_name,
                    'city':            city,
                    'state':           state,
                    'status':          status,
                    'proposed_rate':   proposed_rate,
                    'f_and_b_minimum': f_and_b_minimum,
                    'commission_pct':  commission_pct,
                    'attrition_pct':   None,
                    'cutoff_days':     None,
                    'concessions':     None,
                    'notes':           notes,
                    'contact_name':    contact_name,
                    'contact_email':   contact_email,
                    'contact_phone':   contact_phone,
                    'contact_title':   contact_title,
                    'crf_row_data':    json.dumps(kv),
                    'venue_type':      venue_type,
                    'rental_fee':      None,
                })
            except Exception:
                continue

    return {'rfp_meta': rfp_meta, 'hotels': hotels}


def parse_cc_contract(file_bytes, filename=''):
    """
    Parse a convention/event center License Agreement or Use Agreement.
    Extracts: venue_name, organization, rental_fee, net_rental_fee, fb_minimum,
    move_in_date, move_out_date, critical_dates (deposits + cancellation deadlines).
    Uses the same vision/text pipeline as parse_contract_document.
    Returns dict with those keys + 'error', 'raw_text'.
    """
    import sys, importlib, os, re as _re, json as _json
    import io as _io

    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    raw_text = ''
    try:
        if ext == 'pdf':
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
                raw_text = '\n'.join((p.extract_text() or '') for p in pdf.pages)
        elif ext in ('docx', 'doc'):
            try:
                import docx
                doc = docx.Document(_io.BytesIO(file_bytes))
                raw_text = '\n'.join(p.text for p in doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        raw_text += '\n' + '\t'.join(c.text for c in row.cells)
            except ImportError:
                return {'error': 'python-docx not installed'}
    except Exception as e:
        return {'error': f'Could not read file: {e}', 'raw_text': ''}

    _CONTRACT_KEYWORDS = {'license', 'rental', 'agreement', 'licensee', 'licensor',
                          'convention', 'facility', 'venue', 'deposit', 'event',
                          'operator', 'center', 'hall', 'exhibit', 'meeting'}
    _text_lower = raw_text.lower()
    _plain_words = set(_re.findall(r'\b[a-z]{3,}\b', _text_lower))
    _has_keywords = bool(_CONTRACT_KEYWORDS & _plain_words)
    _text_meaningful = bool(raw_text) and len(raw_text.split()) >= 50 and _has_keywords
    is_scanned = (ext == 'pdf' and not _text_meaningful)

    api_key = ''
    try:
        if 'config' in sys.modules:
            importlib.reload(sys.modules['config'])
        else:
            import config as _cfg
            sys.modules['config'] = _cfg
        api_key = sys.modules['config'].ANTHROPIC_API_KEY.strip()
    except Exception:
        pass
    if not api_key:
        api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
    if not api_key:
        return {'error': 'No Anthropic API key configured', 'raw_text': raw_text}

    import anthropic as _ant
    client = _ant.Anthropic(api_key=api_key)

    _cc_prompt = (
        'Extract ALL of the following from this convention/event center License Agreement or Use Agreement. '
        'Return ONLY valid JSON — no markdown, no explanation.\n\n'
        'Required keys:\n'
        '  "venue_name"       — full venue/facility name, or null\n'
        '  "organization"     — the licensee/client organization name, or null\n'
        '  "venue_contact"    — venue sales/event contact name, or null\n'
        '  "venue_email"      — venue contact email, or null\n'
        '  "rental_fee"       — gross/base rental fee as a number (e.g. 667655.00), or null\n'
        '  "net_rental_fee"   — net rental fee after discounts (e.g. 133531.00), or null\n'
        '  "fb_minimum"       — food and beverage minimum commitment as a number, or null\n'
        '  "move_in_date"     — move-in date as "YYYY-MM-DD", or null\n'
        '  "move_out_date"    — move-out date as "YYYY-MM-DD", or null\n'
        '  "event_start_date" — first event/program day as "YYYY-MM-DD", or null\n'
        '  "event_end_date"   — last event/program day as "YYYY-MM-DD", or null\n'
        '  "critical_dates"   — array of important dates:\n'
        '    [{"date":"YYYY-MM-DD","label":"description","amount":number_or_null}, ...]\n'
        '    Include: ALL deposit due dates + amounts, cancellation deadlines with amounts,\n'
        '    insurance certificate due date, F&B payment due date, contract return deadline.\n'
        '    Return [] if none found.\n'
        '\nReturn ONLY the JSON object.'
    )

    ai_data = {}
    try:
        if is_scanned:
            import fitz, base64
            from PIL import Image as _PILImage
            doc = fitz.open(stream=file_bytes, filetype='pdf')
            content = [{'type': 'text', 'text': _cc_prompt}]
            images_added = 0
            for page_num in range(min(len(doc), 10)):
                page = doc[page_num]
                embedded = sorted(page.get_images(full=True), key=lambda x: x[2]*x[3], reverse=True)
                _added_this_page = False
                for img_info in embedded[:1]:
                    xref, w, h = img_info[0], img_info[2], img_info[3]
                    img_data_e = doc.extract_image(xref)
                    raw_img = img_data_e['image']
                    if w*h < 100_000 or len(raw_img) < 30_000:
                        continue
                    if w > 1600:
                        pil = _PILImage.open(_io.BytesIO(raw_img))
                        scale = 1600/w
                        pil = pil.resize((1600, int(h*scale)), _PILImage.LANCZOS)
                        buf = _io.BytesIO()
                        pil.save(buf, format='JPEG', quality=85)
                        raw_img = buf.getvalue()
                    mime = 'image/jpeg' if raw_img[:3] == b'\xff\xd8\xff' else 'image/png'
                    img_b64 = base64.standard_b64encode(raw_img).decode()
                    content.append({'type': 'image', 'source': {'type': 'base64', 'media_type': mime, 'data': img_b64}})
                    images_added += 1
                    _added_this_page = True
                    break
                # Fallback: render page at 150 DPI when no usable embedded image
                if not _added_this_page:
                    mat = fitz.Matrix(150/72, 150/72)
                    pix = page.get_pixmap(matrix=mat)
                    raw_render = pix.tobytes('jpeg')
                    img_b64 = base64.standard_b64encode(raw_render).decode()
                    content.append({'type': 'image', 'source': {'type': 'base64', 'media_type': 'image/jpeg', 'data': img_b64}})
                    images_added += 1
                if images_added >= 8:
                    break
            doc.close()
            resp = client.messages.create(model='claude-opus-4-5', max_tokens=2048,
                                           messages=[{'role': 'user', 'content': content}])
        else:
            text_for_ai = raw_text[:25000]
            resp = client.messages.create(
                model='claude-opus-4-5', max_tokens=2048,
                messages=[{'role': 'user', 'content': _cc_prompt + '\n\nDocument text:\n' + text_for_ai}])

        raw_j = resp.content[0].text.strip()
        raw_j = _re.sub(r'^```[a-z]*\n?', '', raw_j)
        raw_j = _re.sub(r'\n?```$', '', raw_j)
        ai_data = _json.loads(raw_j)
    except Exception as e:
        return {'error': f'AI extraction failed: {e}', 'raw_text': raw_text}

    def _sf(v):
        try: return float(v) if v is not None else None
        except: return None

    def _norm_date(v):
        """Parse a date string to YYYY-MM-DD, or None for null/invalid values."""
        if not v or not isinstance(v, str): return None
        v = v.strip()
        if not v or v.lower() in ('none', 'null', 'n/a', 'na', 'tbd', 'upon execution',
                                  'at signing', 'upon signing', 'at execution'): return None
        try:
            from datetime import datetime as _dt
            for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m/%d/%y', '%B %d, %Y', '%b %d, %Y'):
                try: return _dt.strptime(v[:len(fmt)+2], fmt).strftime('%Y-%m-%d')
                except: pass
        except: pass
        return None

    raw_cd = ai_data.get('critical_dates') or []
    if isinstance(raw_cd, dict):
        raw_cd = [{'date': v, 'label': k.replace('_', ' ').title(), 'amount': None}
                  for k, v in raw_cd.items() if v]
    critical_dates = []
    seen = set()
    _pending_relative = []
    _rel_pat = _re.compile(r'(\d+)\s+days?\s+(?:prior|before)', _re.IGNORECASE)

    for cd in (raw_cd if isinstance(raw_cd, list) else []):
        if not isinstance(cd, dict): continue
        raw_dt = cd.get('date')
        lbl = str(cd.get('label', '')).strip()
        amt = _sf(cd.get('amount'))
        if raw_dt and isinstance(raw_dt, str) and any(
                kw in raw_dt.lower() for kw in ('signing', 'execution', 'upon')):
            key = ('signing', lbl[:20])
            if key not in seen:
                seen.add(key)
                critical_dates.append({'date': 'At signing', 'label': lbl, 'amount': amt})
            continue
        if not raw_dt and _rel_pat.search(lbl):
            _pending_relative.append({'label': lbl, 'amount': amt})
            continue
        iso = _norm_date(str(raw_dt)) if raw_dt is not None else None
        if not iso: continue
        key = (iso, lbl[:20])
        if key not in seen:
            seen.add(key)
            critical_dates.append({'date': iso, 'label': lbl, 'amount': amt})

    _anchor = (_norm_date(ai_data.get('move_in_date')) or
               _norm_date(ai_data.get('event_start_date')))
    if _anchor:
        from datetime import datetime as _dt_cc, timedelta as _td_cc
        _anchor_dt = _dt_cc.strptime(_anchor, '%Y-%m-%d').date()
        for cd in _pending_relative:
            _rm = _rel_pat.search(cd['label'])
            if _rm:
                try:
                    n = int(_rm.group(1))
                    calc_iso = (_anchor_dt - _td_cc(days=n)).strftime('%Y-%m-%d')
                    key = (calc_iso, cd['label'][:20])
                    if key not in seen:
                        seen.add(key)
                        critical_dates.append({'date': calc_iso, 'label': cd['label'], 'amount': cd['amount']})
                except Exception:
                    pass

    critical_dates.sort(key=lambda x: ('0' if x['date'] == 'At signing' else '1') + (x['date'] or ''))

    return {
        'venue_name':       ai_data.get('venue_name') or None,
        'organization':     ai_data.get('organization') or None,
        'venue_contact':    ai_data.get('venue_contact') or None,
        'venue_email':      ai_data.get('venue_email') or None,
        'rental_fee':       _sf(ai_data.get('rental_fee')),
        'net_rental_fee':   _sf(ai_data.get('net_rental_fee')),
        'fb_minimum':       _sf(ai_data.get('fb_minimum')),
        'move_in_date':     _norm_date(ai_data.get('move_in_date')),
        'move_out_date':    _norm_date(ai_data.get('move_out_date')),
        'event_start_date': _norm_date(ai_data.get('event_start_date')),
        'event_end_date':   _norm_date(ai_data.get('event_end_date')),
        'critical_dates':   critical_dates,
        'raw_text':         raw_text,
        'error':            None,
    }


# ── NCSL Pickup Report Excel Importer ────────────────────────────────────────

def parse_pickup_xlsx(file_bytes):
    """
    Parse an NCSL-style pickup report Excel workbook.
    Each sheet = one event/hotel (two side-by-side grids on multi-hotel sheets).
    Returns list of dicts, one per hotel grid:
      {sheet_name, organization, hotel, event_name, contact_name, contact_email,
       booking_id, contracted_block, contracted_rate, attrition_pct, pickups}
    """
    import io as _io
    import re as _re
    from datetime import date as _date, timedelta as _td, datetime as _dt
    import openpyxl

    wb = openpyxl.load_workbook(_io.BytesIO(file_bytes), data_only=True)
    results = []

    for sheet_name in wb.sheetnames:
        if sheet_name.lower() == 'template':
            continue
        ws = wb[sheet_name]

        # Find all grid start columns: look for "PICK-UP UPDATE" in row 1
        grid_cols = []
        for col in range(1, 40):
            v = str(ws.cell(row=1, column=col).value or '').strip().upper()
            if 'PICK-UP' in v or 'PICKUP' in v:
                grid_cols.append(col)
        if not grid_cols:
            grid_cols = [1]

        for gc in grid_cols:
            grid = _parse_one_grid(ws, sheet_name, gc)
            if grid:
                results.append(grid)

    return results


def _parse_one_grid(ws, sheet_name, start_col):
    """Parse a single hotel grid starting at 1-based start_col."""
    from datetime import date as _date, timedelta as _td, datetime as _dt

    def _sv(v):
        if v is None:
            return ''
        return str(v).strip()

    def _cell(row, col_offset=0):
        return ws.cell(row=row, column=start_col + col_offset).value

    def _to_date(val):
        if val is None:
            return None
        if isinstance(val, _dt):
            return val.strftime('%Y-%m-%d')
        if hasattr(val, 'strftime'):
            return val.strftime('%Y-%m-%d')
        if isinstance(val, (int, float)) and 1000 < val < 100000:
            try:
                return (_date(1899, 12, 30) + _td(days=int(val))).strftime('%Y-%m-%d')
            except Exception:
                pass
        s = _sv(val)
        for fmt in ('%m/%d/%Y', '%Y-%m-%d', '%m/%d/%y', '%B %d, %Y'):
            try:
                from datetime import datetime as _dt2
                return _dt2.strptime(s, fmt).strftime('%Y-%m-%d')
            except Exception:
                pass
        return None

    # ── Scan label rows ───────────────────────────────────────────────────────
    org_row = hotel_row = event_row = contact_row = email_row = phone_row = None
    booking_row = dates_row = block_row = cutoff_row = None
    first_pickup_row = None

    for row in range(1, 55):
        lbl = _sv(_cell(row, 0)).lower().rstrip(':').strip()
        if 'organization' in lbl and org_row is None:
            org_row = row
        elif ('hotel' in lbl or 'location' in lbl) and hotel_row is None and 'contact' not in lbl:
            hotel_row = row
        elif ('name' in lbl and 'date' in lbl) and event_row is None:
            event_row = row
        elif 'contact' in lbl and 'number' not in lbl and 'email' not in lbl and contact_row is None:
            contact_row = row
        elif 'number' in lbl and 'booking' not in lbl and phone_row is None:
            phone_row = row
        elif 'email' in lbl and email_row is None:
            email_row = row
        elif 'booking' in lbl and booking_row is None:
            booking_row = row
        elif lbl == 'dates' and dates_row is None:
            dates_row = row
        elif lbl == 'day' and dates_row is None:
            dates_row = row
        elif 'amended block' in lbl:
            block_row = row          # prefer amended over original
        elif lbl == 'block' and block_row is None:
            block_row = row
        elif ('cut-off' in lbl or 'cutoff' in lbl) and cutoff_row is None:
            cutoff_row = row
        elif 'date called' in lbl and first_pickup_row is None:
            first_pickup_row = row

    if dates_row is None:
        return None

    # ── Header fields ─────────────────────────────────────────────────────────
    organization   = _sv(_cell(org_row, 1))    if org_row      else 'NCSL'
    hotel_name     = _sv(_cell(hotel_row, 1))  if hotel_row    else ''
    event_name     = _sv(_cell(event_row, 1))  if event_row    else ''
    contact_name   = _sv(_cell(contact_row, 1)) if contact_row else ''
    contact_phone  = _sv(_cell(phone_row, 1))  if phone_row    else ''
    contact_email  = _sv(_cell(email_row, 1))  if email_row    else ''
    booking_id     = _sv(_cell(booking_row, 1)) if booking_row else ''
    # Customer/group contact — always at col_offset 9 (label at 8) from grid start
    gc_name  = _sv(_cell(contact_row, 9)) if contact_row else ''
    gc_email = _sv(_cell(email_row, 9))   if email_row   else ''

    # ── Event dates (row 9 equivalent, cols C+ = col_offset 2+) ──────────────
    event_dates = {}   # col_offset → 'YYYY-MM-DD'
    for co in range(2, 30):
        d = _to_date(_cell(dates_row, co))
        if d:
            event_dates[co] = d
        elif event_dates:
            # Allow one None gap in case cells are merged
            d2 = _to_date(_cell(dates_row, co + 1))
            if d2 is None:
                break

    if not event_dates and dates_row is not None:
        # Some sheets have day-names on "Dates:" row and actual dates on the next "Day:" row
        for co in range(2, 30):
            d = _to_date(_cell(dates_row + 1, co))
            if d:
                event_dates[co] = d
            elif event_dates:
                d2 = _to_date(_cell(dates_row + 1, co + 1))
                if d2 is None:
                    break
        if event_dates:
            dates_row = dates_row + 1  # update so block row uses same column mapping

    if not event_dates:
        return None

    # ── Contracted block ──────────────────────────────────────────────────────
    contracted_block = {}
    if block_row:
        for co, ds in event_dates.items():
            v = _cell(block_row, co)
            if isinstance(v, (int, float)) and v > 0:
                contracted_block[ds] = int(v)

    # Remove any entry where the value = sum of all others AND > each individual value
    # (it's a "Total" column that shares a date column with a real checkout date)
    if len(contracted_block) > 2:
        block_vals = list(contracted_block.values())
        total_sum  = sum(block_vals)
        last_date  = max(contracted_block.keys())
        last_val   = contracted_block[last_date]
        rest_vals  = [v for d, v in contracted_block.items() if d != last_date]
        rest_sum   = sum(rest_vals)
        if last_val == rest_sum and last_val > max(rest_vals):
            del contracted_block[last_date]

    total_block_rooms = sum(contracted_block.values()) if contracted_block else 0

    # ── Rate & attrition from cutoff row ─────────────────────────────────────
    contracted_rate = None
    attrition_pct   = None
    attrition_rooms = None
    if cutoff_row:
        for co in range(4, 12):
            v = _cell(cutoff_row, co)
            if isinstance(v, (int, float)) and 50 < v < 2000:
                contracted_rate = float(v)
        att_v = _cell(cutoff_row, 5)
        if isinstance(att_v, (int, float)):
            attrition_pct   = float(att_v) if att_v <= 1 else att_v / 100.0
            attrition_rooms = round(total_block_rooms * attrition_pct) if total_block_rooms else None

    # ── Weekly pickup rows ────────────────────────────────────────────────────
    pickups = []
    if first_pickup_row:
        prev_total = None
        row = first_pickup_row
        max_row = first_pickup_row + 120
        while row <= max_row:
            lbl = _sv(_cell(row, 0)).lower()
            if 'date called' in lbl:
                report_date = _to_date(_cell(row, 1))
                if report_date:
                    # Pickup values are on the NEXT row
                    pr = row + 1
                    pickup_by_night = {}
                    for co, ds in event_dates.items():
                        v = _cell(pr, co)
                        if isinstance(v, (int, float)):
                            pickup_by_night[ds] = int(v)
                    # Total column = first col_offset after last date col
                    total_co = max(event_dates.keys()) + 1
                    total_v  = _cell(pr, total_co)
                    if isinstance(total_v, (int, float)):
                        total_rooms = int(total_v)
                    else:
                        total_rooms = sum(pickup_by_night.values())

                    change = (total_rooms - prev_total) if prev_total is not None else None
                    pct_block    = round(total_rooms / total_block_rooms, 4) if total_block_rooms else None
                    pct_attrition = round(total_rooms / attrition_rooms, 4) if attrition_rooms else None

                    pickups.append({
                        'report_date':    report_date,
                        'pickup_by_night': pickup_by_night,
                        'total_rooms':    total_rooms,
                        'change_from_last': change,
                        'pct_of_block':   pct_block,
                        'pct_of_attrition': pct_attrition,
                    })
                    prev_total = total_rooms
            row += 1

    return {
        'sheet_name':       sheet_name,
        'organization':     organization or 'NCSL',
        'hotel':            hotel_name,
        'event_name':       event_name,
        'contact_name':     contact_name,
        'contact_phone':    contact_phone,
        'contact_email':    contact_email,
        'gc_name':          gc_name,
        'gc_email':         gc_email,
        'booking_id':       booking_id,
        'contracted_block': contracted_block,
        'contracted_rate':  contracted_rate,
        'attrition_pct':    attrition_pct,
        'pickups':          pickups,
    }


def parse_hhr_excel(file_bytes):
    """
    Parse a Cvent Client Post Event Housing History Report (.xlsx).
    Returns a dict of summary stats.
    """
    import io as _io
    import openpyxl

    wb = openpyxl.load_workbook(_io.BytesIO(file_bytes), data_only=True)
    ws = wb.active

    def _sv(v):
        return str(v).strip() if v is not None else ''

    def _num(v):
        try:
            return float(v) if v is not None and str(v).strip() != '' else None
        except (ValueError, TypeError):
            return None

    def _to_date(v):
        if v is None:
            return None
        if hasattr(v, 'strftime'):
            return v.strftime('%Y-%m-%d')
        from datetime import date as _date, timedelta as _td
        if isinstance(v, (int, float)) and 1000 < v < 200000:
            try:
                return (_date(1899, 12, 30) + _td(days=int(v))).strftime('%Y-%m-%d')
            except Exception:
                pass
        s = _sv(v)
        for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m/%d/%y', '%B %d, %Y'):
            try:
                from datetime import datetime as _dt
                return _dt.strptime(s, fmt).strftime('%Y-%m-%d')
            except Exception:
                pass
        return None

    stats = {}
    date_cols = {}   # col_index → 'YYYY-MM-DD'
    contracted_block = {}
    final_pickup_by_night = {}

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
        lbl = _sv(row[0].value).lower().rstrip(':').strip() if row[0].value else ''
        col0 = _sv(row[0].value)
        col2 = row[2].value if len(row) > 2 else None

        if 'organization' in lbl and not stats.get('organization'):
            stats['organization'] = _sv(col2)
        elif 'hotel' in lbl and 'location' not in lbl and not stats.get('hotel'):
            stats['hotel'] = _sv(col2)
        elif 'name' in lbl and 'date' in lbl and not stats.get('event_name'):
            stats['event_name'] = _sv(col2)

        # Booking # can be anywhere in the row — scan all cells
        if not stats.get('booking_id'):
            for ci, cell in enumerate(row):
                cv = _sv(cell.value).lower().rstrip(':')
                if 'booking #' in cv or 'booking id' in cv or 'booking number' in cv:
                    for offset in range(1, 4):
                        if ci + offset >= len(row):
                            break
                        bv = _sv(row[ci + offset].value)
                        if bv:
                            try:
                                stats['booking_id'] = str(int(float(bv)))
                            except Exception:
                                stats['booking_id'] = bv
                            break

        # Date header row — find which columns have dates
        elif 'date' == lbl and not date_cols:
            for i, cell in enumerate(row):
                if i < 2:
                    continue
                d = _to_date(cell.value)
                if d:
                    date_cols[i] = d
                elif date_cols:
                    break

        # Contracted block row
        elif 'contracted block' in lbl and not contracted_block:
            for i, d in date_cols.items():
                if i < len(row):
                    v = _num(row[i].value)
                    if v and v > 0:
                        contracted_block[d] = int(v)
            # Total in col 12 (index 12)
            if len(row) > 12:
                stats['contracted_total'] = _num(row[12].value)
            if len(row) > 13:
                stats['contracted_rate'] = _num(row[13].value)

        # Final total pickup
        elif 'final total pickup' in lbl:
            for i, d in date_cols.items():
                if i < len(row):
                    v = _num(row[i].value)
                    if v and v > 0:
                        final_pickup_by_night[d] = int(v)
            if len(row) > 12:
                stats['final_total_pickup'] = _num(row[12].value)

        # Total pickup inside block
        elif 'total pickup inside block' in lbl:
            if len(row) > 12:
                stats['pickup_inside_block'] = _num(row[12].value)

        # Audit pickup
        elif 'total audit pickup' in lbl:
            if len(row) > 12:
                stats['audit_pickup'] = _num(row[12].value)

        # No shows / Cancellations
        elif 'no show' in lbl:
            if len(row) > 12:
                stats['no_shows'] = _num(row[12].value)
        elif 'cancellation' in lbl:
            if len(row) > 12:
                stats['cancellations'] = _num(row[12].value)

        # Revenue lines — labels may appear in any column, scan the whole row
        row_text = ' '.join(_sv(c.value).lower() for c in row)
        if 'total actualized pickup revenue' in row_text and not stats.get('room_revenue'):
            for i, cell in enumerate(row):
                if 'total actualized pickup revenue' in _sv(cell.value).lower():
                    # value is 2 columns to the right
                    for j in (i+2, i+1, i+3):
                        if j < len(row):
                            v = _num(row[j].value)
                            if v and v > 0:
                                stats['room_revenue'] = v
                                break
                    break
        if ('total food' in row_text or ('food' in row_text and 'beverage' in row_text)) and not stats.get('fb_revenue'):
            for i, cell in enumerate(row):
                if 'total food' in _sv(cell.value).lower() or ('food' in _sv(cell.value).lower() and 'beverage' in _sv(cell.value).lower()):
                    for j in (i+2, i+1, i+3):
                        if j < len(row):
                            v = _num(row[j].value)
                            if v and v > 0:
                                stats['fb_revenue'] = v
                                break
                    break

        # Earned comps
        elif 'total comp amount' in lbl or ('earned comp' in lbl and 'total' in lbl):
            if len(row) > 12:
                stats['earned_comps_value'] = _num(row[12].value)
        elif 'total rns eligible for earned comp' in lbl:
            if len(row) > 12:
                stats['earned_comps_rns'] = _num(row[12].value)

        # Hotel accounting sign-off
        elif 'history report approved' in lbl:
            for i, cell in enumerate(row):
                if _sv(cell.value).lower() == 'email' and i+1 < len(row):
                    stats['hotel_approver_email'] = _sv(row[i+1].value)
                if _sv(cell.value).lower() == 'name' and i+1 < len(row):
                    stats['hotel_approver'] = _sv(row[i+1].value)
                d = _to_date(cell.value)
                if d and not stats.get('report_date'):
                    stats['report_date'] = d

    stats['contracted_block']     = contracted_block
    stats['final_pickup_by_night'] = final_pickup_by_night

    # Derived stats
    ct = stats.get('contracted_total') or sum(contracted_block.values()) or None
    fp = stats.get('final_total_pickup')
    if ct and fp:
        stats['pct_of_block'] = round(fp / ct * 100, 1)
    if ct and fp and stats.get('contracted_rate'):
        stats['room_revenue_calc'] = round(fp * stats['contracted_rate'], 2)

    return stats


def strip_hhr_commission_rows(file_bytes):
    """
    Return a client-safe copy of the HHR Excel file with all commission data removed.

    Removes:
      • Entire row 1 except the report title (col C) — clears Comm %, Booking #, Currency
      • Col P (16): commission % column across all rows
      • Col Q (17): commission value column across all rows (except "Total Actualized Revenue")
      • Cols R–S (18–19): annotation text and "Total Commissionable Revenue"/"Avg Rate" headers
      • Commission label rows in col O (15): clears label + value
      • "Commission Check Generated By" row (col A)
      • AUDIT DETAIL section (all individual guest records)

    Renames:
      • "Total Actualized Pickup Revenue (Inside Block)" → "Total Actualized Revenue"
        and sets its Q value = inside block + audit pickup gross revenue
      • "ROOM REVENUE/COMMISSION" → "ROOM REVENUE"
    """
    import io as _io
    from openpyxl import load_workbook

    LABEL_COL = 15   # O — commission label column
    VALUE_COL = 17   # Q — commission value column

    # Commission labels in col O to clear
    LABELS_TO_CLEAR = {
        'total actualized pickup commission (inside block)',
        'total commissionable audit pickup revenue',
        'total commissionable audit pickup commission',
        'total commissionable no show commission',
        'total commissionable cancellation commission',
        'less housing fee commission',
        'less rebate commission',
        'less earned comp commission',
        'total rooms commission due',
    }

    def _safe_clear(ws, rnum, col):
        try:
            ws.cell(row=rnum, column=col).value = None
        except AttributeError:
            pass  # merged cell slave — master was already cleared

    def _unmerge_row(ws, rnum):
        for mr in list(ws.merged_cells.ranges):
            if mr.min_row <= rnum <= mr.max_row:
                ws.unmerge_cells(str(mr))

    wb = load_workbook(_io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    max_col = ws.max_column

    # ── First pass: scan for values we need to preserve + locate key rows ─────
    inside_revenue   = 0.0
    audit_revenue    = 0.0
    revenue_row      = None   # row with "Total Actualized Pickup Revenue" label
    audit_detail_row = None   # first row of the AUDIT DETAIL section

    for row in ws.iter_rows():
        rnum  = row[0].row
        o_lbl = str(ws.cell(row=rnum, column=LABEL_COL).value or '').strip().lower()
        q_val = ws.cell(row=rnum, column=VALUE_COL).value
        d_val = str(ws.cell(row=rnum, column=4).value or '').strip().lower()

        if d_val == 'audit detail' and audit_detail_row is None:
            audit_detail_row = rnum
        if o_lbl == 'total actualized pickup revenue (inside block)':
            revenue_row    = rnum
            inside_revenue = float(q_val or 0)
        elif o_lbl == 'total commissionable audit pickup revenue':
            audit_revenue  = float(q_val or 0)

    gross_revenue = inside_revenue + audit_revenue

    # ── Delete AUDIT DETAIL section first (row numbers stay stable above it) ──
    if audit_detail_row:
        ws.delete_rows(audit_detail_row, ws.max_row - audit_detail_row + 1)

    # ── Clear entire row 1 except col C (the report title), then merge C1:K1 ──
    _unmerge_row(ws, 1)
    for col in range(1, max_col + 1):
        if col != 3:
            _safe_clear(ws, 1, col)
    ws.merge_cells('C1:K1')
    from openpyxl.styles import Alignment
    ws['C1'].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # ── Clear cols P (16), R (18), S (19) for every row ──────────────────────
    # Clear col Q (17) for every row EXCEPT the "Total Actualized Revenue" row
    for rnum in range(1, ws.max_row + 1):
        for col in (16, 18, 19):
            _safe_clear(ws, rnum, col)
        if rnum != revenue_row:
            _safe_clear(ws, rnum, 17)

    # ── Commission label rows in col O ────────────────────────────────────────
    for row in ws.iter_rows():
        rnum   = row[0].row
        o_cell = ws.cell(row=rnum, column=LABEL_COL)
        a_cell = ws.cell(row=rnum, column=1)
        o_lbl  = str(o_cell.value or '').strip().lower()
        a_lbl  = str(a_cell.value or '').strip().lower()

        if o_lbl == 'room revenue/commission':
            o_cell.value = 'ROOM REVENUE'

        elif o_lbl == 'total actualized pickup revenue (inside block)':
            o_cell.value = 'Total Actualized Revenue'
            ws.cell(row=rnum, column=VALUE_COL).value = gross_revenue if gross_revenue else inside_revenue

        elif o_lbl in LABELS_TO_CLEAR:
            o_cell.value = None

        if 'commission check generated' in a_lbl:
            _unmerge_row(ws, rnum)
            for col in range(1, ws.max_column + 1):
                _safe_clear(ws, rnum, col)

    out = _io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out.read()


def clean_hhr_for_client(file_bytes):
    """
    Apply clean, professional formatting to a client-safe HHR Excel file.

    Removes:
      • All indexed/theme cell fill colors (the colored boxes)
      • Commissionable No Shows row
      • Commissionable Cancellations row
      • Non-Commissionable Audit Pickup row
      • Commissionable Audit Pickup row  (Total Audit Pickup is kept)

    Applies:
      • Navy header bar on DATE row (white text)
      • Light navy tint on DAY row
      • Subtle blue tint on Contracted Block row
      • Soft gray on subtotal rows (Total Inside Block, Total Audit Pickup)
      • Navy bar on FINAL TOTAL PICKUP row (white bold text)
      • Light gray section headers (NOTES TO COLLECTIONS, HOTEL ACCOUNTING, etc.)
      • Commission columns (P–S) hidden
      • Thin grid borders on pickup data section
      • Dates reformatted MM/DD; "Please fill out…" stripped from title
      • Calibri 10pt throughout; consistent column widths; freeze panes
    """
    import io as _io, re as _re, datetime as _dt
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    NAVY       = 'FF1A3A5C'
    NAVY_LIGHT = 'FFE8EDF4'
    MID_GRAY   = 'FFF3F4F6'
    WHITE      = 'FFFFFFFF'
    TOTAL_BG   = 'FFD1D5DB'
    ACCENT_BG  = 'FFE9F0F8'
    _THIN      = Side(style='thin', color='FFD1D5DB')
    _THIN_B    = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)

    def _navy_f():          return PatternFill('solid', fgColor=NAVY)
    def _fill(c):           return PatternFill('solid', fgColor=c)
    def _no():              return PatternFill(fill_type=None)
    def _fnt(bold=False, size=10, color='FF000000', italic=False):
        return Font(name='Calibri', bold=bold, size=size, color=color, italic=italic)
    def _ctr(): return Alignment(horizontal='center', vertical='center')
    def _lft(): return Alignment(horizontal='left',   vertical='center')

    ROWS_TO_DELETE = (
        'commissionable no show',
        'commissionable cancellation',
        'non-commissionable audit pickup',
        'commissionable audit pickup',
    )

    wb = load_workbook(_io.BytesIO(file_bytes), data_only=True)
    ws = wb.active

    # ── Delete unwanted rows (bottom-up so indices stay valid) ────────────────
    for rnum in range(ws.max_row, 0, -1):
        al = str(ws.cell(row=rnum, column=1).value or '').strip().lower()
        if any(al.startswith(lbl) for lbl in ROWS_TO_DELETE):
            ws.delete_rows(rnum)

    max_col = ws.max_column

    # ── Strip all existing fills; reset fonts to clean black Calibri ──────────
    for row in ws.iter_rows():
        for cell in row:
            cell.fill = _no()
            old = cell.font
            cell.font = Font(name='Calibri', bold=old.bold,
                             size=old.size or 10, color='FF000000',
                             italic=old.italic)

    # ── Locate key rows ───────────────────────────────────────────────────────
    date_row = day_row = block_row = total_inside_row = None
    total_audit_row = final_row = notes_row = acct_row = note_footer_row = None
    rate_rows = []

    for rnum in range(1, ws.max_row + 1):
        al = str(ws.cell(row=rnum, column=1).value or '').strip().lower()
        cl = str(ws.cell(row=rnum, column=3).value or '').strip().lower()
        if al == 'date'  and date_row  is None:               date_row         = rnum
        elif al == 'day' and day_row   is None:               day_row          = rnum
        elif 'contracted block' in al  and block_row is None: block_row        = rnum
        elif _re.match(r'rate\s+\d+', al):                    rate_rows.append(rnum)
        elif 'total pickup inside block'  in al:              total_inside_row = rnum
        elif 'total audit pickup'         in al:              total_audit_row  = rnum
        elif 'final total pickup'         in al:              final_row        = rnum
        elif 'notes to collections'       in al:              notes_row        = rnum
        elif 'hotel accounting information' in al:            acct_row         = rnum
        elif 'by submitting' in al or 'note:' in al or 'by submitting' in cl:
            note_footer_row = rnum

    # ── Row 1: title ──────────────────────────────────────────────────────────
    for cnum in range(1, max_col + 1):
        cell = ws.cell(row=1, column=cnum)
        if cell.value:
            val = str(cell.value)
            cell.value = val.split('\n')[0].strip() if '\n' in val else val
            cell.font  = _fnt(bold=True, size=13, color=NAVY)
            cell.alignment = _lft()

    # ── Info rows 2–(date_row-1): Org / Hotel / Event ────────────────────────
    for rnum in range(2, (date_row or 5)):
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=rnum, column=cnum)
            is_lbl = cnum == 1 and bool(cell.value)
            cell.font = _fnt(bold=is_lbl, color=NAVY if is_lbl else 'FF000000')
            cell.alignment = _lft()

    # ── DATE row: navy fill, white text, reformat dates MM/DD ────────────────
    if date_row:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=date_row, column=cnum)
            cell.fill = _navy_f()
            if isinstance(cell.value, _dt.datetime):
                cell.value = cell.value.strftime('%m/%d')
            cell.font      = _fnt(bold=bool(cell.value), color=WHITE, size=9)
            cell.alignment = _ctr()

    # ── DAY row: light navy tint ──────────────────────────────────────────────
    if day_row:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=day_row, column=cnum)
            cell.fill      = _fill(NAVY_LIGHT)
            cell.font      = _fnt(bold=True, size=9, color=NAVY)
            cell.alignment = _ctr()

    # ── Contracted Block ──────────────────────────────────────────────────────
    if block_row:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=block_row, column=cnum)
            cell.fill      = _fill(ACCENT_BG)
            cell.font      = _fnt(bold=True, size=9, color=NAVY)
            cell.alignment = _lft() if cnum == 1 else _ctr()

    # ── Rate 1–7 rows: clean white ────────────────────────────────────────────
    for rnum in rate_rows:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=rnum, column=cnum)
            cell.font      = _fnt(size=9)
            cell.alignment = _lft() if cnum == 1 else _ctr()
            if cnum in (14, 15) and isinstance(cell.value, (int, float)):
                cell.number_format = '$#,##0.00'

    # ── Total Inside Block ────────────────────────────────────────────────────
    if total_inside_row:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=total_inside_row, column=cnum)
            cell.fill      = _fill(TOTAL_BG)
            cell.font      = _fnt(bold=True, size=9)
            cell.alignment = _lft() if cnum == 1 else _ctr()

    # ── Total Audit Pickup ────────────────────────────────────────────────────
    if total_audit_row:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=total_audit_row, column=cnum)
            cell.fill      = _fill(TOTAL_BG)
            cell.font      = _fnt(bold=True, size=9)
            cell.alignment = _lft() if cnum == 1 else _ctr()

    # ── FINAL TOTAL PICKUP: navy bar ──────────────────────────────────────────
    if final_row:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=final_row, column=cnum)
            cell.fill      = _navy_f()
            cell.font      = _fnt(bold=True, color=WHITE, size=10)
            cell.alignment = _lft() if cnum == 1 else _ctr()

    # ── Spacer rows between FINAL and NOTES ───────────────────────────────────
    if final_row and notes_row:
        for rnum in range(final_row + 1, notes_row):
            for cnum in range(1, max_col + 1):
                ws.cell(row=rnum, column=cnum).fill = _no()

    # ── NOTES TO COLLECTIONS section ─────────────────────────────────────────
    if notes_row:
        end = acct_row or (notes_row + 12)
        for rnum in range(notes_row, end):
            a_val = str(ws.cell(row=rnum, column=1).value or '').strip()
            for cnum in range(1, max_col + 1):
                cell  = ws.cell(row=rnum, column=cnum)
                c_val = str(cell.value or '').strip()
                is_hdr = ((cnum == 1 and a_val and a_val.upper() == a_val) or
                          (c_val and c_val.upper() == c_val and len(c_val) > 4))
                if is_hdr:
                    cell.fill = _fill(MID_GRAY)
                    cell.font = _fnt(bold=True, size=9, color=NAVY)
                else:
                    cell.font = _fnt(bold=cell.font.bold, size=9)
                cell.alignment = _lft()

    # ── HOTEL ACCOUNTING INFORMATION ─────────────────────────────────────────
    if acct_row:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=acct_row, column=cnum)
            cell.fill      = _fill(MID_GRAY)
            cell.font      = _fnt(bold=True, size=9, color=NAVY)
            cell.alignment = _lft()
        for rnum in range(acct_row + 1, ws.max_row + 1):
            for cnum in range(1, max_col + 1):
                cell = ws.cell(row=rnum, column=cnum)
                cell.font = _fnt(bold=cell.font.bold, size=9)

    # ── Note footer ───────────────────────────────────────────────────────────
    if note_footer_row:
        for cnum in range(1, max_col + 1):
            cell = ws.cell(row=note_footer_row, column=cnum)
            if cell.value:
                cell.font = _fnt(italic=True, size=8, color='FF6B7280')

    # ── Thin borders on pickup grid (DATE row → FINAL row) ───────────────────
    if date_row and final_row:
        right_col = 15
        for rnum in range(date_row, final_row + 1):
            for cnum in range(max_col, 0, -1):
                if ws.cell(row=rnum, column=cnum).value not in (None, ''):
                    right_col = max(right_col, cnum)
                    break
        for rnum in range(date_row, final_row + 1):
            for cnum in range(1, right_col + 1):
                ws.cell(row=rnum, column=cnum).border = _THIN_B

    # ── Column widths ─────────────────────────────────────────────────────────
    ws.column_dimensions['A'].width = 32
    ws.column_dimensions['B'].width = 4
    for c in range(3, 13):
        ws.column_dimensions[get_column_letter(c)].width = 8
    ws.column_dimensions['M'].width = 12
    ws.column_dimensions['N'].width = 10
    ws.column_dimensions['O'].width = 12
    for c in range(16, 20):
        ws.column_dimensions[get_column_letter(c)].width  = 0
        ws.column_dimensions[get_column_letter(c)].hidden = True

    # ── Row heights ───────────────────────────────────────────────────────────
    ws.row_dimensions[1].height = 28
    if date_row:  ws.row_dimensions[date_row].height  = 22
    if day_row:   ws.row_dimensions[day_row].height   = 16
    if final_row: ws.row_dimensions[final_row].height = 20

    # ── Freeze panes below DATE header ────────────────────────────────────────
    if date_row:
        ws.freeze_panes = ws.cell(row=date_row + 1, column=3)

    out = _io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out.read()



def populate_hhr_template(raw_hhr_bytes, template_path=None):
    """
    Fill the client-facing HHR Excel template (hhr_template.xlsx) with data
    extracted from the raw hotel Housing History Report.

    Template layout (fixed):
      Row 1: title  Row 2: ORGANIZATION  Row 3: HOTEL  Row 4: NAME & DATE
      Row 5: DATE   Row 6: DAY
      Row 7: Contracted Block  Rows 8–14: Rate 1–7
      Row 15: Total Pickup Inside Block  Row 16: Total Audit Pickup
      Row 17: FINAL TOTAL PICKUP  Rows 18+: Notes section

    Date columns C–J (cols 3–10, max 8 nights); Total=M(13), Rate=N(14), Revenue=O(15).

    Falls back to clean_hhr_for_client() if the template file is not found.
    """
    import io as _io, os as _os
    import openpyxl

    if template_path is None:
        _here = _os.path.dirname(_os.path.abspath(__file__))
        template_path = _os.path.join(_here, 'static', 'hhr_template.xlsx')
    if not _os.path.exists(template_path):
        return clean_hhr_for_client(strip_hhr_commission_rows(raw_hhr_bytes))

    def _sv(v):  return str(v).strip() if v is not None else ''
    def _num(v):
        try:    return float(v) if v not in (None, '') else None
        except: return None
    def _to_iso(v):
        if v is None: return None
        if hasattr(v, 'strftime'): return v.strftime('%Y-%m-%d')
        s = _sv(v)
        for fmt in ('%m/%d/%Y', '%m/%d/%y', '%Y-%m-%d'):
            try:
                from datetime import datetime as _dt
                return _dt.strptime(s, fmt).strftime('%Y-%m-%d')
            except Exception:
                pass
        return None

    SKIP_COL_A = {
        'commissionable no show', 'commissionable cancellation',
        'non-commissionable audit pickup', 'commissionable audit pickup',
    }

    PICKUP_TO_TMPL_ROW = {
        'contracted block':          7,
        'rate 1':                    8,
        'rate 2':                    9,
        'rate 3':                   10,
        'rate 4':                   11,
        'rate 5':                   12,
        'rate 6':                   13,
        'rate 7':                   14,
        'total pickup inside block': 15,
        'total audit pickup':        16,
        'final total pickup':        17,
    }

    raw_wb  = openpyxl.load_workbook(_io.BytesIO(raw_hhr_bytes), data_only=True)
    raw_ws  = raw_wb.active
    tmpl_wb = openpyxl.load_workbook(template_path)
    tmpl_ws = tmpl_wb.active

    # ── Pass 1: locate date columns and notes-section start ───────────────────
    date_col_map = {}   # raw col (1-indexed) → date YYYY-MM-DD
    day_col_map  = {}   # raw col (1-indexed) → day name string

    for rnum in range(1, raw_ws.max_row + 1):
        lbl = _sv(raw_ws.cell(rnum, 1).value).lower().rstrip(':').strip()
        if lbl == 'date' and not date_col_map:
            for cnum in range(3, raw_ws.max_column + 1):
                d = _to_iso(raw_ws.cell(rnum, cnum).value)
                if d:
                    date_col_map[cnum] = d
                elif date_col_map:
                    break
        elif lbl == 'day' and not day_col_map:
            for cnum in sorted(date_col_map.keys()):
                day_col_map[cnum] = _sv(raw_ws.cell(rnum, cnum).value)
        elif 'notes to collections' in lbl:
            break

    # Template date columns C–J = cols 3–10 (1-indexed); max 8 nights
    TMPL_DATE_COLS = list(range(3, 11))
    sorted_raw_date_cols = sorted(date_col_map.keys())[:8]
    raw_to_tmpl = {rc: TMPL_DATE_COLS[n] for n, rc in enumerate(sorted_raw_date_cols)}

    # ── Identify contracted nights, add 1 pre-shoulder, fix dates to correct year ──
    pre_shoulder_tmpl_col = None
    _dates_written        = False
    _n_tmpl_used          = len(sorted_raw_date_cols)

    _contracted_by_col = {}
    for _rn in range(1, raw_ws.max_row + 1):
        if _sv(raw_ws.cell(_rn, 1).value).lower().strip().startswith('contracted block'):
            for rc in sorted_raw_date_cols:
                _contracted_by_col[rc] = _num(raw_ws.cell(_rn, rc).value) or 0
            break

    _contracted_raw_cols = [rc for rc in sorted_raw_date_cols
                            if _contracted_by_col.get(rc, 0) > 0]
    if _contracted_raw_cols:
        _c_dates   = sorted([date_col_map[rc] for rc in _contracted_raw_cols])
        _c_start   = datetime.strptime(_c_dates[0], '%Y-%m-%d')
        _pre_dt    = _c_start - timedelta(days=1)
        _new_dates = [_pre_dt.strftime('%Y-%m-%d')] + _c_dates

        # Write corrected date row (5) and day row (6) to template
        for i, _d in enumerate(_new_dates):
            _dto = datetime.strptime(_d, '%Y-%m-%d')
            tmpl_ws.cell(5, TMPL_DATE_COLS[i]).value = f"{_dto.month}/{_dto.day}"
            tmpl_ws.cell(6, TMPL_DATE_COLS[i]).value = _dto.strftime('%a')
        for tc in TMPL_DATE_COLS[len(_new_dates):]:
            tmpl_ws.cell(5, tc).value = None
            tmpl_ws.cell(6, tc).value = None

        _dates_written        = True
        pre_shoulder_tmpl_col = TMPL_DATE_COLS[0]
        _n_tmpl_used          = len(_new_dates)

        # Contracted cols shift right by 1 to make room for pre-shoulder in col C
        raw_to_tmpl = {}
        for i, _d in enumerate(_c_dates):
            for rc, iso in date_col_map.items():
                if iso == _d:
                    raw_to_tmpl[rc] = TMPL_DATE_COLS[i + 1]
                    break

    # ── Spend summary lookup (raw HHR notes section) ──────────────────────────
    # Find a labelled cell and return the first positive number within the next
    # few cells to its right. Used for the F&B / Space Rent / Rebate totals.
    def _find_money(*needles):
        needles = [n.lower() for n in needles]
        for rr in range(1, raw_ws.max_row + 1):
            for cc in range(1, raw_ws.max_column + 1):
                txt = _sv(raw_ws.cell(rr, cc).value).lower()
                if txt and all(n in txt for n in needles):
                    for off in (2, 1, 3):
                        if cc + off <= raw_ws.max_column:
                            val = _num(raw_ws.cell(rr, cc + off).value)
                            if val is not None and val > 0:
                                return val
        return None

    # ── Per-row readers (raw row → template columns) ──────────────────────────
    def _row_cells(rnum):
        """{tmpl_col: per-night value} for a raw pickup row (pre-shoulder = 0)."""
        cells = {}
        if pre_shoulder_tmpl_col is not None:
            cells[pre_shoulder_tmpl_col] = 0
        for rc, tc in raw_to_tmpl.items():
            v = _num(raw_ws.cell(rnum, rc).value)
            cells[tc] = v if v is not None else 0
        return cells

    def _row_total(rnum):
        v = _num(raw_ws.cell(rnum, 13).value)          # M = Total Room Nights
        if v is None:
            return None
        return int(v) if v == int(v) else v

    def _tier_rate(rnum):
        """Best-effort nightly rate: N column, else revenue/nights, else $-label."""
        rate = _num(raw_ws.cell(rnum, 14).value)        # N = Rate
        if not rate:
            o = _num(raw_ws.cell(rnum, 15).value)       # O = Revenue
            m = _num(raw_ws.cell(rnum, 13).value)       # M = Total nights
            if o and m:
                rate = round(o / m, 2)
        if rate is None:
            a = _num(raw_ws.cell(rnum, 1).value)        # some HHRs label the row with the $ rate
            if a is not None:
                rate = a
        return rate

    # ── Locate the pickup-grid anchor rows in the raw HHR ─────────────────────
    r_contracted = r_total_inside = r_audit = r_final = None
    for rnum in range(1, raw_ws.max_row + 1):
        lbl = _sv(raw_ws.cell(rnum, 1).value).lower().rstrip(':').strip()
        if 'notes to collections' in lbl:
            break
        if r_contracted is None and lbl.startswith('contracted block'):
            r_contracted = rnum
        elif lbl.startswith('total pickup inside block'):
            r_total_inside = rnum
        elif lbl.startswith('total audit pickup'):
            r_audit = rnum
        elif lbl.startswith('final total pickup'):
            r_final = rnum

    # Can't locate the grid → fall back to the legacy cleaner.
    if not (r_contracted and r_total_inside):
        return clean_hhr_for_client(strip_hhr_commission_rows(raw_hhr_bytes))

    # Rate-tier rows = rows between Contracted Block and Total Pickup Inside Block
    # with any pickup (handles both 'Rate N' and '$rate' label styles).
    rate_rows = []
    for rnum in range(r_contracted + 1, r_total_inside):
        tot    = _num(raw_ws.cell(rnum, 13).value) or 0
        nights = any((_num(raw_ws.cell(rnum, rc).value) or 0) > 0 for rc in raw_to_tmpl)
        if tot > 0 or nights:
            rate_rows.append(rnum)

    # ── Header (org / hotel / event) ──────────────────────────────────────────
    def _hdr(pred):
        for rnum in range(1, r_contracted):
            lbl = _sv(raw_ws.cell(rnum, 1).value).lower().rstrip(':').strip()
            if pred(lbl):
                return _sv(raw_ws.cell(rnum, 3).value)
        return ''
    _org = _hdr(lambda l: 'organization' in l)
    if _org:   tmpl_ws.cell(2, 3).value = _org
    _hotel = _hdr(lambda l: 'hotel' in l and 'accounting' not in l and 'location' not in l)
    if _hotel: tmpl_ws.cell(3, 3).value = _hotel
    _evt = _hdr(lambda l: 'name' in l and 'date' in l)
    if _evt:   tmpl_ws.cell(4, 3).value = _evt

    # ── Date / day fallback (no contracted dates triggered the Pass-1 write) ──
    if not _dates_written:
        for rnum in range(1, r_contracted):
            lbl = _sv(raw_ws.cell(rnum, 1).value).lower().rstrip(':').strip()
            if lbl == 'date' and date_col_map:
                for rc, tc in raw_to_tmpl.items():
                    raw_v = raw_ws.cell(rnum, rc).value
                    disp = f"{raw_v.month}/{raw_v.day}" if hasattr(raw_v, 'month') else _sv(raw_v).split(' ')[0]
                    tmpl_ws.cell(5, tc).value = disp
            elif lbl == 'day' and date_col_map:
                for rc, tc in raw_to_tmpl.items():
                    tmpl_ws.cell(6, tc).value = _sv(raw_ws.cell(rnum, rc).value)

    # ── Spend totals ──────────────────────────────────────────────────────────
    space_rent = _find_money('total meeting room revenue') or _find_money('meeting room revenue')
    fb_spend   = _find_money('total food', 'beverage') or _find_money('food', 'beverage', 'revenue')
    rebate_amt = _find_money('total rebate amount') or _find_money('rebate', 'amount')

    # ── Dynamic re-layout: rate rows vary per event, so rebuild rows 7 down ────
    from copy import copy as _copy
    from openpyxl.styles import Alignment as _Al, Border as _Bd, Side as _Sd

    def _cap(coord):
        c = tmpl_ws[coord]
        return dict(font=_copy(c.font), fill=_copy(c.fill),
                    border=_copy(c.border), align=_copy(c.alignment), numfmt=c.number_format)

    ST = {
        'ct_lbl': _cap('A7'),  'ct_dat': _cap('C7'),  'ct_tot': _cap('M7'),
        'rt_lbl': _cap('A8'),  'rt_dat': _cap('C8'),  'rt_tot': _cap('M8'),
        'ti_lbl': _cap('A15'), 'ti_dat': _cap('C15'), 'ti_tot': _cap('M15'),
        'au_lbl': _cap('A16'), 'au_dat': _cap('C16'), 'au_tot': _cap('M16'),
        'fn_lbl': _cap('A17'), 'fn_dat': _cap('C17'), 'fn_tot': _cap('M17'),
        'sp_lbl': _cap('A19'), 'sp_val': _cap('C19'), 'hdr5': _cap('M5'),
    }

    def _apply(cell, st):
        cell.font = _copy(st['font']); cell.fill = _copy(st['fill'])
        cell.border = _copy(st['border']); cell.alignment = _copy(st['align'])
        cell.number_format = st['numfmt']

    # Rate (N=14) header at row 5; keep it aligned with the Total Room Nights col
    n5 = tmpl_ws.cell(5, 14, 'Rate'); _apply(n5, ST['hdr5'])
    tmpl_ws.column_dimensions['N'].width = 10

    # Clear everything from row 7 down (unmerge first, then drop cells/dims)
    for rng in list(tmpl_ws.merged_cells.ranges):
        if rng.min_row >= 7:
            tmpl_ws.unmerge_cells(str(rng))
    for key in [k for k in list(tmpl_ws._cells.keys()) if k[0] >= 7]:
        del tmpl_ws._cells[key]
    for rd in [rd for rd in list(tmpl_ws.row_dimensions.keys()) if rd >= 7]:
        del tmpl_ws.row_dimensions[rd]

    DATA_COLS = TMPL_DATE_COLS  # C–J (3–10)

    def _write_row(r, label, rnum, lbl_st, dat_st, tot_st, show_rate=False):
        cells = _row_cells(rnum)
        # per-night values (style + value first, merge label last)
        for tc in DATA_COLS:
            _apply(tmpl_ws.cell(r, tc, cells.get(tc, 0)), dat_st)
        mt = tmpl_ws.cell(r, 13, _row_total(rnum) if _row_total(rnum) is not None else 0)
        _apply(mt, tot_st); mt.number_format = '0'
        nc = tmpl_ws.cell(r, 14)
        _apply(nc, tot_st); nc.number_format = '$#,##0.00'
        nc.alignment = _Al(horizontal='center', vertical='center')
        if show_rate:
            nc.value = _tier_rate(rnum)
        _apply(tmpl_ws.cell(r, 1, label), lbl_st)
        _apply(tmpl_ws.cell(r, 2), lbl_st)
        tmpl_ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)

    r = 7
    _write_row(r, 'Contracted Block ', r_contracted, ST['ct_lbl'], ST['ct_dat'], ST['ct_tot'], show_rate=True); r += 1
    for i, rr in enumerate(rate_rows, start=1):
        _write_row(r, f'Rate {i}', rr, ST['rt_lbl'], ST['rt_dat'], ST['rt_tot'], show_rate=True); r += 1
    _write_row(r, 'Total Pickup Inside Block ', r_total_inside, ST['ti_lbl'], ST['ti_dat'], ST['ti_tot']); r += 1
    if r_audit:
        _write_row(r, 'Total Audit Pickup', r_audit, ST['au_lbl'], ST['au_dat'], ST['au_tot']); r += 1
    if r_final:
        _write_row(r, 'FINAL TOTAL PICKUP (ALL)', r_final, ST['fn_lbl'], ST['fn_dat'], ST['fn_tot']); r += 1

    r += 1  # blank spacer row

    # ── Spend summary: AV manual ($0.00 placeholder), the rest auto-fill ───────
    spend_specs = [('AV Spend', None), ('Space Rent Spend', space_rent),
                   ('F&B Spend', fb_spend), ('Rebate Received', rebate_amt)]
    for label, val in spend_specs:
        _apply(tmpl_ws.cell(r, 1, label), ST['sp_lbl'])
        _apply(tmpl_ws.cell(r, 2), ST['sp_lbl'])
        vcell = tmpl_ws.cell(r, 3, round(val, 2) if (val and val > 0) else 0)
        _apply(vcell, ST['sp_val']); vcell.number_format = '$#,##0.00'
        _apply(tmpl_ws.cell(r, 4), ST['sp_val'])
        _apply(tmpl_ws.cell(r, 5), ST['sp_val'])
        tmpl_ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        tmpl_ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=5)
        tmpl_ws.row_dimensions[r].height = 16
        r += 1

    out = _io.BytesIO()
    tmpl_wb.save(out)
    out.seek(0)
    return out.read()

# ── Contract Document Parser ──────────────────────────────────────────────────

def _pdf_pages_to_images(file_bytes, max_pages=10, dpi=100):
    """Render PDF pages to base64 JPEG strings using pymupdf (no poppler needed).
    Returns a list of base64-encoded JPEG byte strings, up to max_pages pages."""
    try:
        import fitz  # pymupdf
        import base64
        doc = fitz.open(stream=file_bytes, filetype='pdf')
        images = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            jpeg_bytes = pix.tobytes('jpeg')
            images.append(base64.standard_b64encode(jpeg_bytes).decode('ascii'))
        return images
    except Exception:
        return []


def _ai_parse_contract_vision(images, api_key):
    """Send rendered PDF page images to Claude vision API and extract structured
    contract data. Used as fallback for scanned (image-only) PDFs."""
    try:
        import anthropic
    except ImportError:
        return {'error': 'anthropic package not installed'}

    contract_extraction_prompt = """You are extracting structured data from scanned hotel group sales contract pages.

Return ONLY a valid JSON object — no markdown fences, no explanation, nothing else.

The JSON object must have exactly these keys:
  contracted_rate       — nightly room rate as a number (e.g. 219.00), or null if not found
  rebate_per_room       — per-room per-night rebate/credit/allowance paid back to the group
                          (e.g. if contract says "$10 rebate per room per night", return 10.00),
                          or null if no rebate clause exists
  cutoff_date           — room block cutoff/release date in YYYY-MM-DD format, or null
  attrition_pct         — attrition percentage as a decimal 0–1 (e.g. 0.80 for 80%), or null
  hotel                 — hotel name as it appears in the contract, or null
  hotel_contact         — hotel contact person's full name, or null
  hotel_contact_email   — hotel contact email address, or null
  group_contact         — client/group contact person's full name, or null
  group_contact_email   — client/group contact email address, or null
  block_review_date     — room block review date in YYYY-MM-DD format, or null. Look for a
                          "block review" or "pickup review" clause specifying when the hotel
                          will review pickup and may reduce the block.
  contracted_block      — object mapping each night's date (YYYY-MM-DD) to the number of rooms
                          blocked for that night (integer). Only include nights with a room
                          count > 0. Example: {"2026-10-31": 200, "2026-11-01": 350}

Rules:
- For contracted_block, look for a night-by-night room schedule or block grid. Dates are
  typically check-in dates (the night starting on that date).
- If the contract shows a flat block (same rooms every night) without per-night breakdown,
  generate one entry per night of the block period, each with that room count.
- Cutoff date may be labeled "cutoff", "release date", "cut-off date", "room block deadline", etc.
- block_review_date: look for a "room block review" clause, "review date", "pickup review date",
  or similar language specifying when the block will be reviewed/reduced.
- Attrition is often "80% attrition" or "you must use 80% of your block" — convert to 0.80.
- Hotel contact is usually signed by a Sales Manager or Director of Sales at the hotel.
- Group contact is the client or meeting planner who signed or is listed as the customer.
- rebate_per_room: look for language like "rebate", "commission rebate", "credit per room",
  "net rate", "$X per room per night rebate/allowance/credit". Return the dollar amount only.
- If a field genuinely cannot be found, use null (not empty string, not 0)."""

    content = [{'type': 'text', 'text': contract_extraction_prompt}]
    for i, b64 in enumerate(images):
        content.append({
            'type': 'image',
            'source': {'type': 'base64', 'media_type': 'image/jpeg', 'data': b64}
        })

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model='claude-opus-4-5',
            max_tokens=2048,
            messages=[{'role': 'user', 'content': content}]
        )
        raw = message.content[0].text.strip()
        raw = re.sub(r'^```[a-z]*\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        data = json.loads(raw)

        # Normalise contracted_block
        block_raw = data.get('contracted_block') or {}
        block = {}
        for k, v in block_raw.items():
            try:
                datetime.strptime(str(k).strip(), '%Y-%m-%d')
                rooms = int(v)
                if rooms > 0:
                    block[str(k).strip()] = rooms
            except Exception:
                continue
        data['contracted_block'] = block

        # Normalise attrition
        atr = data.get('attrition_pct')
        if atr is not None:
            try:
                atr = float(atr)
                if atr > 1:
                    atr = atr / 100
                data['attrition_pct'] = round(atr, 4)
            except Exception:
                data['attrition_pct'] = None

        # Normalise rate
        rate = data.get('contracted_rate')
        if rate is not None:
            try:
                data['contracted_rate'] = round(float(rate), 2)
            except Exception:
                data['contracted_rate'] = None

        # Normalise rebate
        rebate = data.get('rebate_per_room')
        if rebate is not None:
            try:
                data['rebate_per_room'] = round(float(rebate), 2)
            except Exception:
                data['rebate_per_room'] = None

        # Normalise block_review_date — validate YYYY-MM-DD
        brd = data.get('block_review_date')
        if brd:
            try:
                datetime.strptime(str(brd).strip(), '%Y-%m-%d')
                data['block_review_date'] = str(brd).strip()
            except Exception:
                data['block_review_date'] = None
        else:
            data['block_review_date'] = None

        data.setdefault('error', None)
        return data
    except Exception as e:
        return {'error': str(e), 'contracted_block': {}}


def _extract_text_from_contract(file_bytes, filename=''):
    """Extract plain text from a PDF or Word contract file."""
    fname = (filename or '').lower()

    if fname.endswith('.pdf'):
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            return '\n'.join(text_parts)
        except ImportError:
            return None

    if fname.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            parts.append(t)
            return '\n'.join(parts)
        except ImportError:
            return None

    if fname.endswith('.doc'):
        import subprocess, tempfile, os
        try:
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            result = subprocess.run(
                ['textutil', '-convert', 'txt', '-stdout', tmp_path],
                capture_output=True, text=True, timeout=30
            )
            os.unlink(tmp_path)
            return result.stdout if result.returncode == 0 else None
        except Exception:
            return None

    return None


def _ai_parse_contract(text):
    """
    Send contract text to Claude and extract structured hotel contract data.
    Returns a dict with keys: contracted_block, contracted_rate, cutoff_date,
    attrition_pct, hotel, hotel_contact, hotel_contact_email,
    group_contact, group_contact_email, error.
    """
    api_key = ''
    try:
        import importlib, sys
        if 'config' in sys.modules:
            importlib.reload(sys.modules['config'])
        else:
            import config as _cfg
            sys.modules['config'] = _cfg
        api_key = sys.modules['config'].ANTHROPIC_API_KEY.strip()
    except Exception:
        pass
    if not api_key:
        import os
        api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
    if not api_key:
        return {'error': 'Anthropic API key not configured'}

    try:
        import anthropic
    except ImportError:
        return {'error': 'anthropic package not installed'}

    prompt = """You are extracting structured data from a hotel group sales contract.

Return ONLY a valid JSON object — no markdown fences, no explanation, nothing else.

The JSON object must have exactly these keys:
  contracted_rate       — nightly room rate as a number (e.g. 219.00), or null if not found
  rebate_per_room       — per-room per-night rebate/credit/allowance paid back to the group
                          (e.g. if contract says "$10 rebate per room per night", return 10.00),
                          or null if no rebate clause exists
  cutoff_date           — room block cutoff/release date in YYYY-MM-DD format, or null
  attrition_pct         — attrition percentage as a decimal 0–1 (e.g. 0.80 for 80%), or null
  hotel                 — hotel name as it appears in the contract, or null
  hotel_contact         — hotel contact person's full name, or null
  hotel_contact_email   — hotel contact email address, or null
  group_contact         — client/group contact person's full name, or null
  group_contact_email   — client/group contact email address, or null
  block_review_date     — room block review date in YYYY-MM-DD format, or null. Look for a
                          "block review" or "pickup review" clause specifying a date when the
                          hotel will review pickup and may reduce the block.
  contracted_block      — object mapping each night's date (YYYY-MM-DD) to the number of rooms
                          blocked for that night (integer). Only include nights with a room
                          count > 0. Example: {"2026-10-31": 200, "2026-11-01": 350}
  critical_dates        — array of important dates the client must know about:
                          [{"date":"YYYY-MM-DD","label":"description","amount":number_or_null}, ...]
                          Include: deposit due dates + amounts, cutoff date, block review dates,
                          cancellation deadlines, rooming list due dates, attrition review dates,
                          any other contractual deadlines. Use null for amount if no $ amount.
                          Return [] if none found.

Rules:
- For contracted_block, look for a night-by-night room schedule or block grid. Dates are
  typically check-in dates (the night starting on that date).
- If the contract shows a flat block (same rooms every night) without per-night breakdown,
  generate one entry per night of the block period, each with that room count.
- Cutoff date may be labeled "cutoff", "release date", "cut-off date", "room block deadline", etc.
- block_review_date: look for a "room block review" clause, "review date", "pickup review date",
  or similar language specifying when the block will be reviewed/reduced.
- Attrition is often "80% attrition" or "you must use 80% of your block" — convert to 0.80.
- Hotel contact is usually signed by a Sales Manager or Director of Sales at the hotel.
- Group contact is the client or meeting planner who signed or is listed as the customer.
- rebate_per_room: look for language like "rebate", "commission rebate", "credit per room",
  "net rate", "$X per room per night rebate/allowance/credit". Return the dollar amount only.
- If a field genuinely cannot be found, use null (not empty string, not 0).

Contract text:
""" + text[:15000]

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model='claude-haiku-4-5',
            max_tokens=2048,
            messages=[{'role': 'user', 'content': prompt}]
        )
        raw = message.content[0].text.strip()
        raw = re.sub(r'^```[a-z]*\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        data = json.loads(raw)

        # Normalise contracted_block: ensure all values are ints, keys are YYYY-MM-DD
        block_raw = data.get('contracted_block') or {}
        block = {}
        for k, v in block_raw.items():
            try:
                # Validate date format
                datetime.strptime(str(k).strip(), '%Y-%m-%d')
                rooms = int(v)
                if rooms > 0:
                    block[str(k).strip()] = rooms
            except Exception:
                continue
        data['contracted_block'] = block

        # Normalise attrition
        atr = data.get('attrition_pct')
        if atr is not None:
            try:
                atr = float(atr)
                # If someone passed 80 instead of 0.80, fix it
                if atr > 1:
                    atr = atr / 100
                data['attrition_pct'] = round(atr, 4)
            except Exception:
                data['attrition_pct'] = None

        # Normalise rate
        rate = data.get('contracted_rate')
        if rate is not None:
            try:
                data['contracted_rate'] = round(float(rate), 2)
            except Exception:
                data['contracted_rate'] = None

        # Normalise rebate
        rebate = data.get('rebate_per_room')
        if rebate is not None:
            try:
                data['rebate_per_room'] = round(float(rebate), 2)
            except Exception:
                data['rebate_per_room'] = None

        # Normalise block_review_date — validate YYYY-MM-DD
        brd = data.get('block_review_date')
        if brd:
            try:
                datetime.strptime(str(brd).strip(), '%Y-%m-%d')
                data['block_review_date'] = str(brd).strip()
            except Exception:
                data['block_review_date'] = None
        else:
            data['block_review_date'] = None

        # Normalise critical_dates — must be list of {date, label, amount}
        _raw_cd = data.get('critical_dates') or []
        if not isinstance(_raw_cd, list):
            _raw_cd = []
        critical_dates = []
        for _cd in _raw_cd:
            if isinstance(_cd, dict) and _cd.get('date') and _cd.get('label'):
                critical_dates.append({
                    'date':   str(_cd['date']),
                    'label':  str(_cd['label']),
                    'amount': _cd.get('amount'),
                })
        data['critical_dates'] = critical_dates

        data.setdefault('error', None)
        return data

    except Exception as e:
        return {'error': str(e), 'contracted_block': {}}


def parse_contract_document(file_bytes, filename='', hotel_hint=''):
    """
    Extract the contracted room block and key terms from a hotel group contract
    (PDF or Word .docx).

    Uses pdfplumber for PDFs, python-docx for Word files, then sends the raw
    text to Claude to extract structured data.

    Returns a dict:
      {
        'contracted_block': {'YYYY-MM-DD': n_rooms, ...},
        'contracted_rate':  float or None,
        'cutoff_date':      'YYYY-MM-DD' or None,
        'attrition_pct':    float (0.80 = 80%) or None,
        'hotel':            str or None,
        'organization':     str or None,
        'event_name':       str or None,
        'hotel_contact':    str or None,
        'hotel_contact_email': str or None,
        'group_contact':    str or None,
        'group_contact_email': str or None,
        'critical_dates':   [...],
        'years':            [...],
        'error':            str or None,   # present only on failure
        'raw_text':         str,           # full extracted text (for debugging)
      }
    """
    import sys, importlib, os

    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    # ── Extract raw text ──────────────────────────────────────────────────────
    raw_text = ''
    try:
        if ext == 'pdf':
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                raw_text = '\n'.join(
                    (page.extract_text() or '') for page in pdf.pages
                )
        elif ext in ('docx', 'doc'):
            try:
                import docx, io
                doc = docx.Document(io.BytesIO(file_bytes))
                raw_text = '\n'.join(p.text for p in doc.paragraphs)
                # Also grab table cells
                for table in doc.tables:
                    for row in table.rows:
                        raw_text += '\n' + '\t'.join(c.text for c in row.cells)
            except ImportError:
                return {'contracted_block': {}, 'raw_text': '',
                        'error': 'python-docx not installed — run: pip install python-docx'}
        else:
            return {'contracted_block': {}, 'raw_text': '',
                    'error': f'Unsupported file type .{ext} — use PDF or DOCX'}
    except Exception as e:
        return {'contracted_block': {}, 'raw_text': '', 'error': f'Could not read file: {e}'}

    # Treat as scanned if: no text, too short, or lacks contract keywords
    import re as _re
    _CONTRACT_KEYWORDS = {'rate', 'room', 'block', 'arrival', 'departure', 'cutoff',
                          'night', 'hotel', 'group', 'suite', 'meeting', 'attrition',
                          'reservation', 'check', 'guest', 'contract', 'agreement'}
    _text_lower = (raw_text or '').lower()
    _text_words = _text_lower.split()
    _plain_words = set(_re.findall(r'\b[a-z]{3,}\b', _text_lower))
    _has_keywords = bool(_CONTRACT_KEYWORDS & _plain_words)
    _text_meaningful = raw_text and len(_text_words) >= 50 and _has_keywords
    is_scanned_pdf = (ext == 'pdf' and not _text_meaningful)

    # ── Get API key ───────────────────────────────────────────────────────────
    api_key = ''
    try:
        if 'config' in sys.modules:
            importlib.reload(sys.modules['config'])
        else:
            import config as _cfg
            sys.modules['config'] = _cfg
        api_key = sys.modules['config'].ANTHROPIC_API_KEY.strip()
    except Exception:
        pass
    if not api_key:
        api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
    if not api_key:
        return {'contracted_block': {}, 'raw_text': raw_text,
                'error': 'No Anthropic API key configured — cannot parse contract'}

    # ── Direct block parser: position-based "Date … Cut Off" row search ─────────
    import re as _re

    def _md_to_iso(raw_d):
        """Convert M/D/YY or M/D/YYYY to YYYY-MM-DD, or None."""
        parts = raw_d.split('/')
        if len(parts) != 3:
            return None
        try:
            mo, dy, yr = int(parts[0]), int(parts[1]), int(parts[2])
            if yr < 100:
                yr += 2000
            return f'{yr}-{mo:02d}-{dy:02d}'
        except Exception:
            return None

    _direct_years = []
    _date_row_re  = _re.compile(
        r'Date[ \t]+((?:\d{1,2}/\d{1,2}/\d{2,4}[ \t]+)+)Cut[ \t]+Off[ \t]+(\d{1,2}/\d{1,2}/\d{2,4})',
        _re.IGNORECASE)

    for _drm in _date_row_re.finditer(raw_text):
        _raw_dates  = _drm.group(1).strip().split()
        _cutoff_raw = _drm.group(2).strip()
        _cutoff_iso = _md_to_iso(_cutoff_raw)
        _iso_dates  = [d for d in (_md_to_iso(x) for x in _raw_dates) if d]
        if not _iso_dates:
            continue

        # Look ahead for the Total row (within 600 chars)
        _ahead = raw_text[_drm.end(): _drm.end() + 600]
        _total_m = _re.search(r'(?:^|\n)Total[ \t]+([\d .]+)', _ahead)
        if not _total_m:
            continue
        _total_nums = [int(float(x)) for x in _total_m.group(1).split() if _re.match(r'^\d+\.?\d*$', x)]
        _night_counts = _total_nums[:len(_iso_dates)]
        if len(_night_counts) != len(_iso_dates):
            continue
        _block = {d: c for d, c in zip(_iso_dates, _night_counts)}

        # Look back 800 chars for year, rate, attrition
        _back = raw_text[max(0, _drm.start() - 800): _drm.start()]

        _yr_m   = _re.search(r'(?:Event\s+Year|Year)\s+(\d{4})', _back, _re.IGNORECASE)
        _yr_num = int(_yr_m.group(1)) if _yr_m else None
        # Fallback: infer year from first ISO date
        if not _yr_num and _iso_dates:
            _yr_num = int(_iso_dates[0][:4])

        _rate_m = _re.search(r'Group\s+rate[:\s]+\$(\d+(?:\.\d+)?)\s+per\s+room', _back, _re.IGNORECASE)
        _rate   = float(_rate_m.group(1)) if _rate_m else None

        _atr_m  = _re.search(
            r'(\d{2,3})\s*%[^\n]*(?:attrition|commitment)|(?:attrition|commitment)[^\n]*?(\d{2,3})\s*%',
            _back, _re.IGNORECASE)
        _atr = None
        if _atr_m:
            _atr_raw = _atr_m.group(1) or _atr_m.group(2)
            if _atr_raw:
                _atr_val = int(_atr_raw) / 100.0
                # Flip release % to minimum % (20% release → 80% minimum)
                _atr = (1.0 - _atr_val) if _atr_val < 0.5 else _atr_val

        _direct_years.append({
            'year':              _yr_num,
            'event_name':        None,
            'contracted_block':  _block,
            'contracted_rate':   _rate,
            'rebate_per_room':   None,
            'cutoff_date':       _cutoff_iso,
            'block_review_date': None,
            'attrition_pct':     _atr,
        })

    # ── Ask Claude for contacts/org/attrition only ────────────────────────────
    import json as _json, re as _re2
    ai_data = {}
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)

        hotel_hint_line = (
            f'\nIMPORTANT: This contract covers multiple hotels. '
            f'Extract data ONLY for the hotel named: "{hotel_hint}".\n'
        ) if hotel_hint else ''

        # For long contracts send only pages that contain the header/signature block
        _text_for_ai = raw_text[:15000]

        _critical_dates_instructions = (
            '  "critical_dates"     — array of important dates the client must know about:\n'
            '                         [{"date":"YYYY-MM-DD","label":"description","amount":number_or_null}, ...]\n'
            '                         Include: deposit due dates + amounts, cutoff date, block review dates,\n'
            '                         cancellation deadlines, rooming list due dates, attrition review dates,\n'
            '                         any other contractual deadlines. Use null for amount if no $ amount.\n'
            '                         Return [] if none found.\n'
        )
        if is_scanned_pdf:
            import fitz, base64
            doc  = fitz.open(stream=file_bytes, filetype='pdf')
            content = [{'type': 'text', 'text': (
                'Extract ONLY the fields listed below from this hotel contract. '
                'Return a JSON object with ONLY these keys — do NOT invent or guess room block data.\n'
                + hotel_hint_line
            )}]
            images_added = 0
            for page_num in range(min(len(doc), 10)):
                page = doc[page_num]
                # Prefer large embedded images (actual scanned pages) over re-rendering.
                # Skip tiny images < 50KB (DocuSign badges, logos, etc.)
                embedded = sorted(page.get_images(full=True),
                                  key=lambda x: x[2] * x[3], reverse=True)  # sort by w×h
                added_from_page = False
                for img_info in embedded[:1]:   # only the largest image per page
                    xref = img_info[0]
                    w, h = img_info[2], img_info[3]
                    if w * h < 100_000:
                        continue
                    img_data = doc.extract_image(xref)
                    raw = img_data['image']
                    if len(raw) < 30_000:
                        continue
                    # Downscale large images to ≤1600px wide for the API
                    if w > 1600:
                        import io as _io
                        try:
                            from PIL import Image as _PILImage
                            pil = _PILImage.open(_io.BytesIO(raw))
                            scale = 1600 / w
                            pil = pil.resize((1600, int(h * scale)), _PILImage.LANCZOS)
                            buf = _io.BytesIO()
                            pil.save(buf, format='JPEG', quality=85)
                            raw = buf.getvalue()
                        except ImportError:
                            pass
                    ext_img = img_data.get('ext', 'jpeg')
                    mime = 'image/jpeg' if ext_img in ('jpg', 'jpeg') else 'image/png'
                    img_b64 = base64.standard_b64encode(raw).decode()
                    content.append({'type': 'image',
                                     'source': {'type': 'base64', 'media_type': mime,
                                                'data': img_b64}})
                    images_added += 1
                    added_from_page = True
                    break
                # Fallback: render page at 150 DPI if no usable embedded image
                if not added_from_page:
                    mat = fitz.Matrix(150/72, 150/72)
                    pix = page.get_pixmap(matrix=mat)
                    img_b64 = base64.standard_b64encode(pix.tobytes('jpeg')).decode()
                    content.append({'type': 'image',
                                     'source': {'type': 'base64', 'media_type': 'image/jpeg',
                                                'data': img_b64}})
                    images_added += 1
                if images_added >= 8:
                    break
            doc.close()
            content.append({'type': 'text', 'text': (
                'Extract ALL of the following from this hotel group contract. '
                'Return ONLY valid JSON — no markdown, no explanation.\n'
                + hotel_hint_line +
                '\nRequired keys:\n'
                '  "hotel"               — full hotel name, or null\n'
                '  "organization"        — group/client org name, or null\n'
                '  "hotel_contact"       — hotel contact full name, or null\n'
                '  "hotel_contact_email" — hotel contact email, or null\n'
                '  "group_contact"       — client contact full name, or null\n'
                '  "group_contact_email" — client contact email, or null\n'
                '  "contracted_rate"     — room rate as a number (e.g. 189.00), or null\n'
                '  "contracted_block"    — night-by-night room block as {"YYYY-MM-DD": rooms, ...}, or {}\n'
                '  "cutoff_date"         — cut-off date as "YYYY-MM-DD", or null\n'
                '  "attrition_pct"       — attrition as decimal (80%→0.80), or null\n'
                '  "rebate_per_room"     — $/room/night rebate if any, or null\n'
                '  "block_review_date"   — block review date "YYYY-MM-DD", or null\n'
                + _critical_dates_instructions +
                '\nReturn ONLY the JSON object, no other text.'
            )})
            response = client.messages.create(
                model='claude-opus-4-5', max_tokens=2048,
                messages=[{'role': 'user', 'content': content}])
        else:
            extraction_prompt = (
                'Extract ONLY the fields below from this hotel contract text. '
                'Return ONLY valid JSON — no markdown, no explanation.\n'
                + hotel_hint_line +
                '\nRequired keys:\n'
                '  "hotel"               — full hotel name, or null\n'
                '  "organization"        — group/client org name, or null\n'
                '  "hotel_contact"       — hotel contact full name, or null\n'
                '  "hotel_contact_email" — hotel contact email, or null\n'
                '  "group_contact"       — client contact full name, or null\n'
                '  "group_contact_email" — client contact email, or null\n'
                '  "attrition_pct"       — attrition as decimal (80%→0.80), or null\n'
                '  "rebate_per_room"     — $/room/night rebate if any, or null\n'
                '  "block_review_date"   — block review date YYYY-MM-DD, or null\n'
                + _critical_dates_instructions +
                '\nDO NOT return contracted_block or contracted_rate — those are handled separately.\n'
                '\nContract text:\n' + _text_for_ai
            )
            response = client.messages.create(
                model='claude-opus-4-5', max_tokens=1024,
                messages=[{'role': 'user', 'content': extraction_prompt}])

        raw_json = response.content[0].text.strip()
        raw_json = _re2.sub(r'^```[a-z]*\n?', '', raw_json)
        raw_json = _re2.sub(r'\n?```$',        '', raw_json)
        ai_data  = _json.loads(raw_json)
    except Exception:
        pass   # AI is supplemental — failures are silent, direct parse still used

    # ── Build final result: direct parse is authoritative for block/rate/cutoff ──
    years = []
    base_years = _direct_years if _direct_years else []

    # For scanned PDFs the vision response includes block/rate in ai_data — normalise and use it.
    if not base_years and is_scanned_pdf and ai_data:
        from datetime import datetime as _dt_scan

        # Normalise contracted_block — Claude may return nested room-type dicts
        blk_scan = {}
        for k, v in (ai_data.get('contracted_block') or {}).items():
            try:
                _dt_scan.strptime(k, '%Y-%m-%d')
                if isinstance(v, dict):
                    total = sum(int(n) for n in v.values() if isinstance(n, (int, float)) and n > 0)
                else:
                    total = int(float(v))
                if total > 0:
                    blk_scan[k] = total
            except Exception:
                pass

        # Normalise contracted_rate — Claude may return a dict of room types
        _raw_rate = ai_data.get('contracted_rate')
        if isinstance(_raw_rate, dict):
            _rate_vals = sorted(v for v in [_safe_float(x) for x in _raw_rate.values() if x is not None] if v > 0)
            _contracted_rate = _rate_vals[0] if _rate_vals else None
        else:
            _contracted_rate = _safe_float(_raw_rate) or None

        _atr_raw = _safe_float(ai_data.get('attrition_pct'))
        _atr = (_atr_raw / 100.0) if _atr_raw and _atr_raw > 1 else (_atr_raw or None)

        # Normalise critical_dates from vision — Claude may return dict instead of list
        _vis_cd = ai_data.get('critical_dates') or []
        if isinstance(_vis_cd, dict):
            _vis_cd_list = []
            for _lbl, _val in _vis_cd.items():
                if not _val or not isinstance(_val, str): continue
                _iso = _parse_any_date(_val)
                if _iso:
                    _vis_cd_list.append({'date': _iso, 'label': _lbl.replace('_', ' ').title(), 'amount': None})
            ai_data['critical_dates'] = _vis_cd_list

        if blk_scan or _contracted_rate:
            base_years = [{
                'year':              int(list(blk_scan.keys())[0][:4]) if blk_scan else None,
                'event_name':        None,
                'contracted_block':  blk_scan,
                'contracted_rate':   _contracted_rate,
                'rebate_per_room':   _safe_float(ai_data.get('rebate_per_room')),
                'cutoff_date':       ai_data.get('cutoff_date') or None,
                'block_review_date': ai_data.get('block_review_date') or None,
                'attrition_pct':     _atr,
            }]

    # If no direct years, ask Claude for block too (short single-year contract)
    if not base_years and not is_scanned_pdf and raw_text.strip():
        try:
            full_prompt = (
                'Extract group room block data from this hotel contract. '
                'Return ONLY valid JSON.\n'
                + hotel_hint_line +
                '\nKeys:\n'
                '  "contracted_block": {"YYYY-MM-DD": rooms, ...} or {}\n'
                '  "contracted_rate": number or null\n'
                '  "cutoff_date": "YYYY-MM-DD" or null\n'
                '  "attrition_pct": decimal or null\n'
                '  "year": integer or null\n'
                '  "rebate_per_room": number or null\n'
                '  "block_review_date": "YYYY-MM-DD" or null\n'
                '\nContract text:\n' + raw_text[:20000]
            )
            resp2 = client.messages.create(
                model='claude-opus-4-5', max_tokens=1024,
                messages=[{'role': 'user', 'content': full_prompt}])
            raw2  = resp2.content[0].text.strip()
            raw2  = _re2.sub(r'^```[a-z]*\n?', '', raw2)
            raw2  = _re2.sub(r'\n?```$',        '', raw2)
            blk_data = _json.loads(raw2)
            from datetime import datetime as _dt2
            blk = {}
            for k, v in (blk_data.get('contracted_block') or {}).items():
                try:
                    _dt2.strptime(k, '%Y-%m-%d'); blk[k] = int(v)
                except Exception:
                    pass

            def _sf(val):
                try:
                    return float(val) if val is not None else None
                except (TypeError, ValueError):
                    return None

            base_years = [{
                'year':              blk_data.get('year'),
                'event_name':        None,
                'contracted_block':  blk,
                'contracted_rate':   _sf(blk_data.get('contracted_rate')),
                'rebate_per_room':   _sf(blk_data.get('rebate_per_room')),
                'cutoff_date':       blk_data.get('cutoff_date') or None,
                'block_review_date': blk_data.get('block_review_date') or None,
                'attrition_pct':     _sf(blk_data.get('attrition_pct')),
            }]
        except Exception:
            pass

    def _safe_float_local(val):
        try:
            return float(val) if val is not None else None
        except (TypeError, ValueError):
            return None

    for dy in sorted(base_years, key=lambda y: (y.get('year') or 0)):
        years.append({
            'year':              dy.get('year'),
            'event_name':        None,
            'contracted_block':  dy.get('contracted_block', {}),
            'contracted_rate':   dy.get('contracted_rate'),
            'rebate_per_room':   _safe_float_local(ai_data.get('rebate_per_room')),
            'cutoff_date':       dy.get('cutoff_date'),
            'block_review_date': ai_data.get('block_review_date') or None,
            'attrition_pct':     dy.get('attrition_pct') or _safe_float_local(ai_data.get('attrition_pct')),
        })

    if not years:
        return {'contracted_block': {}, 'raw_text': raw_text,
                'error': 'Could not extract room block — contract format not recognized'}

    # ── Direct regex parser for critical dates (runs over full raw_text) ────────
    # Catches deposit schedules, cancellation deadlines, and other dated clauses
    # that may fall outside Claude's 15k-char window.

    def _parse_any_date(s, min_year=2024):
        """Try to parse a date string in multiple formats → YYYY-MM-DD or None.
        Rejects dates before min_year to filter header/prep-date false positives."""
        import re as _r2
        from datetime import datetime as _DT
        s = s.strip().rstrip(',').strip()
        # Strip ordinal suffixes: "1st" → "1", "2nd" → "2", "3rd" → "3", "28th" → "28"
        s = _r2.sub(r'(\d+)(?:st|nd|rd|th)\b', r'\1', s)
        def _check(dt_obj):
            if dt_obj and dt_obj.year >= min_year:
                return dt_obj.strftime('%Y-%m-%d')
            return None
        # ISO
        try:
            return _check(_DT.strptime(s[:10], '%Y-%m-%d'))
        except Exception: pass
        # M/D/YY or M/D/YYYY
        m = _r2.match(r'^(\d{1,2})/(\d{1,2})/(\d{2,4})$', s)
        if m:
            mo, dy, yr = int(m.group(1)), int(m.group(2)), int(m.group(3))
            if yr < 100: yr += 2000
            try: return _check(_DT(yr, mo, dy))
            except Exception: pass
        # Month DD, YYYY  or  Month DD YYYY (ordinals already stripped above)
        MONTHS = {'january':1,'february':2,'march':3,'april':4,'may':5,'june':6,
                  'july':7,'august':8,'september':9,'october':10,'november':11,'december':12}
        m2 = _r2.match(r'^([A-Za-z]+)\s+(\d{1,2}),?\s+(\d{4})$', s)
        if m2:
            mo_name = m2.group(1).lower()
            if mo_name in MONTHS:
                try: return _check(_DT(int(m2.group(3)), MONTHS[mo_name], int(m2.group(2))))
                except Exception: pass
        # DD Month YYYY
        m3 = _r2.match(r'^(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})$', s)
        if m3:
            mo_name = m3.group(2).lower()
            if mo_name in MONTHS:
                try: return _check(_DT(int(m3.group(3)), MONTHS[mo_name], int(m3.group(1))))
                except Exception: pass
        return None

    import re as _rer
    _direct_critical = []
    _seen_dates = set()  # deduplicate by (date, label[:20])

    # Date token — matches most written and numeric date forms, including ordinals
    # e.g. "August 1st, 2026", "March 28th 2027", "April 18th, 2027"
    _MONTH_NAMES = r'(?:January|February|March|April|May|June|July|August|September|October|November|December)'
    _ORDINAL_SFX = r'(?:st|nd|rd|th)?'
    _DT_PAT = (
        r'(?:'
        + _MONTH_NAMES + r'\s+\d{1,2}' + _ORDINAL_SFX + r',?\s+\d{4}'
        r'|\d{1,2}/\d{1,2}/\d{2,4}'
        r'|\d{4}-\d{2}-\d{2}'
        r'|\d{1,2}\s+' + _MONTH_NAMES + r'\s+\d{4}'
        r')'
    )
    _DOLLAR_PAT = r'\$\s*([\d,]+(?:\.\d{2})?)'

    # ── 1. Deposit amounts with associated dates ──────────────────────────────
    # Looks for patterns like:
    #   "deposit of $X,XXX … due … [date]"
    #   "[date] … deposit … $X,XXX"
    #   "Initial/First/Second/Final deposit: $X,XXX … [date]"
    _deposit_clause_re = _rer.compile(
        r'(?P<prefix>(?:initial|first|second|third|fourth|final|non[-\s]?refundable|additional|advance|partial)?'
        r'\s*deposit[^.]{0,250}?)'
        r'(?P<date>' + _DT_PAT + r')',
        _rer.IGNORECASE | _rer.DOTALL
    )
    for _m in _deposit_clause_re.finditer(raw_text):
        _clause = _m.group('prefix')
        _date_str = _m.group('date').strip()
        _iso = _parse_any_date(_date_str)
        if not _iso:
            continue
        # Find dollar amount nearest to "deposit" in the clause
        _amt = None
        _amts = _rer.findall(_DOLLAR_PAT, _clause)
        if _amts:
            try: _amt = float(_amts[0].replace(',', ''))
            except Exception: pass
        # Build label from context
        _prefix_lower = _clause.lower()
        if 'initial' in _prefix_lower or 'first' in _prefix_lower:
            _label = 'Initial deposit due'
        elif 'second' in _prefix_lower:
            _label = 'Second deposit due'
        elif 'third' in _prefix_lower:
            _label = 'Third deposit due'
        elif 'final' in _prefix_lower:
            _label = 'Final deposit due'
        elif 'non' in _prefix_lower and 'refund' in _prefix_lower:
            _label = 'Non-refundable deposit due'
        else:
            _label = 'Deposit due'
        _key = (_iso, _label[:20])
        if _key not in _seen_dates:
            _seen_dates.add(_key)
            _direct_critical.append({'date': _iso, 'label': _label, 'amount': _amt})

    # Also catch date-first deposit patterns:
    #   "[date] … deposit … $X,XXX"
    _deposit_date_first_re = _rer.compile(
        r'(?P<date>' + _DT_PAT + r')'
        r'(?P<suffix>[^.]{0,200}?deposit[^.]{0,100}?)',
        _rer.IGNORECASE | _rer.DOTALL
    )
    for _m in _deposit_date_first_re.finditer(raw_text):
        _date_str = _m.group('date').strip()
        _iso = _parse_any_date(_date_str)
        if not _iso:
            continue
        _suffix = _m.group('suffix')
        _amts = _rer.findall(_DOLLAR_PAT, _suffix)
        _amt = None
        if _amts:
            try: _amt = float(_amts[0].replace(',', ''))
            except Exception: pass
        _label = 'Deposit due'
        _key = (_iso, _label[:20])
        if _key not in _seen_dates:
            _seen_dates.add(_key)
            _direct_critical.append({'date': _iso, 'label': _label, 'amount': _amt})

    # ── 2. Cancellation deadlines ─────────────────────────────────────────────
    # Strategy A: "cancellation between X and $AMOUNT\nY" (Hilton style)
    # For each "cancellation between" occurrence, find the date that appears
    # AFTER the dollar amount — that's the end-of-range deadline.
    for _cm in _rer.finditer(r'cancellation\s+between', raw_text, _rer.IGNORECASE):
        _win = raw_text[_cm.start():_cm.start()+400]
        _all_dates = _rer.findall(_DT_PAT, _win)
        _all_amts  = _rer.findall(_DOLLAR_PAT, _win)
        if not _all_dates: continue
        _deadline_iso = None
        _amt = None
        _amt_m = _rer.search(_DOLLAR_PAT, _win)
        if _amt_m:
            try: _amt = float(_amt_m.group(1).replace(',', ''))
            except Exception: pass
            if _amt and _amt < 100: _amt = None
            _after_amt = _win[_amt_m.end():]
            _end_dates = _rer.findall(_DT_PAT, _after_amt)
            if _end_dates:
                _deadline_iso = _parse_any_date(_end_dates[0].strip())
        if not _deadline_iso:
            for _ds in reversed(_all_dates):
                _deadline_iso = _parse_any_date(_ds.strip())
                if _deadline_iso: break
        if not _deadline_iso: continue
        _key = (_deadline_iso, 'cancel')
        if _key not in _seen_dates:
            _seen_dates.add(_key)
            _direct_critical.append({'date': _deadline_iso, 'label': 'Cancellation deadline', 'amount': _amt})

    # ── 3. Rooming list due date ──────────────────────────────────────────────
    _rl_re = _rer.compile(
        r'rooming\s+list[^.]{0,200}?' + _DT_PAT,
        _rer.IGNORECASE | _rer.DOTALL
    )
    for _m in _rl_re.finditer(raw_text):
        _clause = _m.group(0)
        _dates_found = _rer.findall(_DT_PAT, _clause)
        for _ds in _dates_found:
            _iso = _parse_any_date(_ds.strip())
            if not _iso: continue
            _key = (_iso, 'Rooming list due')
            if _key not in _seen_dates:
                _seen_dates.add(_key)
                _direct_critical.append({'date': _iso, 'label': 'Rooming list due', 'amount': None})

    # ── 4. Block / attrition review dates ────────────────────────────────────
    _review_re = _rer.compile(
        r'(?:block\s+review|attrition\s+review|pickup\s+review)[^.]{0,200}?' + _DT_PAT,
        _rer.IGNORECASE | _rer.DOTALL
    )
    for _m in _review_re.finditer(raw_text):
        _clause = _m.group(0)
        _dates_found = _rer.findall(_DT_PAT, _clause)
        for _ds in _dates_found:
            _iso = _parse_any_date(_ds.strip())
            if not _iso: continue
            _lbl = ('Attrition review date' if 'attrition' in _clause.lower() else 'Block review date')
            _key = (_iso, _lbl[:20])
            if _key not in _seen_dates:
                _seen_dates.add(_key)
                _direct_critical.append({'date': _iso, 'label': _lbl, 'amount': None})

    # ── 5. Cutoff date — inject from the Exhibit A direct parser first ────────
    # _direct_years already has the most accurate cutoff (parsed from the block table).
    # Add it to critical_dates so it appears even if Claude missed it.
    for _dy in _direct_years:
        _co = _dy.get('cutoff_date')
        if not _co: continue
        _key = (_co, 'Cut-off date')
        if _key not in _seen_dates:
            _seen_dates.add(_key)
            _direct_critical.append({'date': _co, 'label': 'Cut-off date', 'amount': None})

    # Also scan the full text for cut-off mentions not in Exhibit A
    _cutoff_re = _rer.compile(
        r'cut[\s\-]?off[^.]{0,200}?' + _DT_PAT,
        _rer.IGNORECASE | _rer.DOTALL
    )
    for _m in _cutoff_re.finditer(raw_text):
        _clause = _m.group(0)
        _dates_found = _rer.findall(_DT_PAT, _clause)
        for _ds in _dates_found:
            _iso = _parse_any_date(_ds.strip())
            if not _iso: continue
            _key = (_iso, 'Cut-off date')
            if _key not in _seen_dates:
                _seen_dates.add(_key)
                _direct_critical.append({'date': _iso, 'label': 'Cut-off date', 'amount': None})

    # ── 6. Attrition review — require "review" or "damages" in same clause ──────
    # This prevents meeting-block dates from being misidentified as review dates.
    _atr_re = _rer.compile(
        r'attrition[^.]{0,300}?(?:review|damages|waiver)[^.]{0,200}?' + _DT_PAT,
        _rer.IGNORECASE | _rer.DOTALL
    )
    for _m in _atr_re.finditer(raw_text):
        _clause = _m.group(0)
        _dates_found = _rer.findall(_DT_PAT, _clause)
        for _ds in _dates_found:
            _iso = _parse_any_date(_ds.strip())
            if not _iso: continue
            _key = (_iso, 'Attrition review d')
            if _key not in _seen_dates:
                _seen_dates.add(_key)
                _direct_critical.append({'date': _iso, 'label': 'Attrition review date', 'amount': None})

    # ── 7. Cancellation milestone schedule (liquidated damages table) ─────────
    # Looks for date ranges followed by percentages and/or dollar amounts,
    # typically in a cancellation schedule like:
    #   "Date of signing to 3/11/2025  45% × rooms × rate = $65,712"
    #   "3/12/2025 to 12/9/2025  65% × ... = $94,918"
    # Match "Date of signing to [date] ... = $X,XXX" or "[date] to [date] ... = $X,XXX"
    # Use [\s\S] instead of [^.] so decimal points in amounts like $129.00 don't stop the match.
    _cancel_sched_re = _rer.compile(
        r'(?:date\s+of\s+signing|' + _DT_PAT + r')\s+to\s+(' + _DT_PAT + r')'
        r'[\s\S]{0,400}?=\s*\$\s*([\d,]+(?:\.\d{2})?)',
        _rer.IGNORECASE
    )
    for _m in _cancel_sched_re.finditer(raw_text):
        _ds = _m.group(1).strip() if _m.group(1) else ''
        _iso = _parse_any_date(_ds)
        if not _iso: continue
        _raw_amt = _m.group(2) if _m.group(2) else ''
        _amt = None
        try: _amt = float(_raw_amt.replace(',', ''))
        except Exception: pass
        if _amt and _amt < 100: continue   # skip rates like $129.00 — penalty totals are much larger
        _key = (_iso, 'cancel')
        if _key not in _seen_dates:
            _seen_dates.add(_key)
            _direct_critical.append({'date': _iso, 'label': 'Cancellation penalty deadline', 'amount': _amt})

    # ── 8. Relative dates — "N days prior to the Event" ─────────────────────────
    _event_start_iso = None
    if _direct_years:
        _blk0 = _direct_years[0].get('contracted_block', {})
        if _blk0:
            _event_start_iso = sorted(_blk0.keys())[0]
    if not _event_start_iso and base_years:
        _blk0 = base_years[0].get('contracted_block', {})
        if isinstance(_blk0, dict) and _blk0:
            _event_start_iso = sorted(_blk0.keys())[0]
    if not _event_start_iso and ai_data.get('contracted_block'):
        _blk0 = ai_data.get('contracted_block', {})
        if isinstance(_blk0, dict) and _blk0:
            _event_start_iso = sorted(_blk0.keys())[0]

    if _event_start_iso:
        from datetime import datetime as _dt_rel, timedelta as _td_rel
        try:
            _event_start_dt = _dt_rel.strptime(_event_start_iso, '%Y-%m-%d').date()
        except Exception:
            _event_start_dt = None

        if _event_start_dt:
            _relative_re = _rer.compile(
                r'(\d+)\s+days?\s+(?:prior\s+to|before)\s+'
                r'(?:the\s+)?(?:event|meeting|arrival|check[- ]?in|first\s+day|start)',
                _rer.IGNORECASE
            )
            for _rm in _relative_re.finditer(raw_text):
                try:
                    _n_days = int(_rm.group(1))
                except Exception:
                    continue
                if _n_days <= 0 or _n_days > 730:
                    continue
                _calc_date = _event_start_dt - _td_rel(days=_n_days)
                _calc_iso  = _calc_date.strftime('%Y-%m-%d')
                _look_back = raw_text[max(0, _rm.start()-400):_rm.start()]
                _label = None
                # Pass 0: all-caps section heading within 800 chars
                _look_back_wide = raw_text[max(0, _rm.start()-800):_rm.start()]
                _caps_m = list(_rer.finditer(r'\n([A-Z][A-Z /&\-]{3,50})\n', _look_back_wide))
                if _caps_m:
                    _label = _caps_m[-1].group(1).strip().title()
                # Pass 1: quoted label
                _quoted = _rer.findall(r'[""]([^""]{3,60})[""]', _look_back)
                if _quoted:
                    _label = _quoted[-1].strip()
                # Pass 2: walk backward through sentence boundaries for the best label
                _SENT_STOPS = _rer.compile(
                    r'\s+(?:are\s+due|is\s+due|must\s+be|shall\s+be|will\s+be|'
                    r'at\s+that\s+time|hotel\s+will|group\s+must|no\s+later)',
                    _rer.IGNORECASE
                )
                if not _label:
                    _sent_bounds = [m.end() for m in
                                    _rer.finditer(r'(?:[.!?]\s+|\n\n|\n(?=[A-Z]))', _look_back)]
                    _sent_bounds.append(0)
                    _HAS_DATE = _rer.compile(_DT_PAT, _rer.IGNORECASE)
                    for _sb in reversed(_sent_bounds):
                        _sent_text = _look_back[_sb:].strip()
                        _sent_text = _rer.sub(r'\s*\n\s*', ' ', _sent_text).strip()
                        _stop = _SENT_STOPS.search(_sent_text)
                        _candidate = _sent_text[:_stop.start()].strip() if _stop else _sent_text[:100]
                        if (8 < len(_candidate) < 100
                                and not _candidate.endswith(',')
                                and _candidate[0].isupper()
                                and not _HAS_DATE.search(_candidate)):
                            _label = _candidate
                            break
                # Pass 3: last title-like line
                _TITLE_STOPS = _rer.compile(
                    r'\s+(?:at\s+that\s+time|hotel\s+will|hotel\s+shall|group\s+must|'
                    r'group\s+shall|no\s+later|the\s+hotel|all\s+room|prior\s+to)',
                    _rer.IGNORECASE
                )
                _SKIP_LINES = {'timeline action', 'action', 'timeline', 'event', 'description', 'note', 'dates'}
                if not _label:
                    _lines = [ln.strip() for ln in _look_back.split('\n') if ln.strip()]
                    for _ln in reversed(_lines[-6:]):
                        _ll = _ln.lower()
                        if (_ll in _SKIP_LINES or _ll.startswith('the ') or
                                _ll.startswith('if ') or _ll.startswith('in order') or
                                _ll.startswith('should ') or _ll.startswith('upon ')):
                            continue
                        _stop = _TITLE_STOPS.search(_ln)
                        _candidate = _ln[:_stop.start()].strip() if _stop else _ln.strip()
                        if (4 < len(_candidate) < 70
                                and not _candidate.endswith('.')
                                and not _candidate.endswith(',')
                                and _candidate[0].isupper()):
                            _label = _candidate
                            break
                _wider_ctx = raw_text[max(0, _rm.start()-400):_rm.start()+200]
                _is_cancel_ctx = bool(_rer.search(
                    r'cancell|liquidated\s+damage|%\s+of\s+(total\s+)?room\s+revenue',
                    _wider_ctx, _rer.IGNORECASE))
                if _is_cancel_ctx:
                    _after_match = raw_text[_rm.end():_rm.end()+300]
                    _same_line_end = _after_match.find('\n')
                    _row_text = _after_match[:_same_line_end] if _same_line_end >= 0 else _after_match[:150]
                    _pct_m = _rer.search(r'(\d{2,3})\s*%\s+of\s+(?:total\s+)?room\s+revenue',
                                         _row_text, _rer.IGNORECASE)
                    if not _pct_m:
                        _pct_m = _rer.search(r'(\d{2,3})\s*%', _row_text)
                    _pct_str = f' ({_pct_m.group(1)}% of room revenue)' if _pct_m else ''
                    _fab_ctx = _after_match[:300]
                    _fab_m = _rer.search(
                        r'(?:plus|and|\+)\s*(?:(\d{2,3})\s*%\s+of\s+)?'
                        r'(?:f(?:ood)?\s*(?:&|and)\s*b(?:everage)?|f\s*&\s*b)\s*'
                        r'(?:minimum|min\.?|commitment|revenue)?',
                        _fab_ctx, _rer.IGNORECASE
                    )
                    if _fab_m:
                        _fab_pct = _fab_m.group(1)
                        _fab_str = f' + {_fab_pct}% of F&B' if _fab_pct else ' + F&B minimum'
                    else:
                        _fab_str = ''
                    _label = f'Cancellation penalty{_pct_str}{_fab_str}'
                if not _label:
                    _label = f'{_n_days} days prior to event'
                _label_full = f'{_label}  ({_n_days} days prior to event)'
                _immediate_context = raw_text[max(0, _rm.start()-80):_rm.start()]
                if re.search(_DT_PAT, _immediate_context, re.IGNORECASE):
                    continue   # explicit date right before "N days prior" — already captured
                from datetime import date as _today_date
                _today_d = _today_date.today()
                if _calc_date >= _event_start_dt:
                    continue
                if (_today_d - _calc_date).days > 1825:
                    continue
                _key = (_calc_iso, 'rel:' + str(_n_days))
                if _key not in _seen_dates:
                    _seen_dates.add(_key)
                    _direct_critical.append({'date': _calc_iso, 'label': _label_full, 'amount': None})

    # ── 9. Deposit due at signing (no calendar date) ──────────────────────────
    # If a deposit is mentioned without a specific date ("at the time of signing",
    # "upon execution"), record it with date = "At signing".
    _signing_deposit_re = _rer.compile(
        r'(?:'
        # "$X deposit ... at signing"
        r'\$\s*([\d,]+(?:\.\d{2})?)[^.]{0,150}?deposit[^.]{0,150}?(?:at\s+the\s+time|upon|at\s+signing|at\s+execution|when\s+signed)'
        r'|'
        # "deposit ... $X ... at signing"  or  "deposit in the amount of $X ... at/upon signing"
        r'deposit[\s\S]{0,150}?\$\s*([\d,]+(?:\.\d{2})?)[\s\S]{0,150}?(?:at\s+the\s+time|upon|at\s+signing|at\s+execution|when\s+signed|is\s+due)'
        r')',
        _rer.IGNORECASE | _rer.DOTALL
    )
    for _m in _signing_deposit_re.finditer(raw_text):
        _raw_amt = _m.group(1) or (_m.group(2) if _m.lastindex >= 2 else None) or ''
        _amt = None
        try: _amt = float(_raw_amt.replace(',', ''))
        except Exception: pass
        if not _amt or _amt > 500000: continue   # sanity check
        _key = ('signing', 'Deposit due at sign')
        if _key not in _seen_dates:
            _seen_dates.add(_key)
            _direct_critical.append({'date': 'At signing', 'label': 'Deposit due at contract signing', 'amount': _amt})

    # ── Normalise critical_dates from AI + merge with direct results ──────────
    _raw_cd = ai_data.get('critical_dates') or []
    if not isinstance(_raw_cd, list):
        _raw_cd = []
    critical_dates = []
    for _cd in _raw_cd:
        if isinstance(_cd, dict) and _cd.get('date') and _cd.get('label'):
            _iso = _parse_any_date(str(_cd['date']))
            if not _iso:
                continue
            _key = (_iso, str(_cd['label'])[:20])
            if _key not in _seen_dates:
                _seen_dates.add(_key)
                critical_dates.append({
                    'date':   _iso,
                    'label':  str(_cd['label']),
                    'amount': _cd.get('amount'),
                })
    # Merge direct-parsed results (direct takes priority — added first)
    critical_dates = _direct_critical + critical_dates
    # Sort: "At signing" rows first, then chronologically by date
    critical_dates.sort(key=lambda x: ('0' if x['date'] == 'At signing' else '1') + x['date'])

    first = years[0]
    return {
        'contracted_block':     first.get('contracted_block', {}),
        'contracted_rate':      first.get('contracted_rate'),
        'rebate_per_room':      first.get('rebate_per_room'),
        'cutoff_date':          first.get('cutoff_date'),
        'block_review_date':    first.get('block_review_date'),
        'attrition_pct':        first.get('attrition_pct'),
        'event_name':           None,
        'hotel':                ai_data.get('hotel') or None,
        'organization':         ai_data.get('organization') or None,
        'hotel_contact':        ai_data.get('hotel_contact') or None,
        'hotel_contact_email':  ai_data.get('hotel_contact_email') or None,
        'group_contact':        ai_data.get('group_contact') or None,
        'group_contact_email':  ai_data.get('group_contact_email') or None,
        'critical_dates':       critical_dates,
        'raw_text':             raw_text,
        'years':                years,
    }


# ── Amendment / Addendum parsing ──────────────────────────────────────────────

def _ai_parse_amendment(text):
    """Extract structured data from a contract amendment/addendum (text-based)."""
    api_key = ''
    try:
        import importlib, sys
        if 'config' in sys.modules:
            importlib.reload(sys.modules['config'])
        else:
            import config as _cfg
            sys.modules['config'] = _cfg
        api_key = sys.modules['config'].ANTHROPIC_API_KEY.strip()
    except Exception:
        pass
    if not api_key:
        import os
        api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
    if not api_key:
        return {'error': 'Anthropic API key not configured'}

    try:
        import anthropic
    except ImportError:
        return {'error': 'anthropic package not installed'}

    prompt = """You are extracting structured data from a hotel group contract AMENDMENT or ADDENDUM.
An amendment changes specific terms of an existing contract — it does NOT define a full new contract.

Return ONLY a valid JSON object — no markdown fences, no explanation, nothing else.

The JSON object must have exactly these keys:
  description      — a short plain-English summary of what this amendment changes (1-2 sentences), required
  contracted_block — an object mapping date (YYYY-MM-DD) to rooms (integer) for ONLY the nights
                     that are being changed by this amendment. Nights with 0 rooms mean that night
                     is being REMOVED from the block. Null if the block is not changed.
  date_shift       — use this INSTEAD of contracted_block when the amendment shifts the entire
                     date range without providing a new per-night room count table. Format:
                     {"old_start":"YYYY-MM-DD","old_end":"YYYY-MM-DD","new_start":"YYYY-MM-DD","new_end":"YYYY-MM-DD"}
                     Set to null if contracted_block is populated or dates are not shifting.
  contracted_rate  — the new nightly room rate as a number (e.g. 219.00), or null if rate is not changing
  hotel            — the new hotel name if it has changed (e.g. rebrand/rename), or null if not changing
  cutoff_date      — the new cutoff/release date in YYYY-MM-DD format, or null if not changing

Rules:
- ONLY populate a field if the amendment explicitly changes that value.
- If a field is not mentioned or not changing, return null for it — do NOT copy values from elsewhere.
- For contracted_block: only include nights that the amendment adds, increases, decreases, or removes.
  A night with 0 rooms means it is being removed from the block.
- If the amendment shows a flat block change (e.g. "block increased to 50 rooms per night"), include
  all affected nights using the date range stated.
- Use date_shift (not contracted_block) when the amendment says something like "arrival changes
  from [date] to [date]" or "dates are moved from [range] to [range]" with no new per-night
  room count table. The room counts stay the same, only the dates change.
- Never populate both contracted_block and date_shift — use one or the other.
- The description field is always required — summarize what this amendment does.

Contract amendment text:
""" + text[:15000]

    try:
        import anthropic as _ant
        client = _ant.Anthropic(api_key=api_key)
        message = client.messages.create(
            model='claude-haiku-4-5',
            max_tokens=2048,
            messages=[{'role': 'user', 'content': prompt}]
        )
        raw = message.content[0].text.strip()
        raw = re.sub(r'^```[a-z]*\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        data = json.loads(raw)

        # Normalise contracted_block
        block_raw = data.get('contracted_block') or {}
        if block_raw:
            block = {}
            for k, v in block_raw.items():
                try:
                    datetime.strptime(str(k).strip(), '%Y-%m-%d')
                    block[str(k).strip()] = int(v)  # keep 0s — they mean "remove this night"
                except Exception:
                    continue
            data['contracted_block'] = block if block else None
        else:
            data['contracted_block'] = None

        # Normalise rate
        rate = data.get('contracted_rate')
        if rate is not None:
            try:
                data['contracted_rate'] = round(float(rate), 2)
            except Exception:
                data['contracted_rate'] = None

        data.setdefault('error', None)
        data.setdefault('description', '')
        return data

    except Exception as e:
        return {'error': str(e)}


def _ai_parse_amendment_vision(images, api_key):
    """Vision-based fallback for scanned amendment PDFs."""
    try:
        import anthropic
    except ImportError:
        return {'error': 'anthropic package not installed'}

    amendment_prompt = """You are extracting structured data from a scanned hotel contract AMENDMENT or ADDENDUM.
An amendment changes specific terms of an existing contract — it does NOT define a full new contract.

Return ONLY a valid JSON object — no markdown fences, no explanation, nothing else.

The JSON object must have exactly these keys:
  description      — a short plain-English summary of what this amendment changes (1-2 sentences), required
  contracted_block — an object mapping date (YYYY-MM-DD) to rooms (integer) for ONLY the nights
                     being changed. Nights with 0 rooms mean removal from block. Null if not changed.
  date_shift       — use this INSTEAD of contracted_block when the amendment shifts the entire
                     date range without a new per-night room count table. Format:
                     {"old_start":"YYYY-MM-DD","old_end":"YYYY-MM-DD","new_start":"YYYY-MM-DD","new_end":"YYYY-MM-DD"}
                     Null if contracted_block is populated or dates are not shifting.
  contracted_rate  — the new nightly room rate as a number, or null if rate is not changing
  hotel            — the new hotel name if it has changed (rebrand/rename), or null
  cutoff_date      — the new cutoff/release date in YYYY-MM-DD format, or null

Rules:
- ONLY populate a field if the amendment explicitly changes that value. Null = not changing.
- For contracted_block: only include nights being added, changed, or removed.
- Use date_shift (not contracted_block) when the amendment says arrival/departure dates are
  moving from one range to another, with no new per-night room count table provided.
- Never populate both contracted_block and date_shift — use one or the other.
- The description field is always required."""

    content = [{'type': 'text', 'text': amendment_prompt}]
    for b64 in images:
        content.append({
            'type': 'image',
            'source': {'type': 'base64', 'media_type': 'image/jpeg', 'data': b64}
        })

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model='claude-opus-4-5',
            max_tokens=2048,
            messages=[{'role': 'user', 'content': content}]
        )
        raw = message.content[0].text.strip()
        raw = re.sub(r'^```[a-z]*\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        data = json.loads(raw)

        block_raw = data.get('contracted_block') or {}
        if block_raw:
            block = {}
            for k, v in block_raw.items():
                try:
                    datetime.strptime(str(k).strip(), '%Y-%m-%d')
                    block[str(k).strip()] = int(v)
                except Exception:
                    continue
            data['contracted_block'] = block if block else None
        else:
            data['contracted_block'] = None

        rate = data.get('contracted_rate')
        if rate is not None:
            try:
                data['contracted_rate'] = round(float(rate), 2)
            except Exception:
                data['contracted_rate'] = None

        data.setdefault('error', None)
        data.setdefault('description', '')
        return data

    except Exception as e:
        return {'error': str(e)}


def parse_amendment_document(file_bytes, filename=''):
    """
    Parse a hotel contract amendment PDF or Word file.
    Returns a dict with keys:
      description       — str summary of what changed
      contracted_block  — {YYYY-MM-DD: rooms, ...} (only changed nights) or None
      contracted_rate   — float or None
      hotel             — str or None
      cutoff_date       — 'YYYY-MM-DD' or None
      error             — str or None
    """
    empty = {
        'description': '', 'contracted_block': None, 'date_shift': None,
        'contracted_rate': None, 'hotel': None, 'cutoff_date': None, 'error': None,
    }

    text = _extract_text_from_contract(file_bytes, filename)
    fname = (filename or '').lower()

    # Treat as image-only if text is empty, too short, or contains no contract keywords
    # (some PDFs only extract e-signature boilerplate like Sertifi overlay text)
    # Use whole-word matching to avoid false positives like "millenniumhotels.com"
    # containing "hotel" as a substring.
    _CONTRACT_KEYWORDS = {'rate', 'room', 'block', 'arrival', 'departure', 'cutoff',
                          'night', 'hotel', 'group', 'suite', 'meeting', 'attrition',
                          'reservation', 'check', 'guest', 'contract', 'agreement'}
    import re as _re_kw
    text_lower = (text or '').lower()
    text_words = text_lower.split()
    has_keywords = any(_re_kw.search(r'\b' + kw + r'\b', text_lower) for kw in _CONTRACT_KEYWORDS)
    text_meaningful = text and len(text_words) >= 50 and has_keywords
    if not text_meaningful:
        if fname.endswith('.pdf'):
            api_key = ''
            try:
                import importlib, sys
                if 'config' in sys.modules:
                    importlib.reload(sys.modules['config'])
                else:
                    import config as _cfg
                    sys.modules['config'] = _cfg
                api_key = sys.modules['config'].ANTHROPIC_API_KEY.strip()
            except Exception:
                pass
            if not api_key:
                import os
                api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
            if not api_key:
                return {**empty, 'error': 'Anthropic API key not configured'}
            images = _pdf_pages_to_images(file_bytes, max_pages=10, dpi=100)
            if not images:
                return {**empty, 'error': 'Could not extract text or render pages from PDF.'}
            result = _ai_parse_amendment_vision(images, api_key)
            for k in empty:
                if k not in result:
                    result[k] = empty[k]
            return result
        return {**empty, 'error': 'Could not extract text from file.'}

    result = _ai_parse_amendment(text)
    for k in empty:
        if k not in result:
            result[k] = empty[k]
    return result


# ═════════════════════════════════════════════════════════════════════════════
# COST SAVINGS — proposal & contract document parsers
# ═════════════════════════════════════════════════════════════════════════════

def _cs_safe_float(val):
    try:
        if val is None or val == '':
            return None
        return float(str(val).replace(',', '').replace('$', '').strip())
    except Exception:
        return None


def _cs_safe_int(val):
    try:
        if val is None or val == '':
            return None
        return int(float(str(val).replace(',', '').replace('$', '').strip()))
    except Exception:
        return None


def _cs_extract_text(file_bytes, filename):
    import re as _re
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    raw_text = ''
    try:
        if ext == 'pdf':
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                raw_text = '\n'.join((page.extract_text() or '') for page in pdf.pages)
        elif ext in ('docx', 'doc'):
            try:
                import docx, io
                doc = docx.Document(io.BytesIO(file_bytes))
                raw_text = '\n'.join(p.text for p in doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        raw_text += '\n' + '\t'.join(c.text for c in row.cells)
            except ImportError:
                return None, False, 'python-docx not installed'
        else:
            return None, False, f'Unsupported file type .{ext}'
    except Exception as e:
        return None, False, f'Could not read file: {e}'
    _plain_words = set(_re.findall(r'\b[a-z]{3,}\b', (raw_text or '').lower()))
    KEYWORDS = {'rate', 'room', 'block', 'cutoff', 'attrition', 'hotel', 'group',
                'meeting', 'suite', 'agreement', 'proposal', 'rfp'}
    is_scanned = (ext == 'pdf'
                  and (not raw_text or len((raw_text or '').split()) < 50
                       or not (KEYWORDS & _plain_words)))
    return raw_text, is_scanned, None


def _cs_get_anthropic_key():
    import sys, importlib, os
    api_key = ''
    try:
        if 'config' in sys.modules:
            importlib.reload(sys.modules['config'])
        else:
            import config as _cfg
            sys.modules['config'] = _cfg
        api_key = sys.modules['config'].ANTHROPIC_API_KEY.strip()
    except Exception:
        pass
    if not api_key:
        api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
    return api_key


_CS_PROPOSAL_PROMPT = """You are extracting cost-savings data from a hotel SALES PROPOSAL document
(this is the hotel's initial offer, BEFORE contract negotiation).
Return ONLY a valid JSON object — no markdown, no explanation.

Required keys (use null if not found):
  "rack_rate"               — published / non-discounted nightly rack rate (number)
  "group_rate"              — proposed group / staff nightly rate (number)
  "meeting_room_rental"     — total proposed meeting room rental fee (number)
  "f_and_b_minimum"         — proposed F&B minimum (number)
  "attrition_pct"           — proposed attrition as decimal (80% -> 0.80)
  "comp_industry_standard"  — industry-standard comp room ratio denominator (usually 50)
  "comp_negotiated_policy"  — proposed comp room ratio denominator (e.g. 40)
  "internet_gr_price"       — published price per night for in-room internet
  "internet_ms_price"       — published price per day for meeting-space internet
  "valet_price"             — published nightly valet parking price
  "av_discount_pct"         — A/V discount offered as decimal (15% -> 0.15)
  "concessions"             — array of short strings describing each concession offered
  "hotel_brand"             — one of "Hyatt","Hilton","Marriott","IHG","Preferred"
  "total_room_nights"       — total contracted room nights across the pattern (integer)

Rules:
- Numbers only — strip "$" and "," from amounts.
- For percentages, return decimals (10% -> 0.10).
- If genuinely not found, use null (not 0).
"""

_CS_CONTRACT_PROMPT = """You are extracting cost-savings data from a SIGNED HOTEL CONTRACT
(this is the FINAL negotiated agreement).
Return ONLY a valid JSON object — no markdown, no explanation.

Required keys (use null if not found):
  "contracted_rate"          — final negotiated guest-room nightly rate (number)
  "staff_contracted_rate"    — final negotiated staff-room nightly rate (number)
  "total_room_nights"        — total contracted guest room nights (integer)
  "staff_total_nights"       — total contracted staff room nights (integer)
  "meeting_room_rental"      — final meeting room rental fee (number)
  "f_and_b_minimum"          — final F&B minimum (number)
  "attrition_pct"            — final attrition as decimal (70% -> 0.70)
  "comp_negotiated_policy"   — final comp ratio denominator (e.g. "1 per 40" -> 40)
  "internet_gr_price"        — final per-night in-room internet charge (number)
  "internet_ms_price"        — final per-day meeting-space internet charge (number)
  "valet_price"              — final nightly valet parking (number)
  "av_discount_pct"          — A/V discount as decimal
  "concessions"              — array of short strings — each negotiated concession

Rules:
- Numbers only — strip "$" and ",".
- Percentages as decimals.
- If not found, use null.
"""


def _cs_call_claude(prompt, raw_text, file_bytes, is_scanned):
    api_key = _cs_get_anthropic_key()
    if not api_key:
        return None, 'No Anthropic API key configured'
    try:
        import anthropic, json, re
        client = anthropic.Anthropic(api_key=api_key)
        if is_scanned:
            import fitz, base64
            doc = fitz.open(stream=file_bytes, filetype='pdf')
            content = [{'type': 'text', 'text': prompt}]
            for page_num in range(min(len(doc), 8)):
                page = doc[page_num]
                mat = fitz.Matrix(100 / 72, 100 / 72)
                pix = page.get_pixmap(matrix=mat)
                img_b64 = base64.standard_b64encode(pix.tobytes('png')).decode()
                content.append({
                    'type': 'image',
                    'source': {'type': 'base64', 'media_type': 'image/png', 'data': img_b64}
                })
            doc.close()
            content.append({'type': 'text', 'text': 'Extract the data and return ONLY the JSON.'})
            response = client.messages.create(
                model='claude-opus-4-5',
                max_tokens=2048,
                messages=[{'role': 'user', 'content': content}]
            )
        else:
            if not raw_text or not raw_text.strip():
                return None, 'No text could be extracted from the file'
            response = client.messages.create(
                model='claude-opus-4-5',
                max_tokens=2048,
                messages=[{'role': 'user',
                           'content': prompt + '\nDocument text:\n' + raw_text[:12000]}]
            )
        raw_json = response.content[0].text.strip()
        raw_json = re.sub(r'^```[a-z]*\n?', '', raw_json)
        raw_json = re.sub(r'\n?```$', '', raw_json)
        return json.loads(raw_json), None
    except Exception as e:
        return None, f'AI extraction failed: {e}'


def _cs_normalize(data):
    if not data:
        return {}
    out = {}
    for k, v in data.items():
        if v is None:
            out[k] = None
            continue
        if k in ('total_room_nights', 'staff_total_nights',
                 'comp_industry_standard', 'comp_negotiated_policy'):
            out[k] = _cs_safe_int(v)
        elif k == 'concessions':
            if isinstance(v, list):
                out[k] = [str(x) for x in v]
            else:
                out[k] = []
        elif k == 'hotel_brand':
            out[k] = str(v) if v else None
        else:
            out[k] = _cs_safe_float(v)
    return out


def parse_proposal_for_cost_savings(file_bytes, filename=''):
    raw_text, is_scanned, err = _cs_extract_text(file_bytes, filename)
    if err:
        return {'error': err, 'raw_text': ''}
    data, err = _cs_call_claude(_CS_PROPOSAL_PROMPT, raw_text, file_bytes, is_scanned)
    if err:
        return {'error': err, 'raw_text': raw_text or ''}
    normalized = _cs_normalize(data)
    normalized['raw_text'] = raw_text or ''
    return normalized


def parse_contract_for_cost_savings(file_bytes, filename=''):
    raw_text, is_scanned, err = _cs_extract_text(file_bytes, filename)
    if err:
        return {'error': err, 'raw_text': ''}
    data, err = _cs_call_claude(_CS_CONTRACT_PROMPT, raw_text, file_bytes, is_scanned)
    if err:
        return {'error': err, 'raw_text': raw_text or ''}
    normalized = _cs_normalize(data)
    normalized['raw_text'] = raw_text or ''
    return normalized


def parse_columnar_pickup_pdf(file_bytes, filename=''):
    """
    Parse a hotel "Group Pickup" PDF where dates are column headers and
    pickup counts are in rows beneath (Henry Hotel / IHG grppickup format).

    Returns:
      {
        'pairs':  [{'date': 'YYYY-MM-DD', 'count': N}, ...],   # all dates incl. zeros
        'text':   str,     # raw extracted text
        'source': str,     # 'Grand Total' or block name used
        'error':  str|None
      }
    """
    import re
    from datetime import date as _date

    # ── Extract text ───────────────────────────────────────────────────────────
    text = ''
    try:
        import pdfplumber, io as _io
        with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
            text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
    except Exception as e:
        return {'pairs': [], 'text': '', 'source': '', 'error': str(e)}

    if not text.strip():
        return {'pairs': [], 'text': text, 'source': '', 'error': 'No text extracted from PDF'}

    # ── "Group Pickup Detail" format (Opera/Hilton row-per-date report) ─────────
    if 'group pickup detail' in text.lower():
        _gpd_re = re.compile(
            r'^([A-Z]{1,5})\s+(\d{1,2}/\d{1,2}/\d{2,4})\s+'
            r'\d+\s+\d+\s+\d+\s+\d+\s+(\d+)',
            re.MULTILINE
        )
        date_totals = {}
        for m in _gpd_re.finditer(text):
            raw_date = m.group(2)
            count    = int(m.group(3))
            parts = raw_date.split('/')
            if len(parts) == 3:
                mo, dy, yr = int(parts[0]), int(parts[1]), int(parts[2])
                if yr < 100:
                    yr += 2000
                iso = f'{yr}-{mo:02d}-{dy:02d}'
                date_totals[iso] = date_totals.get(iso, 0) + count
        if date_totals:
            pairs = [{'date': d, 'count': c} for d, c in sorted(date_totals.items())]
            return {'pairs': pairs, 'text': text, 'source': 'Group Pickup Detail', 'error': None}

    lines = [l.strip() for l in text.split('\n') if l.strip()]

    MONTHS = {'jan':1,'feb':2,'mar':3,'apr':4,'may':5,'jun':6,
              'jul':7,'aug':8,'sep':9,'oct':10,'nov':11,'dec':12}

    # ── Find the day-number header row ─────────────────────────────────────────
    day_row_idx = None
    day_numbers = []
    for i, line in enumerate(lines):
        tokens = line.split()
        nums = []
        for t in tokens:
            if re.match(r'^\d{1,2}$', t):
                n = int(t)
                if 1 <= n <= 31:
                    nums.append(n)
        if len(nums) >= 4 and len(nums) >= 0.7 * len(tokens):
            day_row_idx = i
            day_numbers = nums
            break

    if day_row_idx is None:
        return {'pairs': [], 'text': text, 'source': '',
                'error': 'Date header row not found — unrecognized pickup report format'}

    # ── Determine year and start month from surrounding context ────────────────
    year = None
    start_month = None

    for line in lines[:day_row_idx + 6]:
        m = re.search(r'(?:From Date|Start Date|Filter From Date)\s+(\d{1,2})[/\-]\d{1,2}[/\-](\d{2,4})',
                      line, re.IGNORECASE)
        if m:
            if start_month is None:
                start_month = int(m.group(1))
            yr_raw = int(m.group(2))
            if year is None:
                year = 2000 + yr_raw if yr_raw < 100 else yr_raw
            break

    if year is None:
        for line in lines[:10]:
            m = re.search(r'\b(20\d{2})\b', line)
            if m:
                year = int(m.group(1))
                break
    if year is None:
        for line in lines:
            m = re.search(r'\b\d{1,2}[/\-]\d{1,2}[/\-](\d{2})\b', line)
            if m:
                year = 2000 + int(m.group(1))
                break
    if year is None:
        year = _date.today().year

    if start_month is None:
        for line in lines[max(0, day_row_idx - 3):day_row_idx + 1]:
            for token in line.split():
                tok_l = token[:3].lower()
                if tok_l in MONTHS:
                    start_month = MONTHS[tok_l]
                    break
            if start_month:
                break
    if start_month is None:
        start_month = _date.today().month

    # ── Build date list, handling month roll-over ──────────────────────────────
    dates = []
    cur_month = start_month
    prev_day = 0
    cur_year = year
    for day in day_numbers:
        if day < prev_day and day <= 7:
            cur_month += 1
            if cur_month > 12:
                cur_month = 1
                cur_year += 1
        dates.append(f'{cur_year}-{cur_month:02d}-{day:02d}')
        prev_day = day

    # ── Find pickup values: prefer Grand Total Pickup ──────────────────────────
    pickup_values = None
    source = ''

    grand_total_idx = None
    for i, line in enumerate(lines):
        if 'grand total' in line.lower():
            grand_total_idx = i
            break

    def _extract_pickup_nums(line, expected_len):
        nums = [int(x) for x in re.findall(r'\b(\d+)\b', line)]
        if len(nums) == expected_len + 1:
            nums = nums[:-1]
        return nums if len(nums) == expected_len else None

    if grand_total_idx is not None:
        for line in lines[grand_total_idx: grand_total_idx + 6]:
            if re.match(r'.*\bpickup\b', line, re.IGNORECASE):
                nums = _extract_pickup_nums(line, len(dates))
                if nums is not None:
                    pickup_values = nums
                    source = 'Grand Total'
                    break

    if pickup_values is None:
        for line in lines:
            if re.match(r'.*\bpickup\b', line, re.IGNORECASE):
                nums = _extract_pickup_nums(line, len(dates))
                if nums is not None:
                    pickup_values = nums
                    source = 'Pickup'
                    break

    if pickup_values is None:
        return {'pairs': [], 'text': text, 'source': '',
                'error': 'Could not find a Pickup row matching the date columns'}

    pairs = [{'date': d, 'count': c} for d, c in zip(dates, pickup_values)]
    return {'pairs': pairs, 'text': text, 'source': source, 'error': None}


def parse_columnar_pickup_xlsx(file_bytes, filename=''):
    """
    Parse a hotel "Group Pickup" XLSX.  Handles two layouts:

    Format A — Omni/IHG (Pick-up sheet):
      Row 3: date headers as "MM/DD\\nDay"
      Row labeled "Pickup Total": pickup counts

    Format B — Hyatt "Block and PickUp Report":
      Row 7: date headers as ISO strings "YYYY-MM-DD"
      Rows labeled "Sold" under "GRAND TOTAL" section: pickup counts

    Returns: {'pairs': [{'date','count'}], 'text':'', 'source': str, 'error': str|None}
    """
    import re, io as _io
    try:
        import openpyxl
        wb = openpyxl.load_workbook(_io.BytesIO(file_bytes), data_only=True)
    except Exception as e:
        return {'pairs': [], 'text': '', 'source': '', 'error': str(e)}

    ws = None
    for sn in wb.sheetnames:
        sn_l = sn.lower()
        if 'pick' in sn_l and 'room' not in sn_l:
            ws = wb[sn]
            break
    if ws is None:
        ws = wb.active

    all_rows = list(ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True))

    header_row_idx = None
    date_col_map = {}
    fmt = None

    year_a = None
    for r in range(min(4, len(all_rows))):
        v = ws.cell(r + 1, 1).value
        if v is None:
            continue
        if hasattr(v, 'year'):
            year_a = v.year; break
        m = re.search(r'\b(20\d{2})\b', str(v))
        if m:
            year_a = int(m.group(1)); break
    if year_a is None:
        from datetime import date as _d
        year_a = _d.today().year

    for row_idx, row in enumerate(all_rows[:15]):
        hits = sum(1 for c in row if c and re.match(r'^\d{1,2}/\d{1,2}', str(c).strip()))
        if hits >= 3:
            for col_idx, cell in enumerate(row):
                if not cell:
                    continue
                m = re.match(r'^(\d{1,2})/(\d{1,2})', str(cell).strip())
                if m:
                    mo, dy = int(m.group(1)), int(m.group(2))
                    date_col_map[col_idx] = f'{year_a}-{mo:02d}-{dy:02d}'
            if date_col_map:
                header_row_idx = row_idx
                fmt = 'A'
                break

    if not date_col_map:
        for row_idx, row in enumerate(all_rows[:15]):
            hits = 0
            for cell in row:
                if cell and re.match(r'^20\d{2}-\d{2}-\d{2}$', str(cell).strip()):
                    hits += 1
            if hits >= 3:
                for col_idx, cell in enumerate(row):
                    if cell and re.match(r'^20\d{2}-\d{2}-\d{2}$', str(cell).strip()):
                        date_col_map[col_idx] = str(cell).strip()
                if date_col_map:
                    header_row_idx = row_idx
                    fmt = 'B'
                    break

    if not date_col_map:
        return {'pairs': [], 'text': '', 'source': '',
                'error': 'No date header row found (tried MM/DD and YYYY-MM-DD formats)'}

    def _extract_row_vals(row):
        vals = {}
        for col_idx, cell in enumerate(row):
            if col_idx in date_col_map:
                try:
                    vals[date_col_map[col_idx]] = int(float(str(cell or 0)))
                except Exception:
                    vals[date_col_map[col_idx]] = 0
        return vals

    pickup_vals = None
    source = ''

    if fmt == 'A':
        for row in all_rows[header_row_idx + 1:]:
            label = (str(row[0] or '') + str(row[1] or '')).strip().lower()
            if 'pickup' in label:
                vals = _extract_row_vals(row)
                if vals:
                    pickup_vals = vals
                    source = str(row[0] or row[1] or 'Pickup').strip()
                    break

    elif fmt == 'B':
        grand_total_row = None
        for row_idx, row in enumerate(all_rows[header_row_idx:], start=header_row_idx):
            label = str(row[0] or '').strip().upper()
            if 'GRAND TOTAL' in label:
                grand_total_row = row_idx
                break

        search_start = grand_total_row if grand_total_row is not None else header_row_idx + 1
        for row in all_rows[search_start:search_start + 10]:
            label = str(row[1] or row[0] or '').strip().lower()
            if label == 'sold':
                vals = _extract_row_vals(row)
                if vals:
                    pickup_vals = vals
                    source = 'Grand Total — Sold'
                    break

        if pickup_vals is None:
            for row in all_rows[header_row_idx + 1:]:
                label = str(row[1] or row[0] or '').strip().lower()
                if label == 'sold':
                    vals = _extract_row_vals(row)
                    if vals:
                        pickup_vals = vals
                        source = str(row[0] or 'Sold').strip() + ' — Sold'
                        break

    if pickup_vals is None:
        return {'pairs': [], 'text': '', 'source': '',
                'error': f'No pickup row found (format {fmt})'}

    pairs = [{'date': d, 'count': c} for d, c in sorted(pickup_vals.items())]
    return {'pairs': pairs, 'text': '', 'source': source, 'error': None}


def _hhr_pdf_date_to_iso(date_str, year):
    """Convert '5-Jun' → '2026-06-05' given year int."""
    from datetime import datetime as _dt
    try:
        return _dt.strptime(f"{date_str.strip()}-{year}", '%d-%b-%Y').strftime('%Y-%m-%d')
    except ValueError:
        return None


def _hhr_pdf_parse_pct(val):
    """'10%' → 0.10"""
    import re as _re
    if not val:
        return None
    m = _re.search(r'(\d+(?:\.\d+)?)\s*%', str(val))
    return float(m.group(1)) / 100.0 if m else None


def _hhr_pdf_parse_currency(val):
    """'($ 269.00)' or '$269' → 269.0"""
    import re as _re
    if not val:
        return None
    cleaned = str(val).replace('(', '').replace(')', '').replace(',', '').replace('$', '').strip()
    m = _re.search(r'[\d]+\.?\d*', cleaned)
    return float(m.group()) if m else None


def parse_hhr_pdf(file_bytes, fallback_year=None):
    """Parse a completed Housing History Report PDF returned by the hotel.

    Returns a dict with:
      organization, hotel, event_name, booking_id, commission_pct,
      contracted_total, contracted_block (dict), contracted_rate,
      final_total_pickup, pickup_by_night (dict iso→int),
      room_revenue, audit_pickup, no_shows, cancellations,
      hotel_approver, hotel_approver_email, error
    """
    import io, re

    result = {
        'organization':         '',
        'hotel':                '',
        'event_name':           '',
        'booking_id':           '',
        'commission_pct':       None,
        'contracted_total':     0,
        'contracted_block':     {},
        'contracted_rate':      None,
        'final_total_pickup':   0,
        'pickup_by_night':      {},
        'room_revenue':         None,
        'audit_pickup':         0,
        'no_shows':             0,
        'cancellations':        0,
        'hotel_approver':       '',
        'hotel_approver_email': '',
        'error':                None,
    }

    try:
        import pdfplumber
    except ImportError:
        result['error'] = 'pdfplumber not installed'
        return result

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                result['error'] = 'No pages in PDF'
                return result

            page = pdf.pages[0]
            tables = page.extract_tables()
            if not tables:
                result['error'] = 'No tables found on page 1'
                return result

            table = tables[0]
            year = fallback_year
            date_cols = []

            for row in table:
                if not row:
                    continue
                label = (row[0] or '').strip()
                if label == 'NAME & DATE OF EVENT' and len(row) > 1:
                    val = (row[1] or '').strip()
                    result['event_name'] = val
                    m = re.search(r'(20\d{2})', val)
                    if m:
                        year = int(m.group(1))
                    break

            for row in table:
                if not row:
                    continue
                label = (row[0] or '').strip()

                if label == '' and len(row) > 8:
                    for ci, cell in enumerate(row):
                        if (cell or '').strip() == 'Comm %' and ci + 2 < len(row):
                            result['commission_pct'] = _hhr_pdf_parse_pct(row[ci + 2])
                        if (cell or '').strip() == 'Booking #:' and ci + 2 < len(row):
                            result['booking_id'] = (row[ci + 2] or '').strip()

                elif label == 'ORGANIZATION' and len(row) > 1:
                    result['organization'] = (row[1] or '').strip()

                elif label == 'HOTEL' and len(row) > 1:
                    result['hotel'] = (row[1] or '').strip()

                elif label == 'DATE':
                    date_pat = re.compile(
                        r'^\d{1,2}-(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)$', re.I
                    )
                    for ci, cell in enumerate(row[1:], start=1):
                        if cell and date_pat.match(cell.strip()):
                            date_cols.append((ci, cell.strip()))

                elif label == 'Contracted Block' and date_cols and year:
                    for ci, date_str in date_cols:
                        val_str = (row[ci] if ci < len(row) else None) or ''
                        try:
                            n = int(val_str.strip()) if val_str.strip() else 0
                        except ValueError:
                            n = 0
                        iso = _hhr_pdf_date_to_iso(date_str, year)
                        if iso and n > 0:
                            result['contracted_block'][iso] = n
                    for cell in reversed(row):
                        rate = _hhr_pdf_parse_currency(cell)
                        if rate:
                            result['contracted_rate'] = rate
                            break
                    result['contracted_total'] = sum(result['contracted_block'].values())

                elif 'FINAL TOTAL PICKUP' in label.upper() and date_cols and year:
                    for ci, date_str in date_cols:
                        val_str = (row[ci] if ci < len(row) else None) or ''
                        try:
                            n = int(val_str.strip()) if val_str.strip() else 0
                        except ValueError:
                            n = 0
                        iso = _hhr_pdf_date_to_iso(date_str, year)
                        if iso:
                            result['pickup_by_night'][iso] = n
                    total_ci = date_cols[-1][0] + 1
                    if total_ci < len(row):
                        try:
                            result['final_total_pickup'] = int((row[total_ci] or '').strip() or '0')
                        except ValueError:
                            result['final_total_pickup'] = sum(result['pickup_by_night'].values())
                    else:
                        result['final_total_pickup'] = sum(result['pickup_by_night'].values())

                elif 'Non-Commissionable Audit' in label or 'Audit Pickup' in label.lower():
                    total_ci = date_cols[-1][0] + 1 if date_cols else None
                    if total_ci and total_ci < len(row):
                        try:
                            result['audit_pickup'] = int((row[total_ci] or '').strip() or '0')
                        except ValueError:
                            pass

                elif 'Commissionable No Shows' in label:
                    total_ci = date_cols[-1][0] + 1 if date_cols else None
                    if total_ci and total_ci < len(row):
                        try:
                            result['no_shows'] = int((row[total_ci] or '').strip() or '0')
                        except ValueError:
                            pass

                elif 'Commissionable Cancellations' in label:
                    total_ci = date_cols[-1][0] + 1 if date_cols else None
                    if total_ci and total_ci < len(row):
                        try:
                            result['cancellations'] = int((row[total_ci] or '').strip() or '0')
                        except ValueError:
                            pass

                elif 'History Report Approved By' in label:
                    name_next = False
                    email_next = False
                    for cell in row[1:]:
                        if cell is None:
                            continue
                        cell_s = cell.strip()
                        if name_next and cell_s:
                            result['hotel_approver'] = cell_s
                            name_next = False
                        elif email_next and cell_s:
                            result['hotel_approver_email'] = cell_s
                            email_next = False
                        elif cell_s == 'Name':
                            name_next = True
                        elif cell_s == 'Email':
                            email_next = True

            if not result['hotel_approver'] and len(pdf.pages) > 1:
                for pg in pdf.pages[1:]:
                    pg_tables = pg.extract_tables()
                    for pg_table in pg_tables:
                        for row in pg_table:
                            if not row:
                                continue
                            label = (row[0] or '').strip()
                            if 'History Report Approved By' in label:
                                name_next = False
                                email_next = False
                                for cell in row[1:]:
                                    if cell is None:
                                        continue
                                    cell_s = cell.strip()
                                    if name_next and cell_s:
                                        result['hotel_approver'] = cell_s
                                        name_next = False
                                    elif email_next and cell_s:
                                        result['hotel_approver_email'] = cell_s
                                        email_next = False
                                    elif cell_s == 'Name':
                                        name_next = True
                                    elif cell_s == 'Email':
                                        email_next = True
                                if result['hotel_approver']:
                                    break
                        if result['hotel_approver']:
                            break

    except Exception as e:
        result['error'] = str(e)

    return result
